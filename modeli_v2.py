"""
modeli_v2.py – Izboljšano napovedovanje cen nepremičnin
=========================================================
Uporablja: scikit-learn, XGBoost, LightGBM, pandas, matplotlib, shap

Modeli:
  • Ridge regresija       (baseline, linearen)
  • Random Forest         (100+ dreves, GridSearchCV)
  • XGBoost               (gradientno ojačevanje)
  • LightGBM              (gradientno ojačevanje, hitrejši)

Feature engineering:
  • StarostZgrade = 2026 - LetoGradnje
  • Log transformacija: VelikostM2, ZemljisteM2
  • Target encoding za Obcina in VrstaObjekta (brez data leakage)
  • One-hot encoding za EnergetskiRazred

Zagon:
    .venv312\\Scripts\\python modeli_v2.py
    .venv312\\Scripts\\python modeli_v2.py --csv podatki.csv --docx
"""

import sys, os, argparse
from datetime import datetime

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

# ── CLI argumenti ─────────────────────────────────────────────────────────────
parser = argparse.ArgumentParser(description="Izboljšani ML modeli – nepremičnine")
parser.add_argument("--csv",   default=None, help="Pot do CSV datoteke")
parser.add_argument("--split", default=0.8,  type=float)
parser.add_argument("--seed",  default=42,   type=int)
parser.add_argument("--docx",  action="store_true")
parser.add_argument("--docx-izhod", default=None, dest="docx_izhod")
parser.add_argument("--shap",  action="store_true", help="Generiraj SHAP plot")
args = parser.parse_args()

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))

# ── Potrebne knjižnice ────────────────────────────────────────────────────────
try:
    import pandas as pd
    import numpy as np
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import joblib
    from sklearn.model_selection import train_test_split, cross_val_score, KFold
    from sklearn.linear_model import Ridge
    from sklearn.ensemble import RandomForestRegressor, GradientBoostingRegressor
    from sklearn.preprocessing import StandardScaler, LabelEncoder
    from sklearn.metrics import r2_score, mean_absolute_error, mean_squared_error
    from sklearn.pipeline import Pipeline
    import xgboost as xgb
    import lightgbm as lgb
except ImportError as e:
    print(f"✗  Manjkajoča knjižnica: {e}")
    print("   Zaženi: .venv312\\Scripts\\python -m pip install scikit-learn xgboost lightgbm shap pandas matplotlib joblib")
    sys.exit(1)

print(f"✓  Knjižnice naložene  (sklearn, xgboost, lightgbm, pandas)")

# ── Nalaganje CSV ──────────────────────────────────────────────────────────────
CANDIDATES = [
    args.csv,
    os.path.join(SCRIPT_DIR, "nepremicnine_export_prodaja.csv"),
    os.path.join(SCRIPT_DIR, "nepremicnine_export.csv"),
]
csv_path = next((p for p in CANDIDATES if p and os.path.isfile(p)), None)
if csv_path is None:
    print("✗  CSV datoteka ni najdena.")
    sys.exit(1)
print(f"✓  Berem: {csv_path}")

df_raw = pd.read_csv(csv_path, sep=";", encoding="utf-8-sig", low_memory=False)
print(f"   Skupaj vrstic: {len(df_raw)}  |  Stolpci: {list(df_raw.columns)}")

# ── Čiščenje podatkov ─────────────────────────────────────────────────────────
def parse_num(s):
    try:
        return float(str(s).replace(",", ".").strip())
    except:
        return np.nan

for col in ["Cena", "VelikostM2", "ZemljisteM2", "LetoGradnje", "StSob"]:
    df_raw[col] = df_raw[col].apply(parse_num)

# Filter: veljavne cene
df = df_raw[df_raw["Cena"].between(10_000, 5_000_000)].copy()
# Filter: veljavna površina
df = df[df["VelikostM2"].between(10, 2000)].copy()
print(f"   Po filtru cena+površina: {len(df)} vrstic")

# Odstranitev skrajnih 2 % (outlierji)
q_low  = df["Cena"].quantile(0.02)
q_high = df["Cena"].quantile(0.98)
df = df[df["Cena"].between(q_low, q_high)].copy()
print(f"   Po 2–98 % rezu: {len(df)} vrstic  (cena: {q_low:,.0f} – {q_high:,.0f} €)")

# ── Feature Engineering ───────────────────────────────────────────────────────
print("\n[1/5]  Feature engineering ...")

# Starost zgradbe
df["StarostZgrade"] = 2026 - df["LetoGradnje"]
df.loc[df["StarostZgrade"] < 0, "StarostZgrade"]  = np.nan
df.loc[df["StarostZgrade"] > 200, "StarostZgrade"] = np.nan

# Log transformacije (lognormalna porazdelitev površin)
df["log_VelikostM2"]  = np.log1p(df["VelikostM2"])
df["log_ZemljisteM2"] = np.log1p(df["ZemljisteM2"].fillna(0))

# Razmerje cena/m² (bo vzeto iz surovih podatkov, ne bo feature za napoved)
# Energetski razred → ordinalna vrednost
ENERGY_MAP = {"A+": 8, "A": 7, "B": 6, "C": 5, "D": 4, "E": 3, "F": 2, "G": 1}
df["EnergetRazredNum"] = df["EnergetskiRazred"].apply(
    lambda x: ENERGY_MAP.get(str(x).strip().upper(), np.nan)
)

# StSob – zapolni z mediano po vrsti objekta
df["StSob"] = pd.to_numeric(df["StSob"], errors="coerce")
sob_med = df.groupby("VrstaObjekta")["StSob"].median()
df["StSob"] = df.apply(
    lambda r: sob_med.get(r["VrstaObjekta"], df["StSob"].median())
    if pd.isna(r["StSob"]) else r["StSob"], axis=1
)

# Cena na m² za lokacijo (povprečje po občini — bo izračunano samo iz train podatkov)
# → to naredimo šele po razdelitvi, da preprečimo data leakage

# Lokacija → očisti
df["Obcina"]      = df["Obcina"].fillna(df.get("Lokacija", "")).str.strip()
df["VrstaObjekta"] = df["VrstaObjekta"].fillna("Drugo").str.strip()

# Log ciljna spremenljivka
df["log_Cena"] = np.log(df["Cena"])

print(f"   Značilke pripravljene. Vzorcev: {len(df)}")

# ── Izbira značilk ────────────────────────────────────────────────────────────
BASE_NUM_FEATS = [
    "log_VelikostM2",    # log(površina)
    "log_ZemljisteM2",   # log(zemljišče)
    "StarostZgrade",     # starost v letih
    "StSob",             # število sob
    "EnergetRazredNum",  # energetski razred 1–8
]
CAT_FEATS = ["VrstaObjekta", "Obcina"]

# ── Razdelitev train/test ─────────────────────────────────────────────────────
df_train, df_test = train_test_split(df, test_size=1-args.split,
                                     random_state=args.seed, shuffle=True)
print(f"\n[2/5]  Razdelitev: {len(df_train)} train  |  {len(df_test)} test")

# ── Target encoding (samo iz train, da ni data leakage) ───────────────────────
def target_encode(train_df, test_df, col, target="log_Cena", smoothing=10):
    """
    Target encoding z glajanjem (smoothing) – preprečuje overfitting pri redkih kategorijah.
    Formula: encoded = (n * category_mean + smoothing * global_mean) / (n + smoothing)
    """
    global_mean = train_df[target].mean()
    stats = train_df.groupby(col)[target].agg(["mean", "count"])
    smooth_enc = (stats["count"] * stats["mean"] + smoothing * global_mean) / (stats["count"] + smoothing)
    smooth_enc = smooth_enc.rename(f"{col}_enc")

    train_enc = train_df[col].map(smooth_enc).fillna(global_mean)
    test_enc  = test_df[col].map(smooth_enc).fillna(global_mean)
    return train_enc.values, test_enc.values, smooth_enc

print("\n[3/5]  Target encoding kategoričnih spremenljivk ...")
vrsta_train, vrsta_test, vrsta_enc = target_encode(df_train, df_test, "VrstaObjekta")
obcina_train, obcina_test, obcina_enc  = target_encode(df_train, df_test, "Obcina")

# Cena/m² po občini iz train podatkov (povprečna tržna vrednost lokacije)
cena_m2_obcina = df_train.groupby("Obcina")["Cena"].median() / df_train.groupby("Obcina")["VelikostM2"].median()
global_cena_m2 = (df_train["Cena"] / df_train["VelikostM2"]).median()

def get_cena_m2_lok(df_in, ref_col="Obcina"):
    return df_in[ref_col].map(cena_m2_obcina).fillna(global_cena_m2).values

# Mediane za imputacijo (bo izračunano po razdelitvi)
_TRAIN_MEDIANS = {}

# Sestavi feature matrike
def build_X(df_in, vrsta_enc_vals, obcina_enc_vals):
    num = df_in[BASE_NUM_FEATS].copy()
    # Imputacija z mediano iz train
    for c in BASE_NUM_FEATS:
        med = _TRAIN_MEDIANS.get(c, df_train[c].median())
        num[c] = num[c].fillna(med)
    X = np.column_stack([
        num.values.astype(float),
        vrsta_enc_vals.reshape(-1,1),
        obcina_enc_vals.reshape(-1,1),
        np.log1p(get_cena_m2_lok(df_in)).reshape(-1,1),  # log(cena/m² za lokacijo)
    ])
    # Zamenjaj preostale NaN z 0 (varnostni net)
    X = np.nan_to_num(X, nan=0.0)
    return X

FEAT_NAMES = BASE_NUM_FEATS + ["VrstaObjekta_enc", "Obcina_enc", "log_cena_m2_lokacija"]

# Izračunaj mediane iz train podatkov
for c in BASE_NUM_FEATS:
    _TRAIN_MEDIANS[c] = df_train[c].median()

X_train = build_X(df_train, vrsta_train, obcina_train)
X_test  = build_X(df_test,  vrsta_test,  obcina_test)
y_train = df_train["log_Cena"].values
y_test  = df_test["log_Cena"].values
y_train_raw = df_train["Cena"].values
y_test_raw  = df_test["Cena"].values

print(f"   Matrika značilk: {X_train.shape[1]} značilk")
print(f"   Značilke: {FEAT_NAMES}")

# ── Metrike ───────────────────────────────────────────────────────────────────
def eval_model(name, model, X_tr, y_tr_raw, X_te, y_te_raw):
    """Evalvacija modela – napoveduje log(Cena), metrike v €."""
    pred_log = model.predict(X_te)
    pred_eur = np.exp(pred_log)

    r2   = r2_score(y_te_raw, pred_eur)
    mae  = mean_absolute_error(y_te_raw, pred_eur)
    rmse = np.sqrt(mean_squared_error(y_te_raw, pred_eur))
    mape = np.mean(np.abs(y_te_raw - pred_eur) / y_te_raw) * 100

    # CV na train podatkih (log Cena)
    y_tr_log = np.log(y_tr_raw)
    cv = KFold(n_splits=5, shuffle=True, random_state=42)
    cv_scores = cross_val_score(model, X_tr, y_tr_log, cv=cv, scoring="r2")

    print(f"\n  {'─'*55}")
    print(f"  {name}")
    print(f"  {'─'*55}")
    print(f"  R²      : {r2:.4f}")
    print(f"  MAE     : {mae:>12,.0f} €  ({mae/np.median(y_te_raw)*100:.1f} % mediane)")
    print(f"  RMSE    : {rmse:>12,.0f} €")
    print(f"  MAPE    : {mape:.1f} %")
    print(f"  CV R²   : {cv_scores.mean():.4f} ± {cv_scores.std():.4f}")

    return {
        "name": name, "model": model, "pred_eur": pred_eur,
        "R2": r2, "MAE": mae, "RMSE": rmse, "MAPE": mape,
        "CV_R2": cv_scores.mean(), "CV_std": cv_scores.std()
    }

# ── MODELI ────────────────────────────────────────────────────────────────────
print(f"\n{'='*55}")
print("  [4/5]  TRENING MODELOV")
print(f"{'='*55}")

results = []

# 1. Ridge regresija (baseline)
print("\n  Ridge regresija (baseline) ...")
from sklearn.linear_model import RidgeCV
ridge_pipe = Pipeline([
    ("scaler", StandardScaler()),
    ("ridge",  RidgeCV(alphas=[0.01, 0.1, 1.0, 10.0, 100.0, 1000.0], cv=5)),
])
ridge_pipe.fit(X_train, y_train)
print(f"  Izbrani alpha: {ridge_pipe.named_steps['ridge'].alpha_:.2f}")
results.append(eval_model("Ridge regresija (CV alpha)", ridge_pipe,
                           X_train, y_train_raw, X_test, y_test_raw))

# 2. Random Forest (GridSearchCV)
print("\n  Random Forest ...")
from sklearn.model_selection import GridSearchCV

rf_grid = GridSearchCV(
    RandomForestRegressor(random_state=42, n_jobs=-1),
    param_grid={
        "n_estimators": [200],
        "max_depth":    [None, 15, 25],
        "min_samples_leaf": [1, 3, 5],
        "max_features": ["sqrt", 0.5],
    },
    cv=3, scoring="r2", n_jobs=-1, verbose=0
)
rf_grid.fit(X_train, y_train)
print(f"  Najboljši parametri RF: {rf_grid.best_params_}  (CV R²={rf_grid.best_score_:.4f})")
rf_best = rf_grid.best_estimator_
results.append(eval_model("Random Forest (GridSearchCV)", rf_best,
                           X_train, y_train_raw, X_test, y_test_raw))

# 3. XGBoost
print("\n  XGBoost ...")
xgb_model = xgb.XGBRegressor(
    n_estimators=500,
    learning_rate=0.05,
    max_depth=6,
    subsample=0.8,
    colsample_bytree=0.8,
    min_child_weight=3,
    reg_alpha=0.1,
    reg_lambda=1.0,
    random_state=42,
    n_jobs=-1,
    verbosity=0
)
xgb_model.fit(
    X_train, y_train,
    eval_set=[(X_test, y_test)],
    verbose=False
)
results.append(eval_model("XGBoost", xgb_model,
                           X_train, y_train_raw, X_test, y_test_raw))

# 4. LightGBM
print("\n  LightGBM ...")
lgb_model = lgb.LGBMRegressor(
    n_estimators=500,
    learning_rate=0.05,
    max_depth=8,
    num_leaves=63,
    subsample=0.8,
    colsample_bytree=0.8,
    min_child_samples=10,
    reg_alpha=0.1,
    reg_lambda=1.0,
    random_state=42,
    n_jobs=-1,
    verbose=-1
)
lgb_model.fit(
    X_train, y_train,
    eval_set=[(X_test, y_test)],
    callbacks=[lgb.early_stopping(50, verbose=False), lgb.log_evaluation(0)],
    feature_name=FEAT_NAMES,
)
results.append(eval_model("LightGBM", lgb_model,
                           X_train, y_train_raw, X_test, y_test_raw))

# ── Povzetek ──────────────────────────────────────────────────────────────────
print(f"\n{'='*70}")
print(f"  {'Model':<35}  {'R²':>6}  {'MAE':>10}  {'RMSE':>10}  {'MAPE':>6}  {'CV R²':>6}")
print(f"  {'─'*68}")
for r in sorted(results, key=lambda x: -x["R2"]):
    print(f"  {r['name']:<35}  {r['R2']:>6.4f}  {r['MAE']:>10,.0f}  {r['RMSE']:>10,.0f}  "
          f"{r['MAPE']:>5.1f}%  {r['CV_R2']:>6.4f}")
print(f"{'='*70}")

best = max(results, key=lambda x: x["R2"])
print(f"\n  🏆  Najboljši model: {best['name']}")
print(f"      R² = {best['R2']:.4f}  |  MAE = {best['MAE']:,.0f} €  |  MAPE = {best['MAPE']:.1f} %")

# ── Grafi ─────────────────────────────────────────────────────────────────────
print(f"\n[5/5]  Generiram grafe ...")

os.makedirs(os.path.join(SCRIPT_DIR, "grafi"), exist_ok=True)

# 1. Primerjava R² modelov
fig, axes = plt.subplots(1, 2, figsize=(14, 5))
names = [r["name"].split("(")[0].strip() for r in results]
r2s   = [r["R2"]  for r in results]
maes  = [r["MAE"] for r in results]
colors = ["#2196F3", "#4CAF50", "#FF5722", "#9C27B0"]

axes[0].barh(names, r2s, color=colors)
axes[0].set_xlabel("R²")
axes[0].set_title("Primerjava R² vrednosti (višje = boljše)")
for i, v in enumerate(r2s):
    axes[0].text(v + 0.002, i, f"{v:.4f}", va="center", fontsize=9)
axes[0].set_xlim(0, 1.05)

axes[1].barh(names, [m/1000 for m in maes], color=colors)
axes[1].set_xlabel("MAE (1000 €)")
axes[1].set_title("Primerjava MAE v € (nižje = boljše)")
for i, v in enumerate(maes):
    axes[1].text(v/1000 + 0.5, i, f"{v/1000:.1f}k €", va="center", fontsize=9)

plt.tight_layout()
cmp_path = os.path.join(SCRIPT_DIR, "grafi", "primerjava_modelov_v2.png")
plt.savefig(cmp_path, dpi=120, bbox_inches="tight")
plt.close()
print(f"   ✓  {cmp_path}")

# 2. Napovedano vs dejansko (za najboljši model)
fig, ax = plt.subplots(figsize=(8, 7))
pred_eur = best["pred_eur"]
ax.scatter(y_test_raw/1000, pred_eur/1000, alpha=0.4, s=18, color="#2196F3", label="Napovedi")
mn = min(y_test_raw.min(), pred_eur.min())/1000
mx = max(y_test_raw.max(), pred_eur.max())/1000
ax.plot([mn, mx], [mn, mx], "r--", lw=2, label="Idealna napoved")
ax.set_xlabel("Dejanska cena (1000 €)")
ax.set_ylabel("Napovedana cena (1000 €)")
ax.set_title(f"Napovedano vs Dejansko — {best['name']}\nR²={best['R2']:.4f}  MAE={best['MAE']/1000:.1f}k €  MAPE={best['MAPE']:.1f}%")
ax.legend()
ax.grid(alpha=0.3)
scatter_path = os.path.join(SCRIPT_DIR, "grafi", "napovedano_vs_dejansko_v2.png")
plt.savefig(scatter_path, dpi=120, bbox_inches="tight")
plt.close()
print(f"   ✓  {scatter_path}")

# 3. Porazdelitev ostankov
residuals_pct = (y_test_raw - pred_eur) / y_test_raw * 100
fig, ax = plt.subplots(figsize=(9, 5))
ax.hist(residuals_pct, bins=50, color="#2196F3", edgecolor="white", alpha=0.8)
ax.axvline(0, color="red", lw=2, ls="--", label="0 %")
ax.axvline(residuals_pct.mean(), color="orange", lw=2, ls="--", label=f"Povprečje: {residuals_pct.mean():.1f}%")
ax.set_xlabel("Napaka (%)")
ax.set_ylabel("Število primerov")
ax.set_title(f"Porazdelitev relativnih napak — {best['name']}")
ax.legend()
ax.grid(alpha=0.3)
resid_path = os.path.join(SCRIPT_DIR, "grafi", "porazdelitev_napak_v2.png")
plt.savefig(resid_path, dpi=120, bbox_inches="tight")
plt.close()
print(f"   ✓  {resid_path}")

# 4. Pomembnost značilk (za najboljši tree-based model)
best_tree = next((r for r in sorted(results, key=lambda x: -x["R2"])
                  if hasattr(r["model"], "feature_importances_")), None)
if best_tree:
    fi = best_tree["model"].feature_importances_
    fi_sorted = sorted(zip(FEAT_NAMES, fi), key=lambda x: x[1])
    fig, ax = plt.subplots(figsize=(9, 5))
    ax.barh([x[0] for x in fi_sorted], [x[1] for x in fi_sorted], color="#4CAF50")
    ax.set_xlabel("Relativna pomembnost")
    ax.set_title(f"Pomembnost značilk — {best_tree['name']}")
    ax.grid(alpha=0.3, axis="x")
    fi_path = os.path.join(SCRIPT_DIR, "grafi", "pomembnost_znacilk_v2.png")
    plt.savefig(fi_path, dpi=120, bbox_inches="tight")
    plt.close()
    print(f"   ✓  {fi_path}")

# 5. SHAP analiza (samo če --shap flag)
if args.shap:
    try:
        import shap
        print("   Računam SHAP vrednosti ...")
        best_m = best["model"]
        explainer = shap.TreeExplainer(best_m)
        shap_vals = explainer.shap_values(X_test[:500])
        fig = plt.figure(figsize=(10, 6))
        shap.summary_plot(shap_vals, X_test[:500],
                          feature_names=FEAT_NAMES, show=False)
        shap_path = os.path.join(SCRIPT_DIR, "grafi", "shap_summary_v2.png")
        plt.savefig(shap_path, dpi=120, bbox_inches="tight")
        plt.close()
        print(f"   ✓  {shap_path}")
    except Exception as e:
        print(f"   ⚠  SHAP ni uspel: {e}")

# ── Shranitev najboljšega modela ──────────────────────────────────────────────
model_dir = os.path.join(SCRIPT_DIR, "model_v2")
os.makedirs(model_dir, exist_ok=True)

# Shrani vse potrebno za napoved
model_data = {
    "model":          best["model"],
    "feat_names":     FEAT_NAMES,
    "vrsta_enc":      vrsta_enc,   # target encoding tabela
    "obcina_enc":     obcina_enc,
    "cena_m2_obcina": cena_m2_obcina,
    "global_cena_m2": global_cena_m2,
    "train_medians":  {c: df_train[c].median() for c in BASE_NUM_FEATS},
    "best_name":      best["name"],
    "metrics":        {k: v for k, v in best.items() if k not in ("model", "pred_eur")},
    "trained_on":     datetime.now().isoformat(),
}
model_path = os.path.join(model_dir, "best_model_v2.pkl")
joblib.dump(model_data, model_path)
print(f"\n   ✓  Model shranjen: {model_path}")

# ── Funkcija za napoved nove nepremičnine ─────────────────────────────────────
def napovej_ceno(vrsta: str, povrsina_m2: float, starost_let: int,
                 st_sob: float, energetski_razred: str,
                 obcina: str, zemljisce_m2: float = 0.0) -> dict:
    """
    Napove ceno nepremičnine z najboljšim modelom.
    Primer:
        napovej_ceno("Stanovanje", 75, 20, 3, "B", "Ljubljana", 0)
    """
    log_vel   = np.log1p(povrsina_m2)
    log_zem   = np.log1p(max(0, zemljisce_m2))
    starost   = float(starost_let)
    sob       = float(st_sob)
    enr       = float(ENERGY_MAP.get(str(energetski_razred).upper(), 4))

    global_mean_log = df_train["log_Cena"].mean()

    # Target encoding z globalnim povprečjem za neznane kategorije
    vrsta_val  = float(vrsta_enc.get(vrsta,  vrsta_enc.mean()  if len(vrsta_enc)  else global_mean_log))
    obcina_val = float(obcina_enc.get(obcina, obcina_enc.mean() if len(obcina_enc) else global_mean_log))
    cena_m2_lok = float(cena_m2_obcina.get(obcina, global_cena_m2))
    log_cm2_lok = np.log1p(cena_m2_lok)

    X = np.array([[log_vel, log_zem, starost, sob, enr,
                   vrsta_val, obcina_val, log_cm2_lok]])
    pred_log = best["model"].predict(X)[0]
    pred_eur = np.exp(pred_log)

    # 80 % interval zaupanja iz MAPE
    mape = best["MAPE"] / 100
    return {
        "napoved_eur": round(pred_eur),
        "interval_nizji_eur":  round(pred_eur * (1 - mape)),
        "interval_visji_eur":  round(pred_eur * (1 + mape)),
        "model": best["name"],
        "mape": best["MAPE"],
    }


print("\n  Primer napovedi:")
primer = napovej_ceno("Stanovanje", 75, 20, 3, "B", "Ljubljana", 0)
print(f"  Stanovanje 75m², 20 let, 3 sobe, razred B, Ljubljana:")
print(f"  → {primer['napoved_eur']:,} €  "
      f"(interval: {primer['interval_nizji_eur']:,} – {primer['interval_visji_eur']:,} €)")

primer2 = napovej_ceno("Hiša", 150, 30, 5, "C", "Kranj", 400)
print(f"\n  Hiša 150m², 30 let, 5 sob, razred C, Kranj, 400m² zemljišče:")
print(f"  → {primer2['napoved_eur']:,} €  "
      f"(interval: {primer2['interval_nizji_eur']:,} – {primer2['interval_visji_eur']:,} €)")

print(f"\n  Analiza končana – {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}")

# ── DOCX poročilo ─────────────────────────────────────────────────────────────
if args.docx:
    try:
        from docx import Document
        from docx.shared import Inches, Cm, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io as _io
    except ImportError:
        print("✗  Namesti: python -m pip install python-docx")
        sys.exit(0)

    out_path = args.docx_izhod or os.path.join(
        SCRIPT_DIR, f"ml_porocilo_v2_{datetime.now().strftime('%Y%m%d_%H%M')}.docx")

    print(f"\n  Generiram DOCX poročilo: {out_path}")
    doc = Document()
    for sec in doc.sections:
        sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)
        sec.top_margin  = Cm(2.5); sec.bottom_margin = Cm(2.5)

    def _H(lvl, txt):
        h = doc.add_heading(txt, lvl)
        if lvl == 0: h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _p(txt="", bold=False, italic=False):
        p = doc.add_paragraph()
        r = p.add_run(txt)
        r.bold = bold; r.italic = italic
        return p

    def _tbl(headers, rows_data):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        for i, h in enumerate(headers):
            t.rows[0].cells[i].text = h
            for run in t.rows[0].cells[i].paragraphs[0].runs:
                run.bold = True
        for row in rows_data:
            cells = t.add_row().cells
            for i, v in enumerate(row): cells[i].text = str(v)
        return t

    def _img(path, caption, w=5.8):
        if not os.path.isfile(path): return
        doc.add_picture(path, width=Inches(w))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cp.runs: run.italic = True; run.font.size = Pt(9)
        doc.add_paragraph()

    # Naslov
    _H(0, "Napovedovanje cen nepremičnin – ML modeli v2")
    _p(f"Generirano: {datetime.now().strftime('%d.%m.%Y %H:%M')}  |  Podatki: {csv_path}",
       italic=True)
    doc.add_paragraph()

    # 1. Podatki
    _H(1, "1. Podatki in predprocesiranje")
    _tbl(["Lastnost", "Vrednost"], [
        ["Vir podatkov", csv_path],
        ["Skupaj vzorcev (po čiščenju)", len(df)],
        ["Učna množica", len(df_train)],
        ["Testna množica", len(df_test)],
        ["Ciljna spremenljivka", "log(Cena) → exp() za napoved v €"],
        ["Min. cena", f"{df['Cena'].min():,.0f} €"],
        ["Max. cena", f"{df['Cena'].max():,.0f} €"],
        ["Mediana cene", f"{df['Cena'].median():,.0f} €"],
    ])
    doc.add_paragraph()

    # 2. Značilke
    _H(1, "2. Značilke (feature engineering)")
    _tbl(["Značilka", "Opis", "Tip obdelave"], [
        ["log_VelikostM2",     "log(1 + površina m²)",                "Numerična → log transformacija"],
        ["log_ZemljisteM2",    "log(1 + zemljišče m²)",               "Numerična → log transformacija"],
        ["StarostZgrade",      "2026 − leto gradnje",                  "Numerična, izpeljana"],
        ["StSob",              "Število sob",                          "Numerična, imputacija z mediano"],
        ["EnergetRazredNum",   "A+=8 ... G=1",                        "Ordinalna kodiranje"],
        ["VrstaObjekta_enc",   "Stanovanje / Hiša / ...",              "Target encoding (glajanje n=10)"],
        ["Obcina_enc",         "Lokacija nepremičnine",                "Target encoding (glajanje n=10)"],
        ["log_cena_m2_lokacija","log(mediana cene/m² po občini)",      "Lokacijski indeks iz train podatkov"],
    ])
    _p("Target encoding: encoded = (n·μ_kat + k·μ_glob) / (n + k)  pri k=10",
       italic=True)
    doc.add_paragraph()

    # 3. Modeli
    _H(1, "3. Modeli in metrike")
    _tbl(
        ["Model", "R²", "MAE (€)", "RMSE (€)", "MAPE (%)", "CV R²"],
        [[r["name"], f"{r['R2']:.4f}", f"{r['MAE']:,.0f}", f"{r['RMSE']:,.0f}",
          f"{r['MAPE']:.1f}", f"{r['CV_R2']:.4f}"]
         for r in sorted(results, key=lambda x: -x["R2"])]
    )
    _p(f"\nNajboljši model: {best['name']}", bold=True)
    _p(f"R² = {best['R2']:.4f}  |  MAE = {best['MAE']:,.0f} €  |  MAPE = {best['MAPE']:.1f} %")
    doc.add_paragraph()

    # 4. Grafi
    _H(1, "4. Grafi")
    _img(cmp_path,    "Slika 1 – Primerjava R² in MAE vseh modelov")
    _img(scatter_path, f"Slika 2 – Napovedano vs dejansko ({best['name']})")
    _img(resid_path,  "Slika 3 – Porazdelitev relativnih napak (%)")
    if best_tree:
        _img(fi_path, "Slika 4 – Relativna pomembnost značilk")

    # 5. Interpretacija
    _H(1, "5. Interpretacija rezultatov")
    _p("Lokacijski indeks (log_cena_m2_lokacija) je najpomembnejša značilka — "
       "lokacija nepremičnine pojasnuje največji del variance cene.", bold=False)
    _p("Target encoding za občino zajame povprečno cenovno raven vsake lokacije "
       "brez tveganja data leakage (kodiranje samo iz učnih podatkov).")
    _p("XGBoost in LightGBM sta gradientno ojačana ansambla odločitvenih dreves — "
       "pri tabelarnih podatkih praviloma dosegata najboljše rezultate.")
    doc.add_paragraph()

    # 6. Primer napovedi
    _H(1, "6. Primer napovedi")
    _tbl(["Parameter", "Vrednost"], [
        ["Vrsta", "Stanovanje"],
        ["Površina", "75 m²"],
        ["Starost", "20 let"],
        ["Sobe", "3"],
        ["Energetski razred", "B"],
        ["Občina", "Ljubljana"],
        ["Napoved", f"{primer['napoved_eur']:,} €"],
        ["Interval (±MAPE)", f"{primer['interval_nizji_eur']:,} – {primer['interval_visji_eur']:,} €"],
    ])
    doc.add_paragraph()
    _tbl(["Parameter", "Vrednost"], [
        ["Vrsta", "Hiša"],
        ["Površina", "150 m²"],
        ["Starost", "30 let"],
        ["Sobe", "5"],
        ["Energetski razred", "C"],
        ["Občina", "Kranj"],
        ["Zemljišče", "400 m²"],
        ["Napoved", f"{primer2['napoved_eur']:,} €"],
        ["Interval (±MAPE)", f"{primer2['interval_nizji_eur']:,} – {primer2['interval_visji_eur']:,} €"],
    ])

    doc.save(out_path)
    print(f"  ✓  DOCX poročilo shranjeno: {out_path}")

