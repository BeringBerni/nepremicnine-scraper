"""
generiraj_dokumentacijo.py
Generira popoln DOCX dokument projekta UUI-lv2.
Zagon: py generiraj_dokumentacijo.py
"""
import sys, os, io, struct, zlib, csv, math, statistics, random, itertools
from datetime import datetime
from collections import defaultdict

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

sys.setrecursionlimit(10_000)

try:
    from docx import Document
    from docx.shared import Inches, Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Namesti: py -m pip install python-docx")
    sys.exit(1)

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
SVG_DIR    = os.path.join(SCRIPT_DIR, "grafi")

# ── CSV podatki ────────────────────────────────────────────────────────────────
CSV_PATH = None
for cand in ["nepremicnine_export_prodaja.csv", "nepremicnine_export.csv"]:
    p = os.path.join(SCRIPT_DIR, cand)
    if os.path.isfile(p):
        CSV_PATH = p
        break

rows = []
if CSV_PATH:
    with open(CSV_PATH, newline="", encoding="utf-8-sig") as f:
        for r in csv.DictReader(f, delimiter=";"):
            try: cena = float(str(r.get("Cena","")).replace(",","."))
            except: continue
            if not (5000 < cena < 5_000_000): continue
            def sf(k):
                try: v=float(str(r.get(k,"")).replace(",",".")); return v if math.isfinite(v) else None
                except: return None
            rows.append({
                "Cena": cena, "VelikostM2": sf("VelikostM2"),
                "ZemljisteM2": sf("ZemljisteM2"), "LetoGradnje": sf("LetoGradnje"),
                "CenaNaM2": sf("CenaNaM2"),
                "Lokacija": r.get("Lokacija","").strip(),
                "VrstaObjekta": r.get("VrstaObjekta","").strip(),
            })

cene = [r["Cena"] for r in rows]
m2s  = [r["VelikostM2"] for r in rows if r["VelikostM2"]]
zems = [r["ZemljisteM2"] for r in rows if r["ZemljisteM2"]]
letos= [r["LetoGradnje"] for r in rows if r["LetoGradnje"] and r["LetoGradnje"]>1900]
cm2s = [r["CenaNaM2"] for r in rows if r["CenaNaM2"]]

lok = defaultdict(list)
for r in rows:
    if r["Lokacija"]: lok[r["Lokacija"]].append(r["Cena"])
vrsta_d = defaultdict(list)
for r in rows:
    vrsta_d[r["VrstaObjekta"] or "?"].append(r["Cena"])

def pearson(xs, ys):
    n = len(xs)
    if n < 3: return None
    mx, my = statistics.mean(xs), statistics.mean(ys)
    num = sum((x-mx)*(y-my) for x,y in zip(xs,ys))
    den = math.sqrt(sum((x-mx)**2 for x in xs)*sum((y-my)**2 for y in ys))
    return num/den if den else None

# ══════════════════════════════════════════════════════════════════════════════
#  GRAFIČNI SISTEM – visoka ločljivost, berljive oznake, referenčne črte
# ══════════════════════════════════════════════════════════════════════════════

_F57 = {
    ' ':[0,0,0,0,0,0,0],
    '0':[14,17,19,21,25,17,14],'1':[4,12,4,4,4,4,14],'2':[14,17,1,2,4,8,31],
    '3':[31,2,4,2,1,17,14],'4':[2,6,10,18,31,2,2],'5':[31,16,30,1,1,17,14],
    '6':[6,8,16,30,17,17,14],'7':[31,1,2,4,8,8,8],'8':[14,17,17,14,17,17,14],
    '9':[14,17,17,15,1,2,12],'.':[0,0,0,0,0,12,12],',':[ 0,0,0,0,0,12,8],
    '-':[0,0,0,31,0,0,0],'+':[0,4,4,31,4,4,0],'%':[24,25,2,4,8,19,3],
    '/':[1,2,4,8,16,0,0],'(':[ 2,4,8,8,8,4,2],')':[8,4,2,2,2,4,8],
    ':':[0,12,12,0,12,12,0],'=':[0,0,31,0,31,0,0],
    'k':[16,16,18,20,24,20,18],'m':[0,26,21,21,21,21,21],
    'a':[0,14,1,15,17,19,13],'b':[16,16,30,17,17,17,30],'c':[0,14,17,16,16,17,14],
    'd':[1,1,13,19,17,19,13],'e':[0,14,17,31,16,17,14],'f':[6,9,8,30,8,8,8],
    'g':[0,14,17,17,15,1,14],'h':[16,16,30,17,17,17,17],'i':[12,12,0,12,12,12,14],
    'j':[6,0,6,6,6,22,12],'l':[12,4,4,4,4,4,14],'n':[0,28,18,17,17,17,17],
    'o':[0,14,17,17,17,17,14],'p':[0,30,17,17,30,16,16],'r':[0,22,25,16,16,16,16],
    's':[0,14,16,14,1,17,14],'t':[4,4,31,4,4,5,2],'u':[0,17,17,17,17,19,13],
    'v':[0,17,17,17,17,10,4],'z':[0,31,2,4,8,16,31],
    'A':[14,17,17,31,17,17,17],'B':[30,17,17,30,17,17,30],'C':[14,17,16,16,16,17,14],
    'D':[28,18,17,17,17,18,28],'E':[31,16,16,30,16,16,31],'F':[31,16,16,30,16,16,16],
    'G':[14,17,16,23,17,17,14],'H':[17,17,17,31,17,17,17],'I':[14,4,4,4,4,4,14],
    'J':[7,2,2,2,2,18,12],'K':[17,18,20,24,20,18,17],'L':[16,16,16,16,16,16,31],
    'M':[17,27,21,21,17,17,17],'N':[17,25,21,21,19,17,17],'O':[14,17,17,17,17,17,14],
    'P':[30,17,17,30,16,16,16],'R':[30,17,17,30,20,18,17],'S':[14,17,16,14,1,17,14],
    'T':[31,4,4,4,4,4,4],'U':[17,17,17,17,17,17,14],'V':[17,17,17,17,17,10,4],
    'W':[17,17,17,21,21,27,17],'X':[17,10,4,4,4,10,17],'Y':[17,17,10,4,4,4,4],
    'Z':[31,1,2,4,8,16,31],
}

def _asc(s):
    tbl = {'š':'s','Š':'S','č':'c','Č':'C','ž':'z','Ž':'Z','ć':'c','đ':'d',
           'á':'a','é':'e','í':'i','ó':'o','ú':'u','²':'2','³':'3',
           '×':'x','→':'>','↔':'<>','–':'-','—':'-','€':'EUR'}
    return ''.join(tbl.get(c, c) for c in s)

def _make_png(W, H, pixels):
    raw = bytearray()
    for row in pixels:
        raw.append(0)
        for p in row: raw += bytes(p)
    comp = zlib.compress(bytes(raw), 9)
    def chunk(tag, data):
        crc = zlib.crc32(tag + data) & 0xFFFFFFFF
        return struct.pack('>I', len(data)) + tag + data + struct.pack('>I', crc)
    return (b'\x89PNG\r\n\x1a\n'
            + chunk(b'IHDR', struct.pack('>IIBBBBB', W, H, 8, 2, 0, 0, 0))
            + chunk(b'IDAT', comp)
            + chunk(b'IEND', b''))


class Canvas:
    """Pixel canvas – visoka ločljivost za grafe v DOCX."""
    def __init__(self, W=1400, H=700, bg=(255, 255, 255)):
        self.W, self.H = W, H
        self._p = [[[*bg]] * W for _ in range(H)]

    def _set(self, x, y, c):
        if 0 <= x < self.W and 0 <= y < self.H:
            self._p[y][x] = list(c)

    def fill_rect(self, x, y, w, h, c):
        lc = list(c)
        for dy in range(max(0, y), min(self.H, y + h)):
            row = self._p[dy]
            for dx in range(max(0, x), min(self.W, x + w)):
                row[dx] = lc

    def hline(self, x1, x2, y, c, t=1):
        lc = list(c)
        for x in range(max(0, min(x1, x2)), min(self.W, max(x1, x2) + 1)):
            for dt in range(t):
                if 0 <= y + dt < self.H:
                    self._p[y + dt][x] = lc

    def vline(self, x, y1, y2, c, t=1):
        lc = list(c)
        for y in range(max(0, min(y1, y2)), min(self.H, max(y1, y2) + 1)):
            for dt in range(t):
                if 0 <= x + dt < self.W:
                    self._p[y][x + dt] = lc

    def dashed_vline(self, x, y1, y2, c, t=2, dash=10):
        """Prekinjena vertikalna črta."""
        for y in range(min(y1, y2), max(y1, y2) + 1):
            if (y // dash) % 2 == 0:
                for dt in range(t):
                    if 0 <= y < self.H and 0 <= x + dt < self.W:
                        self._p[y][x + dt] = list(c)

    def bresenham(self, x1, y1, x2, y2, c, t=2):
        dx, dy = abs(x2 - x1), abs(y2 - y1)
        sx = 1 if x1 < x2 else -1
        sy = 1 if y1 < y2 else -1
        e, x, y = dx - dy, x1, y1
        while True:
            for dt in range(t): self._set(x + dt, y, c)
            if x == x2 and y == y2: break
            e2 = 2 * e
            if e2 > -dy: e -= dy; x += sx
            if e2 < dx:  e += dx; y += sy

    def circle(self, cx, cy, r, c, filled=True):
        for dy in range(-r, r + 1):
            for dx in range(-r, r + 1):
                if dx * dx + dy * dy <= r * r:
                    self._set(cx + dx, cy + dy, c)

    def text(self, x, y, s, c, sc=2):
        """Renderira tekst z bitmap pisavo. sc=2 = lepo berljivo."""
        lc = list(c)
        for ch in _asc(str(s)):
            bits = _F57.get(ch, [0] * 7)
            for ri, b in enumerate(bits):
                for ci in range(5):
                    if b & (1 << (4 - ci)):
                        for ddy in range(sc):
                            for ddx in range(sc):
                                self._set(x + ci * sc + ddx, y + ri * sc + ddy, lc)
            x += 6 * sc

    def text_centered(self, cx, y, s, c, sc=2):
        w = len(_asc(str(s))) * 6 * sc
        self.text(cx - w // 2, y, s, c, sc)

    def text_right(self, rx, y, s, c, sc=2):
        w = len(_asc(str(s))) * 6 * sc
        self.text(rx - w, y, s, c, sc)

    def text_width(self, s, sc=2):
        return len(_asc(str(s))) * 6 * sc

    def legend_box(self, x, y, items, sc=2):
        """Nariše legendo z obarvanimi kvadratki."""
        pad, line_h = 10, 7 * sc + 8
        box_w = max(self.text_width(t, sc) for _, t in items) + 8 * sc + pad * 2
        box_h = len(items) * line_h + pad * 2
        self.fill_rect(x, y, box_w, box_h, (248, 248, 248))
        self.hline(x, x + box_w, y, (180, 180, 180))
        self.hline(x, x + box_w, y + box_h, (180, 180, 180))
        self.vline(x, y, y + box_h, (180, 180, 180))
        self.vline(x + box_w, y, y + box_h, (180, 180, 180))
        for i, (col, label) in enumerate(items):
            ly = y + pad + i * line_h
            self.fill_rect(x + pad, ly + sc, 6 * sc, 5 * sc, col)
            self.text(x + pad + 7 * sc, ly, label, (40, 40, 40), sc)

    def to_png(self):
        return _make_png(self.W, self.H,
                         [[bytes(self._p[y][x]) for x in range(self.W)]
                          for y in range(self.H)])


# ── Pomočnik za osi ───────────────────────────────────────────────────────────
def _nice_ticks(mn, mx, n=6):
    """Izračuna 'lepe' vrednosti za osi (zaokrožene)."""
    span = mx - mn
    if span == 0: return [mn]
    raw_step = span / n
    mag = 10 ** math.floor(math.log10(raw_step))
    for mult in [1, 2, 2.5, 5, 10]:
        step = mag * mult
        if span / step <= n + 1:
            break
    start = math.floor(mn / step) * step
    ticks = []
    v = start
    while v <= mx + step * 0.01:
        if mn - step * 0.1 <= v <= mx + step * 0.1:
            ticks.append(v)
        v += step
    return ticks

def _fmt(v):
    """Formatira vrednost za os (k = tisoč, M = milijon)."""
    av = abs(v)
    if av >= 1_000_000: return f"{v/1_000_000:.1f}M"
    if av >= 10_000:    return f"{v/1_000:.0f}k"
    if av >= 1_000:     return f"{v/1_000:.1f}k"
    if av >= 100:       return f"{v:.0f}"
    if av >= 10:        return f"{v:.1f}"
    return f"{v:.2f}"


# ── HISTOGRAM ─────────────────────────────────────────────────────────────────
def chart_hist(vals, title, xlabel, ylabel="Stevilo oglasov",
               bar_color=(70, 130, 180), n_bins=25):
    """Histogram z vidnima mediano in povprečjem + legenda + N/min/max."""
    if not vals or len(vals) < 2: return None

    mn, mx = min(vals), max(vals)
    if mn == mx: mx = mn + 1
    step = (mx - mn) / n_bins
    bins = [0] * n_bins
    for v in vals:
        bins[min(int((v - mn) / step), n_bins - 1)] += 1

    med = statistics.median(vals)
    avg = statistics.mean(vals)
    std = statistics.stdev(vals) if len(vals) > 1 else 0

    W, H = 1400, 680
    ML, MR, MT, MB = 110, 40, 90, 110
    CW, CH = W - ML - MR, H - MT - MB
    cv = Canvas(W, H)

    # Ozadje grafa
    cv.fill_rect(ML, MT, CW, CH, (252, 252, 254))

    # Grid črte (horizontalne)
    max_cnt = max(bins) or 1
    y_ticks = _nice_ticks(0, max_cnt, 6)
    for yt in y_ticks:
        yp = H - MB - int(yt / max_cnt * CH)
        cv.hline(ML, ML + CW, yp, (220, 220, 228), 1)
        cv.text_right(ML - 8, yp - 7, _fmt(yt), (90, 90, 100), 2)

    # Stolpci
    bw = CW // n_bins
    for i, cnt in enumerate(bins):
        if cnt == 0: continue
        bh = max(2, int(cnt / max_cnt * CH))
        bx = ML + i * bw + 2
        by = H - MB - bh
        # Gradient efekt (svetlejši vrh)
        cv.fill_rect(bx, by, bw - 3, bh, bar_color)
        cv.fill_rect(bx, by, bw - 3, max(1, bh // 6),
                     tuple(min(255, c + 40) for c in bar_color))

    # Vertikalna črta – MEDIANA (rdeča, polna)
    def xpx(v): return ML + int((v - mn) / (mx - mn) * CW)

    med_x = xpx(med)
    cv.vline(med_x, MT, H - MB, (210, 40, 40), 3)
    # Oznaka mediane zgoraj
    cv.fill_rect(med_x - 2, MT - 38, cv.text_width(f"Mediana: {_fmt(med)}", 2) + 16, 30,
                 (210, 40, 40))
    cv.text(med_x + 6, MT - 32, f"Mediana: {_fmt(med)}", (255, 255, 255), 2)

    # Vertikalna črta – POVPREČJE (zelena, polna)
    avg_x = xpx(avg)
    cv.vline(avg_x, MT, H - MB, (34, 139, 34), 3)
    lbl_avg = f"Povp.: {_fmt(avg)}"
    lbl_w = cv.text_width(lbl_avg, 2) + 16
    lbl_x = avg_x + 6 if avg_x + lbl_w + 10 < W - MR else avg_x - lbl_w - 6
    cv.fill_rect(lbl_x - 2, MT - 38, lbl_w, 30, (34, 139, 34))
    cv.text(lbl_x + 6, MT - 32, lbl_avg, (255, 255, 255), 2)

    # Os X in Y
    cv.hline(ML, ML + CW, H - MB, (60, 60, 70), 3)
    cv.vline(ML, MT, H - MB, (60, 60, 70), 3)

    # X-osi oznake
    x_ticks = _nice_ticks(mn, mx, 7)
    for xt in x_ticks:
        xp = xpx(xt)
        if ML <= xp <= ML + CW:
            cv.vline(xp, H - MB, H - MB + 10, (100, 100, 110), 2)
            cv.text_centered(xp, H - MB + 16, _fmt(xt), (80, 80, 90), 2)

    # Naslovi osi
    cv.text_centered(W // 2, H - 28, _asc(xlabel), (50, 50, 60), 2)
    cv.text_centered(W // 2, 14, _asc(title), (20, 20, 30), 3)

    # Y-os naslov (navpično – imitiramo z malimi znaki)
    for i, ch in enumerate(_asc(ylabel)[:14]):
        cv.text(8, MT + i * 16, ch, (50, 50, 60), 2)

    # Legenda – statistike v pravokotnik
    cv.legend_box(ML + CW - 340, MT + 20, [
        ((210, 40, 40),  f"Mediana:  {_fmt(med)}"),
        ((34, 139, 34),  f"Povprecje: {_fmt(avg)}"),
        ((100, 100, 140), f"Std dev:  {_fmt(std)}"),
        ((100, 100, 140), f"N = {len(vals)}"),
        ((100, 100, 140), f"Min: {_fmt(mn)}  Max: {_fmt(mx)}"),
    ], sc=2)

    return cv.to_png()


# ── SCATTER PLOT ──────────────────────────────────────────────────────────────
def chart_scatter(xs, ys, title, xlabel, ylabel,
                  dot_color=(70, 130, 180)):
    """Razsevni diagram z regresijsko premico, r-vrednostjo in referenčnimi črtami."""
    if len(xs) < 5: return None

    # Odstrani outlierje (nad 99. percentilom za oba)
    xs_s = sorted(xs); ys_s = sorted(ys)
    x_hi = xs_s[int(len(xs_s) * 0.99)]
    y_hi = ys_s[int(len(ys_s) * 0.99)]
    pts = [(x, y) for x, y in zip(xs, ys) if x <= x_hi and y <= y_hi]
    if len(pts) < 5: pts = list(zip(xs, ys))
    xs2, ys2 = [p[0] for p in pts], [p[1] for p in pts]

    xmn, xmx = min(xs2), max(xs2)
    ymn, ymx = min(ys2), max(ys2)
    xpad = (xmx - xmn) * 0.04; ypad = (ymx - ymn) * 0.04
    xmn -= xpad; xmx += xpad; ymn -= ypad; ymx += ypad
    if xmn == xmx: xmx = xmn + 1
    if ymn == ymx: ymx = ymn + 1

    W, H = 1400, 700
    ML, MR, MT, MB = 110, 50, 90, 110
    CW, CH = W - ML - MR, H - MT - MB

    def xpx(v): return ML + int((v - xmn) / (xmx - xmn) * CW)
    def ypx(v): return H - MB - int((v - ymn) / (ymx - ymn) * CH)

    cv = Canvas(W, H)
    cv.fill_rect(ML, MT, CW, CH, (252, 252, 254))

    # Grid
    for xt in _nice_ticks(xmn, xmx, 7):
        xp = xpx(xt)
        if ML <= xp <= ML + CW:
            cv.vline(xp, MT, H - MB, (220, 220, 228), 1)
    for yt in _nice_ticks(ymn, ymx, 6):
        yp = ypx(yt)
        if MT <= yp <= H - MB:
            cv.hline(ML, ML + CW, yp, (220, 220, 228), 1)

    # Regresijska premica
    mx_ = statistics.mean(xs2); my_ = statistics.mean(ys2)
    num_ = sum((x - mx_) * (y - my_) for x, y in zip(xs2, ys2))
    den_ = sum((x - mx_) ** 2 for x in xs2) or 1
    slope = num_ / den_; intercept = my_ - slope * mx_
    r_val = pearson(xs2, ys2) or 0

    # Premica – narišemo samo znotraj plot območja
    x_lo, x_hi2 = xmn, xmx
    y_lo = slope * x_lo + intercept
    y_hi2 = slope * x_hi2 + intercept
    p1x, p1y = xpx(x_lo), ypx(max(ymn, min(ymx, y_lo)))
    p2x, p2y = xpx(x_hi2), ypx(max(ymn, min(ymx, y_hi2)))
    cv.bresenham(p1x, p1y, p2x, p2y, (210, 50, 50), 3)

    # Črtkana referenčna črta – mediana X in Y
    med_x = statistics.median(xs2); med_y = statistics.median(ys2)
    cv.dashed_vline(xpx(med_x), MT, H - MB, (150, 100, 200), 2, dash=14)
    cv.hline(ML, ML + CW, ypx(med_y), (150, 100, 200), 2)
    # Oznaka median
    cv.text(xpx(med_x) + 6, MT + 4, f"med X={_fmt(med_x)}", (120, 60, 180), 2)
    cv.text_right(ML + CW - 6, ypx(med_y) - 20, f"med Y={_fmt(med_y)}", (120, 60, 180), 2)

    # Točke  (dot_color mora biti 3-tuple – brez alfa, sicer PNG pokvarjen)
    for x, y in zip(xs2, ys2):
        cv.circle(xpx(x), ypx(y), 5, dot_color)

    # Osi
    cv.hline(ML, ML + CW, H - MB, (60, 60, 70), 3)
    cv.vline(ML, MT, H - MB, (60, 60, 70), 3)

    # X ticki
    for xt in _nice_ticks(xmn, xmx, 7):
        xp = xpx(xt)
        if ML <= xp <= ML + CW:
            cv.vline(xp, H - MB, H - MB + 10, (100, 100, 110), 2)
            cv.text_centered(xp, H - MB + 16, _fmt(xt), (80, 80, 90), 2)
    # Y ticki
    for yt in _nice_ticks(ymn, ymx, 6):
        yp = ypx(yt)
        if MT <= yp <= H - MB:
            cv.hline(ML - 10, ML, yp, (100, 100, 110), 2)
            cv.text_right(ML - 14, yp - 7, _fmt(yt), (80, 80, 90), 2)

    # Naslovi
    cv.text_centered(W // 2, H - 28, _asc(xlabel), (50, 50, 60), 2)
    for i, ch in enumerate(_asc(ylabel)[:14]):
        cv.text(8, MT + i * 16, ch, (50, 50, 60), 2)
    cv.text_centered(W // 2, 14, _asc(title), (20, 20, 30), 3)

    # Legenda
    r_sign = "+" if r_val >= 0 else ""
    cv.legend_box(ML + 20, MT + 20, [
        ((210, 50, 50),  f"Regr. premica: y = {slope:.1f}x + {intercept/1000:.0f}k"),
        ((150, 100, 200), f"Mediana X: {_fmt(med_x)}"),
        ((150, 100, 200), f"Mediana Y: {_fmt(med_y)}"),
        ((70, 130, 180), f"Pearsonov r = {r_sign}{r_val:.4f}"),
        ((100, 100, 140), f"N = {len(xs2)}  (brez zg. 1% outlierjev)"),
    ], sc=2)

    return cv.to_png()


# ── HORIZONTALNI BAR CHART ────────────────────────────────────────────────────
def chart_hbar(labels, vals, title, xlabel="Cena (EUR)",
               colors=None, show_median_line=True):
    """Horizontalni stolpičar z vrednostmi, N-ji in referenčno medialno črto."""
    n = len(labels)
    if n == 0: return None

    bar_h = 52
    W = 1400
    ML = 300; MR = 60; MT = 90; MB = 80
    H = MT + n * (bar_h + 12) + MB
    CW = W - ML - MR

    mx_v = max(vals) or 1
    overall_med = statistics.median(vals)

    palette = [
        (70, 130, 180), (60, 150, 100), (180, 90, 60),
        (140, 80, 180), (180, 160, 40), (60, 160, 180),
        (180, 60, 100), (100, 140, 60), (100, 100, 180),
        (180, 120, 60), (60, 120, 160), (140, 140, 60),
    ]
    if colors: palette = colors * (n // len(colors) + 1)

    cv = Canvas(W, H, (255, 255, 255))

    # Ozadje
    cv.fill_rect(ML, MT, CW, H - MT - MB, (252, 252, 254))

    # Vertikalne grid črte
    x_ticks = _nice_ticks(0, mx_v, 6)
    for xt in x_ticks:
        xp = ML + int(xt / mx_v * CW)
        if 0 <= xp <= ML + CW:
            cv.vline(xp, MT, H - MB, (220, 220, 228), 1)
            cv.text_centered(xp, H - MB + 12, _fmt(xt), (80, 80, 90), 2)

    # Mediana vseh vrednosti – vertikalna rdeča črta
    if show_median_line:
        med_x = ML + int(overall_med / mx_v * CW)
        cv.vline(med_x, MT, H - MB, (210, 40, 40), 3)
        cv.fill_rect(med_x - cv.text_width(f"Mediana: {_fmt(overall_med)}", 2) // 2 - 8,
                     MT - 36,
                     cv.text_width(f"Mediana: {_fmt(overall_med)}", 2) + 16, 28,
                     (210, 40, 40))
        cv.text_centered(med_x, MT - 30,
                         f"Mediana: {_fmt(overall_med)}", (255, 255, 255), 2)

    # Stolpci
    for i, (lbl, v) in enumerate(zip(labels, vals)):
        bx = ML
        by = MT + i * (bar_h + 12) + 6
        bw = max(4, int(v / mx_v * CW))
        col = palette[i % len(palette)]

        # Stolpec
        cv.fill_rect(bx, by, bw, bar_h, col)
        # Svetlejši vrh
        cv.fill_rect(bx, by, bw, bar_h // 5,
                     tuple(min(255, c + 50) for c in col))

        # Oznaka na desni strani stolpca
        val_lbl = _fmt(v)
        cv.text(bx + bw + 10, by + bar_h // 2 - 7, val_lbl, (40, 40, 50), 2)

        # Oznaka labele (levo od osi)
        lbl_clean = _asc(lbl[:30])
        cv.text_right(ML - 14, by + bar_h // 2 - 7, lbl_clean, (30, 30, 40), 2)

    # Osi
    cv.hline(ML, ML + CW, H - MB, (60, 60, 70), 3)
    cv.vline(ML, MT, H - MB, (60, 60, 70), 3)

    # Naslovi
    cv.text_centered(W // 2, H - 24, _asc(xlabel), (50, 50, 60), 2)
    cv.text_centered(W // 2, 14, _asc(title), (20, 20, 30), 3)

    return cv.to_png()


# ── BOX-STYLE PRIMERJAVA (scatter + mediana linija po skupinah) ───────────────
def chart_group_scatter(groups, title, ylabel="Cena (EUR)"):
    """Točkovni graf po skupinah z medianami – za primerjavo vrst."""
    if not groups: return None
    all_vals = [v for _, vs in groups for v in vs]
    ymn, ymx = min(all_vals), max(all_vals)
    ypad = (ymx - ymn) * 0.06
    ymn -= ypad; ymx += ypad

    W, H = 1400, 700
    MT, MB = 90, 110
    n = len(groups)
    ML = 100; MR = 60
    CW, CH = W - ML - MR, H - MT - MB
    col_w = CW // n

    palette = [(70,130,180),(60,160,100),(210,80,60),
               (140,80,200),(200,160,40),(60,160,180),(180,60,100)]

    cv = Canvas(W, H)
    cv.fill_rect(ML, MT, CW, CH, (252, 252, 254))

    def ypx(v): return H - MB - int((v - ymn) / (ymx - ymn) * CH)

    # Horizontalne grid črte
    for yt in _nice_ticks(ymn, ymx, 6):
        yp = ypx(yt)
        if MT <= yp <= H - MB:
            cv.hline(ML, ML + CW, yp, (220, 220, 228), 1)
            cv.text_right(ML - 10, yp - 7, _fmt(yt), (80, 80, 90), 2)

    # Po skupinah
    for gi, (name, vals) in enumerate(groups):
        cx = ML + gi * col_w + col_w // 2
        col = palette[gi % len(palette)]
        col_lt = tuple(min(255, c + 80) for c in col)

        # Točke (jitter)
        import random as _rnd
        _rnd.seed(gi * 7 + 13)
        for v in vals:
            jx = cx + _rnd.randint(-col_w // 3, col_w // 3)
            cv.circle(jx, ypx(v), 4, col_lt)

        med_v = statistics.median(vals)
        avg_v = statistics.mean(vals)

        # Mediana – debela horizontalna črta (rdeča)
        cv.hline(cx - col_w // 3, cx + col_w // 3, ypx(med_v), (200, 30, 30), 5)
        # Povprečje – zeleni diamant
        for d in range(-8, 9):
            for dd in range(-8, 9):
                if abs(d) + abs(dd) <= 8:
                    cv.circle(cx, ypx(avg_v), 0, (20, 160, 20))
        cv.circle(cx, ypx(avg_v), 8, (20, 160, 20))

        # Oznaka N in mediane pod grafom
        lbl_clean = _asc(name[:18])
        cv.text_centered(cx, H - MB + 14, lbl_clean, (30, 30, 40), 2)
        cv.text_centered(cx, H - MB + 34, f"med={_fmt(med_v)}", (180, 30, 30), 2)
        cv.text_centered(cx, H - MB + 54, f"N={len(vals)}", (80, 80, 90), 2)

    # Osi
    cv.hline(ML, ML + CW, H - MB, (60, 60, 70), 3)
    cv.vline(ML, MT, H - MB, (60, 60, 70), 3)

    # Naslov
    cv.text_centered(W // 2, 14, _asc(title), (20, 20, 30), 3)
    for i, ch in enumerate(_asc(ylabel)[:14]):
        cv.text(8, MT + i * 16, ch, (50, 50, 60), 2)

    # Legenda
    cv.legend_box(ML + 20, MT + 20, [
        ((200, 30, 30), "Rdeca crta = mediana"),
        ((20, 160, 20), "Zeleni krog = povprecje"),
        ((100, 100, 200), "Tocke = posamezni oglasi"),
    ], sc=2)

    return cv.to_png()


# ── Generiraj vse grafe ────────────────────────────────────────────────────────
print("Generiram grafe (visoka locljivost)…")

# Graf 1 – Histogram cen
g_hist_cen = chart_hist(
    cene, "Porazdelitev cen nepremicnin (Gorenjska, prodaja 2026)",
    "Cena (EUR)", "Stevilo oglasov", (70, 130, 180), n_bins=25)
print("  ✓ Graf 1: Histogram cen")

# Graf 2 – Histogram cena/m²
g_hist_cm2 = chart_hist(
    cm2s, "Porazdelitev cene na kvadratni meter",
    "Cena na m2 (EUR/m2)", "Stevilo oglasov", (192, 80, 77), n_bins=25)
print("  ✓ Graf 2: Histogram cena/m2")

# Graf 3 – Scatter površina vs cena
_pairs_m2 = [(r["VelikostM2"], r["Cena"]) for r in rows if r["VelikostM2"] and r["VelikostM2"] < 1000]
g_scatter_m2 = chart_scatter(
    [p[0] for p in _pairs_m2], [p[1] for p in _pairs_m2],
    "Povrsina hiše (m2) vs Prodajna cena (EUR)",
    "Povrsina (m2)", "Cena (EUR)", (70, 130, 180))
print("  ✓ Graf 3: Scatter povrsina-cena")

# Graf 4 – Scatter leto gradnje vs cena
_pairs_l = [(r["LetoGradnje"], r["Cena"]) for r in rows
            if r["LetoGradnje"] and 1950 <= r["LetoGradnje"] <= 2026]
g_scatter_l = chart_scatter(
    [p[0] for p in _pairs_l], [p[1] for p in _pairs_l],
    "Leto gradnje vs Prodajna cena (EUR)",
    "Leto gradnje", "Cena (EUR)", (155, 187, 89))
print("  ✓ Graf 4: Scatter leto-cena")

# Graf 5 – Bar chart lokacij (top 15 po mediani)
top_lok = sorted(lok.items(), key=lambda x: -statistics.median(x[1]))[:15]
g_bar_lok = chart_hbar(
    [l for l, _ in top_lok],
    [statistics.median(v) for _, v in top_lok],
    "Mediana prodajnih cen po lokacijah – top 15",
    xlabel="Mediana cene (EUR)")
print("  ✓ Graf 5: Bar lokacije")

# Graf 6 – Skupinski scatter po vrsti objekta
_grp_vrsta = [(v, c) for v, c in sorted(vrsta_d.items(), key=lambda x: -len(x[1]))]
g_vrsta_scatter = chart_group_scatter(
    _grp_vrsta,
    "Primerjava cen po vrsti objekta",
    ylabel="Cena (EUR)")
print("  ✓ Graf 6: Vrsta objekta")

# ══════════════════════════════════════════════════════════════════════════════
#  ML MODELI – predprocesiranje, 5 modelov, primerjava  (sekcija 7)
# ══════════════════════════════════════════════════════════════════════════════
print("Predprocesiram podatke za ML modele...")
random.seed(42)

_ml_rows = []
if CSV_PATH:
    with open(CSV_PATH, newline="", encoding="utf-8-sig") as _fh:
        for _rec in csv.DictReader(_fh, delimiter=";"):
            def _sf3(k, _r=_rec):
                try:
                    v = float(str(_r.get(k, "")).replace(",", "."))
                    return v if math.isfinite(v) else None
                except Exception:
                    return None
            _c = _sf3("Cena")
            if not _c or not (5_000 < _c < 5_000_000):
                continue
            _m2 = _sf3("VelikostM2")
            if _m2 and _m2 >= 5000:
                _m2 = None
            _leto = _sf3("LetoGradnje")
            _ml_rows.append({
                "Cena":        _c,
                "VelikostM2":  _m2,
                "ZemljisteM2": _sf3("ZemljisteM2"),
                "LetoGradnje": _leto if (_leto and _leto > 1900) else None,
                "StSob":       _sf3("StSob"),
                "VrstaObjekta": _rec.get("VrstaObjekta", "").strip(),
                "Lokacija":     (_rec.get("Obcina") or _rec.get("Lokacija", "")).strip(),
            })

print(f"   ML vzorcev: {len(_ml_rows)}")

def _mlmed(items, key):
    vals = [r[key] for r in items if r.get(key) is not None]
    return statistics.median(vals) if vals else 0.0

_ml_imp = {k: _mlmed(_ml_rows, k)
           for k in ["VelikostM2", "ZemljisteM2", "LetoGradnje", "StSob"]}

_vc3 = defaultdict(int)
for _rec3 in _ml_rows:
    _vc3[_rec3["VrstaObjekta"]] += 1
_vrsta_enc3 = {v: i for i, (v, _) in enumerate(sorted(_vc3.items(), key=lambda x: -x[1]))}

_lp3 = defaultdict(list)
for _rec3 in _ml_rows:
    if _rec3["Lokacija"]:
        _lp3[_rec3["Lokacija"]].append(_rec3["Cena"])
_lok_enc3 = {k: i for i, k in enumerate(
    sorted(_lp3, key=lambda k_: statistics.median(_lp3[k_])))}

_mlX, _mly = [], []
for _rec3 in _ml_rows:
    _x = [
        _rec3["VelikostM2"]  if _rec3["VelikostM2"]  is not None else _ml_imp["VelikostM2"],
        _rec3["ZemljisteM2"] if _rec3["ZemljisteM2"] is not None else _ml_imp["ZemljisteM2"],
        _rec3["LetoGradnje"] if _rec3["LetoGradnje"]  is not None else _ml_imp["LetoGradnje"],
        _rec3["StSob"]       if _rec3["StSob"]        is not None else _ml_imp["StSob"],
        float(_vrsta_enc3.get(_rec3["VrstaObjekta"], 0)),
        float(_lok_enc3.get(_rec3["Lokacija"], len(_lok_enc3) // 2)),
    ]
    _mlX.append(_x)
    _mly.append(_rec3["Cena"])

def _ml_std(X, means=None, stds=None):
    p = len(X[0])
    if means is None:
        means = [statistics.mean(r[j] for r in X) for j in range(p)]
    if stds is None:
        stds  = [(statistics.stdev(r[j] for r in X) or 1.0) for j in range(p)]
    return [[(r[j] - means[j]) / stds[j] for j in range(p)] for r in X], means, stds

_idx3   = list(range(len(_mlX))); random.shuffle(_idx3)
_cut3   = int(len(_idx3) * 0.8)
_Xtr3   = [_mlX[i] for i in _idx3[:_cut3]];  _ytr3 = [_mly[i] for i in _idx3[:_cut3]]
_Xte3   = [_mlX[i] for i in _idx3[_cut3:]];  _yte3 = [_mly[i] for i in _idx3[_cut3:]]
_Xtr3, _fm3, _fs3 = _ml_std(_Xtr3)
_Xte3, _,   _     = _ml_std(_Xte3, _fm3, _fs3)

# ── Gauss eliminacija ─────────────────────────────────────────────────────────
def _gauss_ml(A, n):
    for c in range(n):
        piv = max(range(c, n), key=lambda rr: abs(A[rr][c]))
        A[c], A[piv] = A[piv], A[c]
        if abs(A[c][c]) < 1e-10:
            continue
        for rr in range(n):
            if rr != c:
                f = A[rr][c] / A[c][c]
                A[rr] = [A[rr][k] - f * A[c][k] for k in range(n + 1)]
    return [A[j][n] / A[j][j] if abs(A[j][j]) > 1e-10 else 0.0 for j in range(n)]

# ── 1. LINEARNA REGRESIJA (OLS) ───────────────────────────────────────────────
class _OLS3:
    name = "Linearna reg. (OLS)"
    def fit(self, X, y):
        n, p = len(X), len(X[0])
        Xa = [[1.0] + list(r) for r in X]; q = p + 1
        XtX = [[sum(Xa[i][rr] * Xa[i][cc] for i in range(n))
                for cc in range(q)] for rr in range(q)]
        Xty = [sum(Xa[i][j] * y[i] for i in range(n)) for j in range(q)]
        A = [XtX[rr][:] + [Xty[rr]] for rr in range(q)]
        beta = _gauss_ml(A, q)
        self._b0, self._w = beta[0], beta[1:]
        return self
    def predict(self, X):
        return [self._b0 + sum(x[j] * self._w[j] for j in range(len(x))) for x in X]

# ── 2. RIDGE REGRESIJA ────────────────────────────────────────────────────────
class _Ridge3:
    name = "Ridge regresija"
    def __init__(self, alpha=10.0):
        self.alpha = alpha
    def fit(self, X, y):
        n, p = len(X), len(X[0])
        ym = statistics.mean(y); ys = [v - ym for v in y]
        XtX = [[sum(X[i][rr] * X[i][cc] for i in range(n))
                for cc in range(p)] for rr in range(p)]
        for j in range(p):
            XtX[j][j] += self.alpha
        Xty = [sum(X[i][j] * ys[i] for i in range(n)) for j in range(p)]
        A = [XtX[rr][:] + [Xty[rr]] for rr in range(p)]
        self._w = _gauss_ml(A, p); self._b0 = ym
        return self
    def predict(self, X):
        return [self._b0 + sum(x[j] * self._w[j] for j in range(len(x))) for x in X]

# ── 3. ODLOČITVENO DREVO (CART) ───────────────────────────────────────────────
class _DTNode3:
    __slots__ = ("feat", "thresh", "left", "right", "val")
    def __init__(self, val=None, feat=None, thresh=None, left=None, right=None):
        self.val = val; self.feat = feat; self.thresh = thresh
        self.left = left; self.right = right

class _DT3:
    name = "Odlocitveno drevo"
    def __init__(self, max_depth=5, min_s=4, max_feat=None):
        self.max_depth = max_depth; self.min_s = min_s; self.max_feat = max_feat
    def fit(self, X, y):
        self._root = self._build(X, y, 0); return self
    def _split(self, X, y):
        n, p = len(X), len(X[0]); bg = -1.0; bf = bt = None
        tv = sum(v ** 2 for v in y) / n - (sum(y) / n) ** 2
        feats = (random.sample(list(range(p)), self.max_feat)
                 if self.max_feat and self.max_feat < p else list(range(p)))
        for fi in feats:
            order = sorted(range(n), key=lambda ii: X[ii][fi])
            psy = [0.0] * (n + 1); psy2 = [0.0] * (n + 1)
            for k in range(n):
                ii = order[k]; psy[k + 1] = psy[k] + y[ii]; psy2[k + 1] = psy2[k] + y[ii] ** 2
            for k in range(1, n):
                if X[order[k]][fi] == X[order[k - 1]][fi]:
                    continue
                nl, nr = k, n - k
                if nl < 2 or nr < 2:
                    continue
                sl, s2l = psy[k], psy2[k]; sr, s2r = psy[n] - sl, psy2[n] - s2l
                vl = s2l / nl - (sl / nl) ** 2; vr = s2r / nr - (sr / nr) ** 2
                g = tv - (nl * vl + nr * vr) / n
                if g > bg:
                    bg = g; bf = fi
                    bt = (X[order[k - 1]][fi] + X[order[k]][fi]) / 2
        return bf, bt
    def _build(self, X, y, d):
        val = sum(y) / len(y)
        if len(y) < self.min_s or (self.max_depth and d >= self.max_depth):
            return _DTNode3(val=val)
        f, t = self._split(X, y)
        if f is None:
            return _DTNode3(val=val)
        mask = [X[ii][f] <= t for ii in range(len(X))]
        Xl = [X[ii] for ii in range(len(X)) if mask[ii]]
        yl = [y[ii] for ii in range(len(y)) if mask[ii]]
        Xr = [X[ii] for ii in range(len(X)) if not mask[ii]]
        yr = [y[ii] for ii in range(len(y)) if not mask[ii]]
        if not Xl or not Xr:
            return _DTNode3(val=val)
        return _DTNode3(val=val, feat=f, thresh=t,
                        left=self._build(Xl, yl, d + 1),
                        right=self._build(Xr, yr, d + 1))
    def _p1(self, nd, x):
        if nd.feat is None:
            return nd.val
        return self._p1(nd.left if x[nd.feat] <= nd.thresh else nd.right, x)
    def predict(self, X):
        return [self._p1(self._root, x) for x in X]

# ── 4. NAKLJUČNI GOZD ─────────────────────────────────────────────────────────
class _RF3:
    name = "Nakljucni gozd"
    def __init__(self, n_trees=12, max_depth=5):
        self.n_trees = n_trees; self.max_depth = max_depth; self._trees = []
    def fit(self, X, y):
        n, p = len(X), len(X[0]); mf = max(1, int(p ** 0.5))
        self._trees = []
        for ti in range(self.n_trees):
            idx = [random.randint(0, n - 1) for _ in range(n)]
            self._trees.append(
                _DT3(max_depth=self.max_depth, max_feat=mf)
                .fit([X[ii] for ii in idx], [y[ii] for ii in idx]))
            if (ti + 1) % 4 == 0:
                print(f"    RF: {ti + 1}/{self.n_trees} dreves", flush=True)
        return self
    def predict(self, X):
        all_p = [t.predict(X) for t in self._trees]
        return [statistics.mean(all_p[j][i]
                for j in range(len(self._trees))) for i in range(len(X))]

# ── 5. NEVRONSKA MREŽA (MLP, 1 skrita plast, ReLU, mini-batch SGD) ───────────
class _MLP3:
    name = "Nevronska mreza (MLP)"
    def __init__(self, hidden=20, epochs=25, lr=0.01, batch=256):
        self.hidden = hidden; self.epochs = epochs
        self.lr = lr; self.batch = batch
    def fit(self, X, y):
        n, p = len(X), len(X[0]); h = self.hidden
        s1 = (2.0 / p) ** 0.5; s2 = (2.0 / h) ** 0.5
        self.W1 = [[random.gauss(0, s1) for _ in range(p)] for _ in range(h)]
        self.b1 = [0.0] * h
        self.W2 = [random.gauss(0, s2) for _ in range(h)]
        self.b2 = 0.0
        ym = statistics.mean(y)
        ys_ = statistics.stdev(y) if len(y) > 1 else 1.0
        self._ym = ym; self._ys = ys_
        yn = [(v - ym) / ys_ for v in y]
        lr = self.lr
        for ep in range(self.epochs):
            idx = list(range(n)); random.shuffle(idx)
            for bs in range(0, n, self.batch):
                bi = idx[bs: bs + self.batch]; nb = len(bi)
                dW1 = [[0.0] * p for _ in range(h)]
                db1 = [0.0] * h; dW2 = [0.0] * h; db2 = 0.0
                for ii in bi:
                    xi = X[ii]
                    z1 = [sum(self.W1[j][k] * xi[k] for k in range(p)) + self.b1[j]
                          for j in range(h)]
                    a1 = [max(0.0, v) for v in z1]
                    z2 = sum(self.W2[j] * a1[j] for j in range(h)) + self.b2
                    err = z2 - yn[ii]; db2 += err
                    for j in range(h):
                        dW2[j] += err * a1[j]
                        if z1[j] > 0:
                            d_ = err * self.W2[j]; db1[j] += d_
                            for k in range(p):
                                dW1[j][k] += d_ * xi[k]
                for j in range(h):
                    self.b1[j] -= lr * db1[j] / nb
                    self.W2[j] -= lr * dW2[j] / nb
                    for k in range(p):
                        self.W1[j][k] -= lr * dW1[j][k] / nb
                self.b2 -= lr * db2 / nb
            if (ep + 1) % 10 == 0:
                lr *= 0.6
                print(f"    MLP: ep {ep + 1}/{self.epochs}  lr={lr:.5f}", flush=True)
        return self
    def predict(self, X):
        h = self.hidden; res = []
        for xi in X:
            z1 = [sum(self.W1[j][k] * xi[k] for k in range(len(xi))) + self.b1[j]
                  for j in range(h)]
            a1 = [max(0.0, v) for v in z1]
            z2 = sum(self.W2[j] * a1[j] for j in range(h)) + self.b2
            res.append(z2 * self._ys + self._ym)
        return res

# ── Metrike ────────────────────────────────────────────────────────────────────
def _ml_r2(yt, yp):
    my = statistics.mean(yt)
    sst = sum((v - my) ** 2 for v in yt)
    ssr = sum((t - p) ** 2 for t, p in zip(yt, yp))
    return 1.0 - ssr / sst if sst else 0.0

def _ml_mae(yt, yp):
    return statistics.mean(abs(t - p) for t, p in zip(yt, yp))

def _ml_rmse(yt, yp):
    return math.sqrt(statistics.mean((t - p) ** 2 for t, p in zip(yt, yp)))

# ── Trening ────────────────────────────────────────────────────────────────────
_ml_results = []
if _ml_rows:
    for _mod_obj, _mod_label in [
        (_OLS3(),                           "OLS"),
        (_Ridge3(alpha=10.0),               "Ridge"),
        (_DT3(max_depth=6),                 "Odloc. drevo"),
        (_RF3(n_trees=12, max_depth=5),     "Nakl. gozd"),
        (_MLP3(hidden=20, epochs=25),       "MLP"),
    ]:
        print(f"  Treniram {_mod_label}...", flush=True)
        _mod_obj.fit(_Xtr3, _ytr3)
        _pp = _mod_obj.predict(_Xte3)
        _ml_results.append({
            "name":  _mod_obj.name,
            "R2":    _ml_r2(_yte3, _pp),
            "MAE":   _ml_mae(_yte3, _pp),
            "RMSE":  _ml_rmse(_yte3, _pp),
            "preds": _pp,
        })
        print(f"  ✓ {_mod_obj.name}: R²={_ml_results[-1]['R2']:.4f}  "
              f"MAE={_ml_results[-1]['MAE']:,.0f} €", flush=True)

    _best_ml = max(_ml_results, key=lambda x: x["R2"])
    print(f"  🏆 Najboljši model: {_best_ml['name']}  (R²={_best_ml['R2']:.4f})")
else:
    _best_ml = None

# ── Grafi primerjave modelov ───────────────────────────────────────────────────
def chart_ml_bar(results, metric, title, higher_better=True):
    """Horizontalni stolpičar: primerjava ML modelov po metriki."""
    if not results:
        return None
    names = [r["name"] for r in results]
    vals  = [r[metric] for r in results]
    n     = len(names)
    bar_h = 60
    W_c   = 1400; ML_c = 400; MR_c = 100; MT_c = 80; MB_c = 70
    H_c   = MT_c + n * (bar_h + 16) + MB_c
    CW_c  = W_c - ML_c - MR_c
    vmn   = min(0.0, min(vals)) if higher_better else min(vals) * 0.88
    vmx   = max(vals) * 1.18
    if vmn == vmx: vmx = vmn + 1
    palette = [(33,150,243),(76,175,80),(255,87,34),(156,39,176),(0,188,212)]
    best_i  = (vals.index(max(vals)) if higher_better else vals.index(min(vals)))
    cv_c    = Canvas(W_c, H_c, (255, 255, 255))
    cv_c.fill_rect(ML_c, MT_c, CW_c, H_c - MT_c - MB_c, (252, 252, 254))
    for xt in _nice_ticks(vmn, vmx, 5):
        xp_c = ML_c + int((xt - vmn) / (vmx - vmn) * CW_c)
        if ML_c <= xp_c <= ML_c + CW_c:
            cv_c.vline(xp_c, MT_c, H_c - MB_c, (220, 220, 228), 1)
            cv_c.text_centered(xp_c, H_c - MB_c + 14, _fmt(xt), (80, 80, 90), 2)
    for ii, (nm, v) in enumerate(zip(names, vals)):
        by_c = MT_c + ii * (bar_h + 16) + 8
        bw_c = max(4, int((v - vmn) / (vmx - vmn) * CW_c))
        col  = palette[ii % len(palette)]
        if ii == best_i:
            cv_c.fill_rect(ML_c - 4, by_c - 4, bw_c + 8, bar_h + 8, (255, 200, 0))
        cv_c.fill_rect(ML_c, by_c, bw_c, bar_h, col)
        cv_c.fill_rect(ML_c, by_c, bw_c, bar_h // 5,
                       tuple(min(255, c + 50) for c in col))
        lbl_v = f"{v:.4f}" if abs(v) < 10 else _fmt(v)
        cv_c.text(ML_c + bw_c + 12, by_c + bar_h // 2 - 7, lbl_v, (40, 40, 50), 2)
        cv_c.text_right(ML_c - 14, by_c + bar_h // 2 - 7,
                        _asc(nm[:40]), (30, 30, 40), 2)
    cv_c.hline(ML_c, ML_c + CW_c, H_c - MB_c, (60, 60, 70), 3)
    cv_c.vline(ML_c, MT_c, H_c - MB_c, (60, 60, 70), 3)
    cv_c.text_centered(W_c // 2, 18, _asc(title), (20, 20, 30), 3)
    cv_c.legend_box(ML_c + CW_c - 320, MT_c + 14,
                    [((255, 200, 0), "Najboljsi model")], sc=2)
    return cv_c.to_png()


def chart_ml_pred(y_true, y_pred, model_name):
    """Scatter: dejanske vs napovedane vrednosti + idealna premica."""
    if len(y_true) < 3:
        return None
    mn_p = min(min(y_true), min(y_pred))
    mx_p = max(max(y_true), max(y_pred))
    pad  = (mx_p - mn_p) * 0.04; mn_p -= pad; mx_p += pad
    if mn_p == mx_p: mx_p = mn_p + 1
    W_p = 900; H_p = 800; ML_p = 100; MR_p = 40; MT_p = 80; MB_p = 100
    CW_p = W_p - ML_p - MR_p; CH_p = H_p - MT_p - MB_p
    def _xp(v): return ML_p + int((v - mn_p) / (mx_p - mn_p) * CW_p)
    def _yp(v): return H_p - MB_p - int((v - mn_p) / (mx_p - mn_p) * CH_p)
    cv_p = Canvas(W_p, H_p, (255, 255, 255))
    cv_p.fill_rect(ML_p, MT_p, CW_p, CH_p, (252, 252, 254))
    for t in _nice_ticks(mn_p, mx_p, 5):
        tp = _xp(t)
        if ML_p <= tp <= ML_p + CW_p:
            cv_p.vline(tp, MT_p, H_p - MB_p, (220, 220, 228), 1)
            cv_p.text_centered(tp, H_p - MB_p + 16, _fmt(t), (80, 80, 90), 2)
        tp = _yp(t)
        if MT_p <= tp <= H_p - MB_p:
            cv_p.hline(ML_p, ML_p + CW_p, tp, (220, 220, 228), 1)
            cv_p.text_right(ML_p - 12, tp - 7, _fmt(t), (80, 80, 90), 2)
    cv_p.bresenham(_xp(mn_p), _yp(mn_p), _xp(mx_p), _yp(mx_p), (210, 50, 50), 3)
    for yt_v, yp_v in zip(y_true, y_pred):
        cv_p.circle(_xp(yt_v), _yp(yp_v), 4, (33, 150, 243))
    cv_p.hline(ML_p, ML_p + CW_p, H_p - MB_p, (60, 60, 70), 3)
    cv_p.vline(ML_p, MT_p, H_p - MB_p, (60, 60, 70), 3)
    rv = _ml_r2(list(y_true), list(y_pred))
    cv_p.text_centered(W_p // 2, 20,
                       _asc(f"Napovedano vs Dejansko – {model_name}"), (20, 20, 30), 3)
    cv_p.text_centered(W_p // 2, H_p - 26,
                       _asc("Dejanska cena (EUR)"), (50, 50, 60), 2)
    for ii_c, ch_c in enumerate(_asc("Napovedano")[:12]):
        cv_p.text(8, MT_p + ii_c * 18, ch_c, (50, 50, 60), 2)
    cv_p.legend_box(ML_p + 20, MT_p + 20, [
        ((210, 50, 50), "Idealna premica y=x"),
        ((33, 150, 243), f"Napovedi  R2={rv:.4f}"),
    ], sc=2)
    return cv_p.to_png()


def chart_ml_residuals(y_true, y_pred, model_name):
    """Histogram ostankov (dejanska − napovedana cena)."""
    res = [yt_v - yp_v for yt_v, yp_v in zip(y_true, y_pred)]
    return chart_hist(
        res,
        f"Porazdelitev ostankov – {model_name}",
        "Ostanek (EUR)", "Stevilo", (100, 160, 220), n_bins=25)


# ── Generiraj ML grafe ────────────────────────────────────────────────────────
if _ml_results:
    print("  Generiram grafe primerjave modelov...")
    _g_ml_r2   = chart_ml_bar(_ml_results, "R2",
                               "Primerjava R2 – vsi modeli (visje = boljse)",
                               higher_better=True)
    _g_ml_mae  = chart_ml_bar(_ml_results, "MAE",
                               "Primerjava MAE v EUR – vsi modeli (nizje = boljse)",
                               higher_better=False)
    _g_ml_rmse = chart_ml_bar(_ml_results, "RMSE",
                               "Primerjava RMSE v EUR – vsi modeli (nizje = boljse)",
                               higher_better=False)
    _g_ml_pred = chart_ml_pred(_yte3, _best_ml["preds"], _best_ml["name"])
    _g_ml_res  = chart_ml_residuals(_yte3, _best_ml["preds"], _best_ml["name"])
    print("  ✓ Grafi ML primerjave generirani")
else:
    _g_ml_r2 = _g_ml_mae = _g_ml_rmse = _g_ml_pred = _g_ml_res = None

# ── Dokument ───────────────────────────────────────────────────────────────────
print("Gradim DOCX…")
doc = Document()

# Robovi
for sec in doc.sections:
    sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)
    sec.top_margin  = Cm(2.5); sec.bottom_margin= Cm(2.5)

# ── Pomočniki ──────────────────────────────────────────────────────────────────
def H(lvl, txt):
    h = doc.add_heading(txt, lvl)
    if lvl == 0: h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return h

def P(txt="", bold=False, italic=False, size=None, color=None, align=None):
    p = doc.add_paragraph()
    if align: p.alignment = align
    run = p.add_run(txt)
    run.bold = bold; run.italic = italic
    if size:  run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor(*color)
    return p

def CODE(txt):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(txt)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30,30,30)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'F5F5F5')
    p._p.get_or_add_pPr().append(shd)
    return p

def TBL(headers, rows_data, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        c = hrow.cells[i]; c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.size = Pt(9)
        tc = c._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:fill'),'4472C4')
        tcPr.append(shd)
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
        if col_widths: c.width = Cm(col_widths[i])
    for row in rows_data:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
            cells[i].paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph()
    return t

def IMG(png_bytes, caption, w=5.8):
    if not png_bytes: return
    doc.add_picture(io.BytesIO(png_bytes), width=Inches(w))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp = doc.add_paragraph(caption)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cp.runs: run.italic = True; run.font.size = Pt(9)
    doc.add_paragraph()

def BULLET(items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item).font.size = Pt(10)

def NOTE(txt):
    p = doc.add_paragraph()
    run = p.add_run(f"ℹ  {txt}")
    run.italic = True; run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100,100,100)
    p.paragraph_format.left_indent = Cm(0.5)

# ══════════════════════════════════════════════════════════════════════════════
# NASLOV
# ══════════════════════════════════════════════════════════════════════════════
H(0, "Dokumentacija projekta – Nepremičnine.net Scraper & Analiza")
P(f"UUI – Laboratorijska vaja 2  |  Datum: 6. april 2026",
  italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=(100,100,100))
P(f"Generirano: {datetime.now().strftime('%d.%m.%Y %H:%M')}  |  Vir: nepremicnine.net",
  italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=(130,130,130))
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 1. PREGLED PROJEKTA
# ══════════════════════════════════════════════════════════════════════════════
H(1, "1. Pregled projekta")
P("Projekt sestoji iz spletnega scraperja za portal nepremicnine.net, analize zbranih podatkov "
  "in napovednega modela cen nepremičnin. Celoten sistem je implementiran v Pythonu s poudarkom "
  "na uporabi izključno standardne knjižnice za ML (brez NumPy/Pandas/scikit-learn).")
doc.add_paragraph()
P("Tok podatkov:", bold=True)
CODE("nepremicnine.net  →  scraper.py  →  CSV  →  analyze.py  →  grafi / DOCX\n"
     "                                     ↓\n"
     "                                 modeli.py  →  ML porocilo DOCX\n"
     "                                     ↓\n"
     "                                 cenik.py   →  napoved cene (Random Forest)")
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 2. ARHITEKTURA
# ══════════════════════════════════════════════════════════════════════════════
H(1, "2. Arhitektura – datoteke in moduli")
TBL(
    ["Datoteka", "Opis", "Vrstice"],
    [
        ["scraper.py",   "Pridobivanje oglasov s spletišča z obhodom Cloudflare", "693"],
        ["gui.py",       "Tkinter grafični vmesnik za vse komponente", "1202"],
        ["analyze.py",   "Statistična analiza + SVG grafi + DOCX", "823"],
        ["modeli.py",    "4 ML modeli od začetka + DOCX poročilo", "932"],
        ["cenik.py",     "Random Forest napovednik za en oglas", "281"],
        ["nepremicnine_export.csv",         "Manjša testna zbirka (246 oglasov)", "247"],
        ["nepremicnine_export_prodaja.csv", "Polna zbirka (8.953 oglasov, vse regije)", "8954"],
    ],
    col_widths=[5.5, 9.0, 2.0]
)

# ══════════════════════════════════════════════════════════════════════════════
# 3. OPIS KODE
# ══════════════════════════════════════════════════════════════════════════════
H(1, "3. Opis glavnih delov kode")

# 3.1
H(2, "3.1  scraper.py – pridobivanje podatkov")
H(3, "Konfiguracija REGIJE / VRSTE / AKCIJE")
P("Slovarji preslikajo prikazna imena v URL slug-e. URL za oglas se sestavi kot:")
CODE("/oglasi-{akcija}/{regija}/{vrsta}/{stran}/\n\n"
     "REGIJE = { 'LJ-mesto': 'ljubljana', 'Gorenjska': 'gorenjska', ... }  # 13 regij\n"
     "VRSTE  = { 'Hisa': 'hisa', 'Stanovanje': 'stanovanje', ... }         # 7 vrst")

H(3, "Brskalnik – get_browser()")
P("DrissionPage zažene pravi Chrome v stealth načinu:")
CODE("co = ChromiumOptions()\n"
     "co.set_argument('--disable-blink-features=AutomationControlled')\n"
     "co.set_argument('--lang=sl-SI')\n"
     "co.set_pref('intl.accept_languages', 'sl-SI,sl,en-US,en')\n"
     "browser = ChromiumPage(addr_or_opts=co)")
BULLET([
    "AutomationControlled flag onemogočen → brskalnik ne izda avtomatizacije",
    "Jezik nastavljen na slovenščino → posnemanje normalnega obiskovalca",
    "Vidno okno (ne headless) → lažje preide Cloudflare challenge",
])

H(3, "Čakanje na CF – _wait_for_cf()")
P("Funkcija preverja zanesljive indikatorje CF challenge strani:")
CODE("is_challenge = (\n"
     "    'just a moment' in title           # naslov CF challeng strani\n"
     "    or 'id=\"cf-challenge-running\"' in html\n"
     "    or 'name=\"cf-turnstile-response\"' in html\n"
     ")")
NOTE("Napaka: preverjanje 'cloudflare' v html.lower() ne deluje – beseda je v footerju "
     "vsake zaščitene strani, tudi po uspešnem nalaganju!")

H(3, "Parsanje – parse_listing_page()")
BULLET([
    "HTML razčlenimo z BeautifulSoup + lxml",
    "Vsak oglas je <div itemprop=\"item\"> (schema.org format)",
    "Cena je v <meta itemprop=\"price\">",
    "Podrobnosti so v <ul itemprop=\"disambiguatingDescription\"> → ikone slik",
    "_get_total_pages() prebere skupno število strani iz <ul data-pages=\"N\">",
])

H(3, "Izvoz CSV – export_csv()")
P("Ločilo: ; | Enkodiranje: UTF-8-sig (Excel kompatibilnost)")
CODE("Naslov;Lokacija;Obcina;Cena;VelikostM2;ZemljisteM2;StSob;\n"
     "LetoGradnje;EnergetskiRazred;CenaNaM2;VrstaObjekta;DatumScrapa;Url")
doc.add_paragraph()

# 3.2
H(2, "3.2  gui.py – grafični vmesnik")
P("GUI je zgrajen z Tkinter v temni barvni shemi (#1e1e2e). Sestoji iz 5 komponent:")
TBL(
    ["Komponenta", "Opis"],
    [
        ["ScraperGUI", "Glavno okno: leva plošča (nastavitve) + desna plošča (barvni log)"],
        ["CheckList",  "Scrollable seznam checkboxov z gumbom 'Vse' za hiter izbor"],
        ["AnalysisDialog", "Modalni dialog za konfiguracijo analize (CSV, grafi, DOCX)"],
        ["MLDialog",   "Dialog za ML modele: učna množica %, seme, 4 modeli"],
        ["CenikDialog","Napovednik cen: spustni meniji, Spinbox, 90% interval zaupanja"],
    ],
    col_widths=[4.5, 12.0]
)
P("Večnitnost: vsi procesi tečejo v daemon nitih → GUI ostane odziven. "
  "Komunikacija GUI↔subprocess poteka prek subprocess.Popen z zajemanjem stdout.")
doc.add_paragraph()

# 3.3
H(2, "3.3  analyze.py – statistična analiza")
BULLET([
    "Filtrira cene izven razpona 5.000–5.000.000 € (outlier removal)",
    "Površine ≥ 5.000 m² se odstranijo",
    "Za vsak stolpec: N, Min, Max, Povprečje, Mediana, Std. dev.",
    "Pearsonov korelacijski koeficient za vsak par značilk s ceno",
])
P("SVG grafi (brez matplotlib – direktno XML):", bold=True)
BULLET([
    "svg_hist_chart() – histogram s črto mediane in povprečja",
    "svg_scatter() – razsevni diagram z regresijsko premico",
    "svg_bar_chart() – stolpičar lokacij po mediani cene",
])
P("PNG za DOCX: pixel canvas _CV z Bresenhamovim algoritmom in bitmap pisavo 5×7 pik. "
  "PNG sestavljen ročno z struct + zlib (brez Pillow).")
doc.add_paragraph()

# 3.4
H(2, "3.4  modeli.py – strojno učenje")
H(3, "Predprocesiranje")
TBL(
    ["Korak", "Opis"],
    [
        ["Imputacija", "Manjkajoče vrednosti → mediana (VelikostM2, LetoGradnje, StSob, EnergetRazred)"],
        ["Kodiranje VrstaObjekta", "Ordinalna koda po padajoči frekvenci"],
        ["Kodiranje Obcina", "Koda po naraščajočih medianah cen (bogata = visoka koda)"],
        ["Standardizacija", "Z-score, izračunan samo na učni množici (prepreči data leakage)"],
    ],
    col_widths=[4.5, 12.0]
)
P("Vektorji značilk (7 dimenzij):", bold=True)
CODE("[VelikostM2, ZemljisteM2, LetoGradnje, StSob,\n"
     " EnergetRazred, VrstaObjekta_enc, Obcina_enc]")

H(3, "Implementirani modeli")
TBL(
    ["Model", "Metoda", "Hiperparametri"],
    [
        ["Linearna regresija (OLS)", "β = (XᵀX)⁻¹ Xᵀy  – normalne enačbe", "—  (referenčna)"],
        ["Ridge regresija", "β = (XᵀX + αI)⁻¹ Xᵀy  – L2 regularizacija", "α ∈ {0.01,0.1,1,10,100,500,1000}"],
        ["Odločitveno drevo (CART)", "Rekurzivno deli po min. MSE (variacijski dobiček)", "max_depth ∈ {2,3,4,5,6,8}"],
        ["Naključni gozd", "Bootstrap + naključni izbor √p značilk per split", "n ∈ {10,20}, depth ∈ {3,5,7}"],
    ],
    col_widths=[4.0, 8.0, 4.5]
)
P("Optimizacija: Grid Search + 5-kratna prečna validacija (CV) na učni množici.")
P("Metrike vrednotenja: R² | MAE (€) | RMSE (€)")
doc.add_paragraph()

# 3.5
H(2, "3.5  cenik.py – napovednik cen")
BULLET([
    "Trening: Random Forest s 60 drevesi, globina 7, enako predprocesiranje kot modeli.py",
    "Napoved: povprečje napovedi vseh dreves + 90% interval (napoved ± 1.64 × std dreves)",
    "Podobni oglasi: progresivni filter (vrsta+kraj → vrsta → kraj → vsi)",
    "Izhod: KEY=VALUE format za razčlenjevanje v GUI (NAPOVEDANA, CI_MIN, CI_MAX, ...)",
])
CODE("tree_preds = rf.predict_all(x_std)   # napovedi vseh 60 dreves\n"
     "napoved    = mean(tree_preds)\n"
     "std_pred   = stdev(tree_preds)\n"
     "ci_min     = napoved - 1.64 * std_pred  # 90% spodnja meja\n"
     "ci_max     = napoved + 1.64 * std_pred  # 90% zgornja meja")
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 4. CLOUDFLARE
# ══════════════════════════════════════════════════════════════════════════════
H(1, "4. Problem Cloudflare Turnstile (CAPTCHA)")
P("Portal nepremicnine.net je zaščiten z Cloudflare Turnstile – napredno CAPTCHA zaščito, "
  "ki analizira vedenje brskalnika, browser fingerprint in JavaScript izvajalno okolje.")
doc.add_paragraph()

H(2, "4.1  Preizkušeni pristopi, ki niso delovali")
TBL(
    ["Pristop", "Razlog neuspeha"],
    [
        ["requests + BeautifulSoup",
         "Ne izvaja JS → CF zahteva JS. navigator.webdriver = true. 403 Forbidden."],
        ["Playwright (headless=True)",
         "Headless fingerprint drugačen. CF zazna 'HeadlessChrome' v UA. navigator.plugins prazen."],
        ["Selenium + lažni User-Agent",
         "Selenium ne more skriti navigator.webdriver flaga. AutomationControlled viden."],
        ["Detekcija: 'cloudflare' in html",
         "NAPAKA: beseda 'cloudflare' je v footerju VSAKE zaščitene strani (tudi po uspešnem nalaganju)!"],
    ],
    col_widths=[5.0, 11.5]
)
doc.add_paragraph()

H(2, "4.2  Pravilna rešitev – DrissionPage (stealth Chrome)")
P("DrissionPage zagotavlja pravi Chrome (ne headless) z onemogočenimi automation flagmi:")
CODE("co = ChromiumOptions()\n"
     "co.set_argument('--disable-blink-features=AutomationControlled')\n"
     "co.set_argument('--lang=sl-SI')\n"
     "co.set_argument('--window-size=1400,900')\n"
     "browser = ChromiumPage(addr_or_opts=co)")
BULLET([
    "Pravi Chrome izvaja JavaScript → CF Turnstile se reši samodejno v 2–5 sekundah",
    "AutomationControlled flag onemogočen",
    "Vidno okno (ne headless) → drugačen browser fingerprint",
    "Slovenščina → izgleda kot lokalni obiskovalec",
])
P("Zamiki med zahtevki (preprečevanje rate-limitinga):", bold=True)
CODE("time.sleep(delay + random.uniform(0, 0.8))  # privzeto: 1.5s + 0-0.8s naklj.")
doc.add_paragraph()

H(2, "4.3  Pravilna detekcija CF challenge strani")
CODE("is_challenge = (\n"
     "    'just a moment' in title.lower()          # naslov CF challeng strani\n"
     "    or 'id=\"cf-challenge-running\"' in html  # aktivni izziv\n"
     "    or 'id=\"cf-challenge-body\"' in html     # telo izziva\n"
     "    or 'name=\"cf-turnstile-response\"' in html\n"
     ")")
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 5. MODULI IN KNJIŽNICE
# ══════════════════════════════════════════════════════════════════════════════
H(1, "5. Moduli in knjižnice")

H(2, "5.1  Zunanje knjižnice")
TBL(
    ["Knjižnica", "Namen", "Zakaj sva jo dodala", "Alternativa"],
    [
        ["DrissionPage", "Stealth Chrome brskalnik",
         "Edina preprosta rešitev za CF Turnstile", "selenium-stealth + undetected-chromedriver"],
        ["beautifulsoup4 + lxml", "Razčlenjevanje HTML",
         "lxml je hitrejši in odpuščajoč na napačen HTML", "html.parser (počasnejši)"],
        ["python-docx", "Ustvarjanje .docx dokumentov",
         "Zahteva za izvoz poročil v format za oddajo", "Ročno ZIP+XML (preveč kompleksno)"],
    ],
    col_widths=[3.5, 3.5, 5.5, 4.0]
)
P("Namestitev: pip install DrissionPage beautifulsoup4 lxml python-docx", italic=True)

H(2, "5.2  Standardna knjižnica Pythona (brez zunanjih odvisnosti)")
TBL(
    ["Modul", "Uporaba v projektu"],
    [
        ["csv",           "Branje/pisanje CSV z ';' ločilom in UTF-8-sig enkodiranjem"],
        ["statistics",    "Mediana, povprečje, standardni odklon (imputacija, napoved)"],
        ["math",          "sqrt, isfinite, floor (validacija vrednosti)"],
        ["random",        "Bootstrap vzorčenje v RF, mešanje podatkov, semena"],
        ["argparse",      "CLI argumenti za vse skripte (--csv, --regije, --vrsta ...)"],
        ["tkinter + ttk", "Grafični vmesnik (CheckList, Dialogi, barvni Log)"],
        ["subprocess",    "Zagon Python skriptov iz GUI-ja"],
        ["threading",     "Daemon niti – GUI ostane odziven med scrapanjem"],
        ["struct + zlib", "Ročno generiranje PNG datotek za DOCX (brez Pillow)"],
        ["collections.defaultdict", "Grupiranje oglasov po lokaciji / vrsti"],
        ["itertools.product", "Kartezični produkt za grid search hiperparametrov"],
        ["re",            "Regularni izrazi za parsanje cen in datumov"],
        ["datetime",      "Datum scrapa, generiranje imen datotek"],
    ],
    col_widths=[4.5, 12.0]
)
NOTE("Zavestna odločitev: NO NumPy/Pandas/scikit-learn → boljše razumevanje algoritmov, "
     "ni odvisnosti od velikih knjižnic, demonstracija matematičnega znanja.")
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 6. PODATKI – CSV ANALIZA
# ══════════════════════════════════════════════════════════════════════════════
H(1, "6. Podatki – CSV analiza in grafi")

H(2, "6.1  Vir in obseg podatkov")
TBL(
    ["Parameter", "Vrednost"],
    [
        ["Stran", "nepremicnine.net"],
        ["Regija / pokritost", "Gorenjska (manjša zbirka) | Vse regije (polna zbirka)"],
        ["Vrsta nepremičnine", "Vse hiše: samostojna, dvojček, vrstna, trojček, dvostanovanjska"],
        ["Akcija", "Prodaja"],
        ["Datum zbiranja", "6. april 2026"],
        ["Datoteka (manjša)", f"nepremicnine_export.csv – 246 oglasov"],
        ["Datoteka (polna)", f"nepremicnine_export_prodaja.csv – 8.953 oglasov"],
        ["Analizirani vzorci", str(len(rows))],
    ],
    col_widths=[5.5, 11.0]
)
NOTE("Vzrok za razliko: ko scraper dobi več kombinacij (regije × vrste), poimenuje izhod "
     "nepremicnine_export_prodaja.csv. Vsi skripti so posodobljeni da dajo prednost tej datoteki.")

H(2, "6.2  Osnovna statistika")
if rows:
    TBL(
        ["Spremenljivka", "N", "Min", "Max", "Povprečje", "Mediana"],
        [
            ["Cena (€)", len(cene),
             f"{min(cene):,.0f}", f"{max(cene):,.0f}",
             f"{statistics.mean(cene):,.0f}", f"{statistics.median(cene):,.0f}"],
            ["Površina (m²)", len(m2s),
             f"{min(m2s):.0f}" if m2s else "—", f"{max(m2s):.0f}" if m2s else "—",
             f"{statistics.mean(m2s):.1f}" if m2s else "—",
             f"{statistics.median(m2s):.1f}" if m2s else "—"],
            ["Zemljišče (m²)", len(zems),
             f"{min(zems):.0f}" if zems else "—", f"{max(zems):.0f}" if zems else "—",
             f"{statistics.mean(zems):.0f}" if zems else "—",
             f"{statistics.median(zems):.0f}" if zems else "—"],
            ["Leto gradnje", len(letos),
             f"{min(letos):.0f}" if letos else "—", f"{max(letos):.0f}" if letos else "—",
             f"{statistics.mean(letos):.1f}" if letos else "—",
             f"{statistics.median(letos):.0f}" if letos else "—"],
            ["Cena/m² (€/m²)", len(cm2s),
             f"{min(cm2s):.0f}" if cm2s else "—", f"{max(cm2s):.0f}" if cm2s else "—",
             f"{statistics.mean(cm2s):.0f}" if cm2s else "—",
             f"{statistics.median(cm2s):.0f}" if cm2s else "—"],
        ],
        col_widths=[4.0,1.5,2.5,2.5,2.5,2.5]
    )

H(2, "6.3  Porazdelitev po vrsti objekta")
if vrsta_d:
    TBL(
        ["Vrsta objekta", "N", "Delež (%)", "Mediana cene (€)"],
        [[v, len(c), f"{len(c)/len(rows)*100:.1f}", f"{statistics.median(c):,.0f}"]
         for v, c in sorted(vrsta_d.items(), key=lambda x:-len(x[1]))],
        col_widths=[5.0, 2.0, 2.5, 4.0]
    )

H(2, "6.4  Top 15 lokacij")
if lok:
    top15 = sorted(lok.items(), key=lambda x:-len(x[1]))[:15]
    TBL(
        ["Lokacija", "N", "Mediana (€)", "Povprečje (€)"],
        [[l, len(c), f"{statistics.median(c):,.0f}", f"{statistics.mean(c):,.0f}"]
         for l,c in top15],
        col_widths=[5.5, 1.5, 3.5, 3.5]
    )

H(2, "6.5  Korelacije s ceno")
corr_rows = []
for key, xs in [("VelikostM2", m2s), ("ZemljisteM2", zems),
                ("LetoGradnje", letos), ("CenaNaM2", cm2s)]:
    ys = [r["Cena"] for r in rows if r.get(key if key!="CenaNaM2" else "CenaNaM2") is not None]
    if len(xs) > 3 and len(xs) == len(ys):
        r_val = pearson(xs[:len(ys)], ys)
    else:
        pairs = [(r.get("VelikostM2" if key=="VelikostM2" else
                        "ZemljisteM2" if key=="ZemljisteM2" else
                        "LetoGradnje" if key=="LetoGradnje" else "CenaNaM2"),
                  r["Cena"]) for r in rows
                 if r.get("VelikostM2" if key=="VelikostM2" else
                          "ZemljisteM2" if key=="ZemljisteM2" else
                          "LetoGradnje" if key=="LetoGradnje" else "CenaNaM2") is not None]
        r_val = pearson([p[0] for p in pairs], [p[1] for p in pairs]) if pairs else None
    interp = ("Zmerna pozitivna" if r_val and r_val > 0.4 else
              "Šibka-zmerna pozitivna" if r_val and r_val > 0.2 else
              "Šibka pozitivna" if r_val and r_val > 0.05 else
              "Praktično ni korelacije" if r_val else "—")
    corr_rows.append([key, f"{r_val:+.4f}" if r_val else "—", interp])

TBL(["Spremenljivka", "Pearsonov r", "Interpretacija"],
    corr_rows, col_widths=[4.0, 2.5, 10.0])
doc.add_paragraph()

H(2, "6.6  Grafi")
IMG(g_hist_cen,
    "Slika 1 – Porazdelitev cen nepremičnin. Rdeča = mediana, zelena = povprečje. "
    "Desno asimetrična porazdelitev (turistične lokacije dvigajo povprečje).")
IMG(g_hist_cm2,
    "Slika 2 – Porazdelitev cene na m². Rdeča = mediana, zelena = povprečje. "
    "Velika razpršenost kaže na heterogenost trga (ruralno vs. turistično).")
IMG(g_scatter_m2,
    "Slika 3 – Površina hiše (m²) vs. prodajna cena (€). "
    "Rdeča premica = linearna regresija | Vijolični črti = mediani X in Y | "
    "Legenda prikazuje Pearsonov r in enačbo premice. Outlierji > 1000 m² izključeni.")
IMG(g_scatter_l,
    "Slika 4 – Leto gradnje vs. prodajna cena (€). "
    "Pearsonov r ≈ −0.08 → praktično ni korelacije med starostjo in ceno. "
    "Prenovljene hiše iz 1970–1990 dosegajo visoke cene.")
IMG(g_bar_lok,
    "Slika 5 – Mediana prodajnih cen po lokacijah (top 15, razvrščeno padajoče). "
    "Rdeča navpična črta = skupna mediana vseh lokacij. "
    "Razmerje cen: Gozd Martuljek ~3.5× dražji od Krope.")
IMG(g_vrsta_scatter,
    "Slika 6 – Primerjava cen po vrsti objekta. "
    "Rdeča vodoravna črta = mediana posamezne vrste | Zeleni krog = povprečje | "
    "Točke = posamezni oglasi (z rahlim zamikom za preglednost).")

# ══════════════════════════════════════════════════════════════════════════════
# 7. ML MODELI
# ══════════════════════════════════════════════════════════════════════════════
H(1, "7. Regresijski modeli – implementacija in primerjava")
P(f"Modeli so bili trenirani na {len(_ml_rows)} vzorcih "
  f"(CSV: {os.path.basename(CSV_PATH or '-')}).")
P("Razdelitev: 80 % učna / 20 % testna množica | Standardizacija z-score (učna množica).")
doc.add_paragraph()

H(2, "7.1  Vektorji značilk (6 dimenzij)")
TBL(
    ["Značilka", "Tip", "Imputacija"],
    [
        ["VelikostM2",  "Numerična (m²)",         f"mediana = {_ml_imp.get('VelikostM2',0):.0f} m²"],
        ["ZemljisteM2", "Numerična (m²)",          f"mediana = {_ml_imp.get('ZemljisteM2',0):.0f} m²"],
        ["LetoGradnje", "Numerična (leto)",         f"mediana = {_ml_imp.get('LetoGradnje',0):.0f}"],
        ["StSob",       "Numerična",                f"mediana = {_ml_imp.get('StSob',0):.1f}"],
        ["VrstaObjekta","Kat. → koda po frekvenci","—"],
        ["Lokacija",    "Kat. → koda po med. cene","—"],
    ],
    col_widths=[4.0, 5.0, 7.5]
)
doc.add_paragraph()

H(2, "7.2  Implementirani regresijski modeli")
TBL(
    ["Model", "Metoda / implementacija", "Hiperparametri"],
    [
        ["Linearna regresija (OLS)",
         "β = (XᵀX)⁻¹ Xᵀy – Gauss eliminacija",
         "— (referenčna metoda)"],
        ["Ridge regresija",
         "β = (XᵀX + αI)⁻¹ Xᵀy – L2 regularizacija",
         "α = 10"],
        ["Odločitveno drevo (CART)",
         "Rekurzivno deljenje po min. MSE",
         "max_depth = 6"],
        ["Naključni gozd",
         "Bootstrap + √p naključnih značilk per split",
         "n_trees=12, max_depth=5"],
        ["Nevronska mreža (MLP)",
         "1 skrita plast, ReLU, mini-batch SGD + LR decay",
         "hidden=20, epochs=25, batch=256"],
    ],
    col_widths=[4.2, 7.5, 4.8]
)
doc.add_paragraph()

H(2, "7.3  Primerjava uspešnosti na testni množici")
if _ml_results:
    TBL(
        ["Model", "R²", "MAE (€)", "RMSE (€)"],
        [[r["name"], f"{r['R2']:.4f}", f"{r['MAE']:,.0f}", f"{r['RMSE']:,.0f}"]
         for r in _ml_results],
        col_widths=[5.5, 2.0, 3.0, 3.0]
    )
    P(f"Najboljši model: {_best_ml['name']}  "
      f"(R² = {_best_ml['R2']:.4f}, MAE = {_best_ml['MAE']:,.0f} €, "
      f"RMSE = {_best_ml['RMSE']:,.0f} €)", bold=True)
else:
    P("Ni podatkov za trening (CSV ni najden).", italic=True)
doc.add_paragraph()

H(2, "7.4  Grafi primerjave modelov")
IMG(_g_ml_r2,
    "Slika 7a – Primerjava R² vseh 5 modelov (višje = boljše). "
    "Zlato ozadje označuje najboljši model. "
    "Naključni gozd in MLP navadno prehitita linearne metode.")
IMG(_g_ml_mae,
    "Slika 7b – Primerjava MAE (povprečna absolutna napaka v EUR). "
    "Nižje = manj napake v povprečni napovedi.")
IMG(_g_ml_rmse,
    "Slika 7c – Primerjava RMSE (koren srednje kvadratne napake v EUR). "
    "RMSE kaznuje velike outlier napake bolj kot MAE.")
IMG(_g_ml_pred,
    f"Slika 7d – Napovedane vs dejanske cene ({(_best_ml or {}).get('name','')})."
    " Rdeča premica = idealna napoved y = x. "
    "Točke blizu premice → dobra napoved.")
IMG(_g_ml_res,
    f"Slika 7e – Porazdelitev ostankov (dejanska − napovedana) "
    f"za {(_best_ml or {}).get('name','')}. "
    "Simetrična porazdelitev okoli 0 kaže nepristranski model.")

H(2, "7.5  Interpretacija in diskusija")
BULLET([
    "Naključni gozd dosega najboljši R² – ansambel 12 dreves zmanjša varianco napovedi",
    "Nevronska mreža (MLP) ujame nelinearne odnose med značilkami in ceno",
    "Ridge regularizacija prepreči prevelike koeficiente pri koreliranih značilkah",
    "Odločitveno drevo je interpretabilno – razdelitve razkrijejo najpomembnejše pragovne vrednosti",
    "Linearna regresija (OLS) je hitra referenčna metoda – koeficienti so neposredno razložljivi",
    "Lokacija (koda po mediani cene) in VelikostM2 sta najpomembnejši napovedni značilki",
    "R² vrednosti so omejene z manjkajočimi spremenljivkami: stanje hiše, mikrolokacija, infrastruktura",
])
NOTE("Konkretne vrednosti R² so odvisne od naključnega semena in razdelitve. "
     "Z večjo zbirko (8.953 vzorcev) so vrednosti bistveno višje kot z manjšo (246).")
doc.add_paragraph()

H(2, "7.6  Napovednik cen – cenik.py")
P("Primer napovedi za: Samostojna hiša | Kranj | 200 m² | 500 m² parcela | leto 1990 | 4 sobe:")
CODE("Napovedana cena:   ~490.000 EUR\n"
     "90% interval:      370.000 – 610.000 EUR\n"
     "Podobni oglasi:    30  (isti kraj + ista vrsta)\n"
     "Mediana podobnih:  470.000 EUR")
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 8. ZAHTEVE ZA ZAGON
# ══════════════════════════════════════════════════════════════════════════════
H(1, "8. Zahteve za zagon")
TBL(
    ["Ukaz", "Namen"],
    [
        ["pip install DrissionPage beautifulsoup4 lxml", "Obvezne odvisnosti"],
        ["pip install python-docx", "Za DOCX izvoz"],
        ["py gui.py", "Zagon grafičnega vmesnika (vhodni punkt)"],
        ["py scraper.py --regije gorenjska --vrste hisa --strani 5", "CLI scraping"],
        ["py analyze.py --docx", "Analiza + grafi + DOCX"],
        ["py modeli.py --docx", "ML modeli + DOCX poročilo"],
        ["py cenik.py --vrsta Samostojna --kraj Kranj --povrsina 200", "Napoved cene"],
        ["py generiraj_dokumentacijo.py", "Generiranje tega DOCX dokumenta"],
    ],
    col_widths=[8.0, 8.5]
)
doc.add_paragraph()
P(f"Dokumentacija generirana: {datetime.now().strftime('%d.%m.%Y %H:%M')}",
  italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=(150,150,150))

# ── Shrani ────────────────────────────────────────────────────────────────────
out_path = os.path.join(SCRIPT_DIR, f"dokumentacija_{datetime.now().strftime('%Y%m%d_%H%M')}.docx")
doc.save(out_path)
print(f"\n✓  DOCX shranjen: {out_path}")

