"""
analyze.py – Analiza nepremičnin (Gorenjska, hiše)
====================================================
Samo standardna knjižnica Pythona – brez numpy/pandas/matplotlib.
Ustvari besedilno poročilo in SVG grafe.

Zagon:
    python analyze.py
    python analyze.py --csv pot/do/datoteke.csv
"""

import sys, os, csv, math, argparse, statistics, random, struct, zlib
from datetime import datetime
from collections import defaultdict

# ── Argumenti ────────────────────────────────────────────────────────────────
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

parser = argparse.ArgumentParser()
parser.add_argument("--csv",   default=None,  help="Pot do CSV datoteke")
parser.add_argument("--vrste", default=None,
                    help="Filter VrstaObjekta (vejica-ločeno, npr. Hiša,Stanovanje)")
parser.add_argument("--grafi", default="all",
                    help="Grafi za generiranje: all ali vejica-ločeno iz "
                         "hist_cen,hist_cm2,scatter_m2,scatter_leto,bar_lok")
parser.add_argument("--docx",  action="store_true",
                    help="Izvozi analizo v Word dokument (.docx)")
parser.add_argument("--docx-izhod", default=None, dest="docx_izhod",
                    help="Pot za shranjevanje DOCX (privzeto: analiza_DATUM.docx)")
args = parser.parse_args()

CSV_CANDIDATES = [
    args.csv,
    "nepremicnine_export_prodaja.csv",
    "nepremicnine_export_najem.csv",
    "nepremicnine_export.csv",
    os.path.join("UUI-lv2", "bin", "Debug", "net10.0", "nepremicnine_export.csv"),
]
csv_path = next((p for p in CSV_CANDIDATES if p and os.path.isfile(p)), None)
if csv_path is None:
    print("✗  CSV datoteka ni najdena.")
    sys.exit(1)

print(f"✓  Berem podatke: {csv_path}")

# ── Nalaganje ────────────────────────────────────────────────────────────────
def to_float(s):
    try:
        v = float(str(s).replace(",", ".").strip())
        return v if math.isfinite(v) else None
    except (ValueError, TypeError):
        return None

rows = []
with open(csv_path, newline="", encoding="utf-8-sig") as f:
    reader = csv.DictReader(f, delimiter=";")
    for r in reader:
        cena = to_float(r.get("Cena"))
        m2   = to_float(r.get("VelikostM2"))
        zem  = to_float(r.get("ZemljisteM2"))
        leto = to_float(r.get("LetoGradnje"))
        if cena is None or not (5_000 < cena < 5_000_000):
            continue
        if m2 is not None and m2 >= 5_000:
            m2 = None
        cm2 = round(cena / m2, 2) if (m2 and m2 > 0) else None
        if cm2 is not None and not (100 < cm2 < 25_000):
            cm2 = None
        rows.append({
            "Naslov":      r.get("Naslov", "").strip(),
            "Lokacija":    r.get("Lokacija", "").strip(),
            "VrstaObjekta":r.get("VrstaObjekta", "").strip(),
            "Cena":        cena,
            "VelikostM2":  m2,
            "ZemljisteM2": zem,
            "LetoGradnje": leto,
            "CenaNaM2":    cm2,
        })

print(f"   Skupaj vrstic: {len(rows)}\n")

# ── Filter po vrsti objekta ───────────────────────────────────────────────────
if args.vrste:
    filter_set = {v.strip().lower() for v in args.vrste.split(",")}
    before = len(rows)
    rows = [r for r in rows if r["VrstaObjekta"].lower() in filter_set]
    vrste_str = ", ".join(sorted(filter_set))
    print(f"   Filter vrst ({vrste_str}): {before} → {len(rows)} vrstic\n")
    if not rows:
        print("✗  Ni podatkov po filtru vrst.")
        sys.exit(0)

# ── Izbira grafov ─────────────────────────────────────────────────────────────
_VSI_GRAFI = {"hist_cen", "hist_cm2", "scatter_m2", "scatter_leto", "bar_lok"}
if args.grafi.strip().lower() == "all":
    sel_grafi = _VSI_GRAFI
else:
    sel_grafi = {g.strip() for g in args.grafi.split(",")} & _VSI_GRAFI

# ── Pomožne funkcije ──────────────────────────────────────────────────────────
def col(data, key):
    return [r[key] for r in data if r[key] is not None]

def fmt(v, suf=""):
    return f"{v:,.0f}{suf}" if v is not None else "n/a"

def stats_block(vals, label, suf=""):
    if not vals:
        return
    n   = len(vals)
    avg = statistics.mean(vals)
    med = statistics.median(vals)
    sd  = statistics.stdev(vals) if n > 1 else 0.0
    print(f"\n  {label}")
    print(f"    N:          {n}")
    print(f"    Min:        {fmt(min(vals), suf)}")
    print(f"    Max:        {fmt(max(vals), suf)}")
    print(f"    Povprečje:  {fmt(avg, suf)}")
    print(f"    Mediana:    {fmt(med, suf)}")
    print(f"    Std. dev.:  {fmt(sd, suf)}")

def pearson(xs, ys):
    n = len(xs)
    if n < 3:
        return None
    mx, my = statistics.mean(xs), statistics.mean(ys)
    num = sum((x - mx)*(y - my) for x, y in zip(xs, ys))
    den = math.sqrt(sum((x-mx)**2 for x in xs) * sum((y-my)**2 for y in ys))
    return num / den if den else None

# ── Statistike ────────────────────────────────────────────────────────────────
print("=" * 58)
print("  OSNOVNA STATISTIKA")
print("=" * 58)
stats_block(col(rows, "Cena"),       "Cena (€)",        " €")
stats_block(col(rows, "VelikostM2"), "Površina (m²)",   " m²")
stats_block(col(rows, "CenaNaM2"),   "Cena/m² (€/m²)",  " €/m²")
stats_block([r["LetoGradnje"] for r in rows if r["LetoGradnje"] and r["LetoGradnje"] > 1900],
            "Leto gradnje", "")

# Top lokacije
lok_groups: dict[str, list] = defaultdict(list)
for r in rows:
    if r["Lokacija"]:
        lok_groups[r["Lokacija"]].append(r["Cena"])

print("\n  TOP LOKACIJE (po številu oglasov):")
print(f"  {'Lokacija':<28}  {'N':>4}  {'Povp. cena':>12}  {'Mediana':>12}")
print("  " + "-" * 62)
for lok, cene in sorted(lok_groups.items(), key=lambda x: -len(x[1]))[:15]:
    print(f"  {lok:<28}  {len(cene):>4}  "
          f"{statistics.mean(cene):>12,.0f} €  "
          f"{statistics.median(cene):>12,.0f} €")

# Vrsta objekta
print("\n  PO VRSTI OBJEKTA:")
vrsta_groups: dict[str, list] = defaultdict(list)
for r in rows:
    vrsta_groups[r["VrstaObjekta"] or "?"].append(r["Cena"])
for v, cene in sorted(vrsta_groups.items(), key=lambda x: -len(x[1])):
    print(f"  {v:<22} N={len(cene):>3}  povp={statistics.mean(cene):>10,.0f} €")

print("\n" + "=" * 58)

# ── Korelacije ────────────────────────────────────────────────────────────────
print("\n  KORELACIJE S CENO:")
for key, label in [("VelikostM2","Površina"),("ZemljisteM2","Zemljišče"),
                   ("LetoGradnje","Leto gradnje"),("CenaNaM2","Cena/m²")]:
    pairs = [(r[key], r["Cena"]) for r in rows if r[key] is not None]
    if len(pairs) > 3:
        r_val = pearson([p[0] for p in pairs], [p[1] for p in pairs])
        if r_val is not None:
            print(f"  {label:<18}  r = {r_val:+.4f}")

# ── Ridge regresija (standardna knjižnica) ────────────────────────────────────
print("\n" + "=" * 58)
print("  LINEARNI REGRESIJSKI MODEL")
print("  (cena ~ površina + zemljišče + leto gradnje)")
print("=" * 58)

FEATS = ["VelikostM2", "ZemljisteM2", "LetoGradnje"]
mrows = [r for r in rows if all(r[f] for f in FEATS) and r["LetoGradnje"] > 1900]

if len(mrows) >= 20:
    def get_XY(data):
        return [[r[f] for f in FEATS] for r in data], [r["Cena"] for r in data]

    def standardize(X):
        p = len(X[0])
        means = [statistics.mean(row[j] for row in X) for j in range(p)]
        stds  = [(statistics.stdev(row[j] for row in X) or 1.0) for j in range(p)]
        return [[(row[j]-means[j])/stds[j] for j in range(p)] for row in X], means, stds

    def ridge(X, y, alpha=1.0):
        n, p = len(X), len(X[0])
        XtX = [[sum(X[i][r]*X[i][c] for i in range(n)) for c in range(p)] for r in range(p)]
        for j in range(p): XtX[j][j] += alpha
        Xty = [sum(X[i][j]*y[i] for i in range(n)) for j in range(p)]
        A = [row[:] + [Xty[j]] for j, row in enumerate(XtX)]
        for c in range(p):
            piv = max(range(c, p), key=lambda r: abs(A[r][c]))
            A[c], A[piv] = A[piv], A[c]
            if abs(A[c][c]) < 1e-12: continue
            for r in range(p):
                if r != c:
                    f = A[r][c] / A[c][c]
                    A[r] = [A[r][k] - f*A[c][k] for k in range(p+1)]
        coefs = [A[j][p] / A[j][j] for j in range(p)]
        ic = statistics.mean(y) - sum(c*statistics.mean(X[i][j] for i in range(n))
                                      for j,c in enumerate(coefs))
        return coefs, ic

    def predict(X, coefs, ic):
        return [sum(x[j]*coefs[j] for j in range(len(coefs))) + ic for x in X]

    def r2_score(yt, yp):
        my = statistics.mean(yt)
        ss_t = sum((v-my)**2 for v in yt)
        ss_r = sum((t-p)**2 for t,p in zip(yt,yp))
        return 1 - ss_r/ss_t if ss_t else 0.0

    random.seed(42)
    shuffled = mrows[:]
    random.shuffle(shuffled)
    sp = int(0.8*len(shuffled))
    tr, te = shuffled[:sp], shuffled[sp:]

    Xtr, ytr = get_XY(tr); Xte, yte = get_XY(te)
    Xtr_s, fm, fs = standardize(Xtr)
    Xte_s = [[(x[j]-fm[j])/fs[j] for j in range(len(FEATS))] for x in Xte]
    coefs, ic = ridge(Xtr_s, ytr)
    ypred = predict(Xte_s, coefs, ic)

    print(f"\n  Učna množica:  {len(tr)}  |  Testna: {len(te)}")
    print(f"  R² (test):     {r2_score(yte, ypred):.4f}")
    print(f"  MAE:           {statistics.mean(abs(t-p) for t,p in zip(yte,ypred)):,.0f} €")
    print(f"  RMSE:          {math.sqrt(statistics.mean((t-p)**2 for t,p in zip(yte,ypred))):,.0f} €")
    print("\n  Koeficienti (standardizirani):")
    for f, c in zip(FEATS, coefs):
        print(f"    {f:<16}: {c:>12,.2f}")

    # 5-fold CV
    fold = len(mrows)//5
    cv_r2s = []
    for k in range(5):
        val   = mrows[k*fold:(k+1)*fold]
        train = mrows[:k*fold] + mrows[(k+1)*fold:]
        if len(val)<2 or len(train)<10: continue
        Xt,yt_ = get_XY(train); Xv,yv_ = get_XY(val)
        Xts,fm2,fs2 = standardize(Xt)
        Xvs = [[(x[j]-fm2[j])/fs2[j] for j in range(len(FEATS))] for x in Xv]
        c2,ic2 = ridge(Xts,yt_)
        cv_r2s.append(r2_score(yv_, predict(Xvs,c2,ic2)))
    if cv_r2s:
        print(f"  R² (CV 5-fold): {statistics.mean(cv_r2s):.4f}")
else:
    print(f"\n  Premalo podatkov ({len(mrows)} vrstic, potrebnih vsaj 20).")

# ── ASCII histogrami ──────────────────────────────────────────────────────────
def ascii_hist(vals, title, bins=15, W=38, fmt_fn=lambda v: f"{v:,.0f}"):
    if not vals: return
    mn, mx = min(vals), max(vals)
    if mn == mx: return
    step = (mx-mn)/bins
    counts = [0]*bins
    for v in vals:
        b = min(int((v-mn)/step), bins-1)
        counts[b] += 1
    max_c = max(counts) or 1
    print(f"\n  {title}")
    print(f"  {'─'*(W+27)}")
    for i, c in enumerate(counts):
        lo = mn + i*step
        bar = "█" * int(c/max_c*W)
        print(f"  {fmt_fn(lo):>12} │{bar:<{W}}│ {c}")
    print(f"  {'─'*(W+27)}")

print("\n" + "=" * 58)
print("  PORAZDELITEV CEN")
print("=" * 58)
ascii_hist(col(rows,"Cena"), "Cena hiše (€)", fmt_fn=lambda v: f"{v/1000:.0f}k€")

print("\n" + "=" * 58)
print("  PORAZDELITEV CENA/m²")
print("=" * 58)
ascii_hist(col(rows,"CenaNaM2"), "Cena/m² (€/m²)")

# ── SVG grafi ─────────────────────────────────────────────────────────────────

def _svg_header(W, H, title):
    return [f'<svg xmlns="http://www.w3.org/2000/svg" width="{W}" height="{H}">',
            f'<rect width="{W}" height="{H}" fill="#fafafa" rx="6"/>',
            f'<text x="{W//2}" y="24" text-anchor="middle" font-size="14" '
            f'font-weight="bold" font-family="Arial,sans-serif">{title}</text>']

def svg_scatter(xs, ys, title, xlabel, ylabel, path, color="#2196F3"):
    if not xs: return
    W, H, M = 620, 400, 56
    xmin,xmax = min(xs),max(xs); ymin,ymax = min(ys),max(ys)
    dx,dy = xmax-xmin or 1, ymax-ymin or 1
    def px(x): return M + (x-xmin)/dx*(W-M-20)
    def py(y): return H-M - (y-ymin)/dy*(H-M-40)
    xm,ym = statistics.mean(xs),statistics.mean(ys)
    slope = sum((x-xm)*(y-ym) for x,y in zip(xs,ys)) / (sum((x-xm)**2 for x in xs) or 1)
    inter = ym - slope*xm
    r_val = pearson(xs,ys)
    r_str = f"  r = {r_val:+.3f}" if r_val else ""
    lines = _svg_header(W, H, f"{title}{r_str}")
    lines += [f'<line x1="{M}" y1="{H-M}" x2="{W-20}" y2="{H-M}" stroke="#ccc"/>',
              f'<line x1="{M}" y1="36" x2="{M}" y2="{H-M}" stroke="#ccc"/>',
              f'<text x="{(W+M)//2}" y="{H-8}" text-anchor="middle" font-size="11" font-family="Arial">{xlabel}</text>',
              f'<text x="14" y="{H//2}" text-anchor="middle" font-size="11" font-family="Arial" '
              f'transform="rotate(-90,14,{H//2})">{ylabel}</text>']
    for x,y in zip(xs,ys):
        lines.append(f'<circle cx="{px(x):.1f}" cy="{py(y):.1f}" r="3.5" '
                     f'fill="{color}" opacity="0.55"/>')
    rx1,rx2 = min(xs),max(xs)
    lines.append(f'<line x1="{px(rx1):.1f}" y1="{py(slope*rx1+inter):.1f}" '
                 f'x2="{px(rx2):.1f}" y2="{py(slope*rx2+inter):.1f}" '
                 f'stroke="crimson" stroke-width="2" stroke-dasharray="6,3"/>')
    lines.append("</svg>")
    with open(path, "w", encoding="utf-8") as f: f.write("\n".join(lines))

def svg_hist_chart(vals, title, xlabel, path, bins=20, color="#2196F3"):
    if not vals: return
    mn,mx = min(vals),max(vals); step = (mx-mn)/bins or 1
    counts = [0]*bins
    for v in vals:
        b = min(int((v-mn)/step), bins-1)
        counts[b] += 1
    W, H, ML, M = 680, 400, 60, 50
    bw = (W-ML-20)/bins
    max_c = max(counts) or 1
    def bx(i): return ML + i*bw
    def by(c): return H-M - c/max_c*(H-M-44)
    lines = _svg_header(W, H, title)
    lines += [f'<line x1="{ML}" y1="{H-M}" x2="{W-20}" y2="{H-M}" stroke="#ccc"/>',
              f'<line x1="{ML}" y1="36" x2="{ML}" y2="{H-M}" stroke="#ccc"/>',
              f'<text x="{(W+ML)//2}" y="{H-8}" text-anchor="middle" font-size="11" font-family="Arial">{xlabel}</text>']
    med = statistics.median(vals); avg = statistics.mean(vals)
    def xv(v): return ML + (v-mn)/(mx-mn)*(W-ML-20)
    lines += [f'<line x1="{xv(med):.1f}" y1="36" x2="{xv(med):.1f}" y2="{H-M}" '
              f'stroke="red" stroke-width="1.5" stroke-dasharray="4,3"/>',
              f'<text x="{xv(med):.1f}" y="50" text-anchor="middle" font-size="9" '
              f'fill="red" font-family="Arial">Med {med/1000:.0f}k</text>',
              f'<line x1="{xv(avg):.1f}" y1="36" x2="{xv(avg):.1f}" y2="{H-M}" '
              f'stroke="orange" stroke-width="1.5" stroke-dasharray="4,3"/>',
              f'<text x="{xv(avg):.1f}" y="60" text-anchor="middle" font-size="9" '
              f'fill="darkorange" font-family="Arial">Avg {avg/1000:.0f}k</text>']
    for i,c in enumerate(counts):
        y = by(c); h = H-M-y
        lines.append(f'<rect x="{bx(i):.1f}" y="{y:.1f}" width="{bw*0.9:.1f}" '
                     f'height="{h:.1f}" fill="{color}" opacity="0.82" rx="2"/>')
    lines.append("</svg>")
    with open(path, "w", encoding="utf-8") as f: f.write("\n".join(lines))

def svg_bar_chart(labels, values, title, xlabel, ylabel, path, color="#4CAF50"):
    if not labels: return
    W, H, ML, M = 700, 440, 170, 50
    vmax = max(values) or 1
    bw = (W-ML-20)/len(labels)
    def bx(i): return ML + i*bw
    def by(v): return H-M - v/vmax*(H-M-44)
    lines = _svg_header(W, H, title)
    lines += [f'<line x1="{ML}" y1="{H-M}" x2="{W-20}" y2="{H-M}" stroke="#ccc"/>',
              f'<line x1="{ML}" y1="36" x2="{ML}" y2="{H-M}" stroke="#ccc"/>',
              f'<text x="{(W+ML)//2}" y="{H-6}" text-anchor="middle" font-size="11" font-family="Arial">{xlabel}</text>',
              f'<text x="14" y="{H//2}" text-anchor="middle" font-size="10" font-family="Arial" '
              f'transform="rotate(-90,14,{H//2})">{ylabel}</text>']
    for i,(lbl,v) in enumerate(zip(labels,values)):
        x=bx(i); y=by(v); h=H-M-y
        lines += [f'<rect x="{x:.1f}" y="{y:.1f}" width="{bw*0.82:.1f}" height="{h:.1f}" '
                  f'fill="{color}" opacity="0.85" rx="2"/>',
                  f'<text x="{x+bw*0.41:.1f}" y="{H-M+12}" text-anchor="end" font-size="9" '
                  f'font-family="Arial" transform="rotate(-38,{x+bw*0.41:.1f},{H-M+12})">{lbl[:20]}</text>',
                  f'<text x="{x+bw*0.41:.1f}" y="{y-4:.1f}" text-anchor="middle" font-size="9" '
                  f'font-family="Arial">{v/1000:.0f}k</text>']
    lines.append("</svg>")
    with open(path, "w", encoding="utf-8") as f: f.write("\n".join(lines))

# ── DOCX IZVOZ  (struct + zlib + python-docx) ────────────────────────────────

_F57: dict = {       # 5×7 bitmap font: 7 ints, each = 5-bit row (bit4=leftmost)
    ' ':[0,0,0,0,0,0,0],
    '0':[14,17,19,21,25,17,14],'1':[4,12,4,4,4,4,14],'2':[14,17,1,2,4,8,31],
    '3':[31,2,4,2,1,17,14],'4':[2,6,10,18,31,2,2],'5':[31,16,30,1,1,17,14],
    '6':[6,8,16,30,17,17,14],'7':[31,1,2,4,8,8,8],'8':[14,17,17,14,17,17,14],
    '9':[14,17,17,15,1,2,12],'.':[0,0,0,0,0,12,12],'-':[0,0,0,31,0,0,0],
    '+':[0,4,4,31,4,4,0],'=':[0,0,31,0,31,0,0],'/':[1,2,4,8,16,0,0],
    '(':[2,4,8,8,8,4,2],')':[8,4,2,2,2,4,8],'%':[24,25,2,4,8,19,3],
    'k':[16,16,18,20,24,20,18],'m':[0,26,21,21,21,21,21],
    'a':[0,14,1,15,17,19,13],'b':[16,16,30,17,17,17,30],'c':[0,14,17,16,16,17,14],
    'd':[1,1,13,19,17,19,13],'e':[0,14,17,31,16,17,14],'f':[6,9,8,30,8,8,8],
    'g':[0,14,17,17,15,1,14],'h':[16,16,30,17,17,17,17],'i':[12,12,0,12,12,12,14],
    'j':[6,0,6,6,6,22,12],'l':[12,4,4,4,4,4,14],'n':[0,28,18,17,17,17,17],
    'o':[0,14,17,17,17,17,14],'p':[0,30,17,17,30,16,16],'r':[0,22,25,16,16,16,16],
    's':[0,14,16,14,1,17,14],'t':[4,4,31,4,4,5,2],'u':[0,17,17,17,17,19,13],
    'v':[0,17,17,17,17,10,4],'z':[0,31,2,4,8,16,31],
    'A':[14,17,17,31,17,17,17],'B':[30,17,17,30,17,17,30],'C':[14,17,16,16,16,17,14],
    'D':[28,18,17,17,17,18,28],'E':[31,16,16,30,16,16,31],'F':[31,16,16,30,16,16,16],
    'G':[14,17,16,23,17,17,14],'H':[17,17,17,31,17,17,17],'I':[14,4,4,4,4,4,14],
    'J':[7,2,2,2,2,18,12],'K':[17,18,20,24,20,18,17],'L':[16,16,16,16,16,16,31],
    'M':[17,27,21,21,17,17,17],'N':[17,25,21,21,19,17,17],'O':[14,17,17,17,17,17,14],
    'P':[30,17,17,30,16,16,16],'R':[30,17,17,30,20,18,17],'S':[14,17,16,14,1,17,14],
    'T':[31,4,4,4,4,4,4],'U':[17,17,17,17,17,17,14],'V':[17,17,17,17,17,10,4],
    'W':[17,17,17,21,21,27,17],'X':[17,10,4,4,4,10,17],'Y':[17,17,10,4,4,4,4],
    'Z':[31,1,2,4,8,16,31],
}

def _asc(s: str) -> str:
    """Transliteracija za znakovje brez diakritike (za izris v grafu)."""
    return (s.replace('š','s').replace('Š','S').replace('č','c').replace('Č','C')
             .replace('ž','z').replace('Ž','Z').replace('ć','c').replace('đ','d')
             .replace('á','a').replace('é','e').replace('í','i').replace('ó','o')
             .replace('ú','u').replace('²','2'))

def _make_png(W: int, H: int, px) -> bytes:
    raw = bytearray()
    for row in px:
        raw.append(0)
        for p in row:
            raw += bytes(p)
    comp = zlib.compress(bytes(raw), 6)
    def ck(n, d):
        crc = zlib.crc32(n + d) & 0xFFFFFFFF
        return struct.pack('>I', len(d)) + n + d + struct.pack('>I', crc)
    return (b'\x89PNG\r\n\x1a\n'
            + ck(b'IHDR', struct.pack('>IIBBBBB', W, H, 8, 2, 0, 0, 0))
            + ck(b'IDAT', comp)
            + ck(b'IEND', b''))

class _CV:
    """Miniaturni pixel canvas za risanje grafov → PNG."""
    def __init__(self, W=700, H=400, bg=(248, 249, 250)):
        self.W, self.H = W, H
        self._p = [[[*bg]] * W for _ in range(H)]

    def _s(self, x, y, c):
        if 0 <= x < self.W and 0 <= y < self.H:
            self._p[y][x] = list(c)

    def rect(self, x, y, w, h, c):
        lc = list(c)
        for dy in range(max(0, y), min(self.H, y + h)):
            row = self._p[dy]
            for dx in range(max(0, x), min(self.W, x + w)):
                row[dx] = lc

    def hl(self, x1, x2, y, c, t=1):
        lc = list(c)
        for x in range(min(x1, x2), max(x1, x2) + 1):
            for dt in range(t):
                if 0 <= y + dt < self.H and 0 <= x < self.W:
                    self._p[y + dt][x] = lc

    def vl(self, x, y1, y2, c, t=1):
        lc = list(c)
        for y in range(min(y1, y2), max(y1, y2) + 1):
            for dt in range(t):
                if 0 <= y < self.H and 0 <= x + dt < self.W:
                    self._p[y][x + dt] = lc

    def line(self, x1, y1, x2, y2, c, t=1):
        dx, dy = abs(x2 - x1), abs(y2 - y1)
        sx = 1 if x1 < x2 else -1
        sy = 1 if y1 < y2 else -1
        e, x, y = dx - dy, x1, y1
        while True:
            for dt in range(t): self._s(x + dt, y, c)
            if x == x2 and y == y2:
                break
            e2 = 2 * e
            if e2 > -dy: e -= dy; x += sx
            if e2 < dx:  e += dx; y += sy

    def dot(self, cx, cy, r, c):
        for dy in range(-r, r + 1):
            for dx in range(-r, r + 1):
                if dx * dx + dy * dy <= r * r:
                    self._s(cx + dx, cy + dy, c)

    def txt(self, x, y, s, c, sc=1):
        lc = list(c)
        for ch in _asc(s):
            bits = _F57.get(ch, [0] * 7)
            for ri, b in enumerate(bits):
                for ci in range(5):
                    if b & (1 << (4 - ci)):
                        for dy in range(sc):
                            for dx in range(sc):
                                self._s(x + ci * sc + dx, y + ri * sc + dy, lc)
            x += 6 * sc

    def tw(self, s, sc=1): return len(s) * 6 * sc

    def to_png(self) -> bytes:
        return _make_png(self.W, self.H, [
            [bytes(self._p[y][x]) for x in range(self.W)]
            for y in range(self.H)
        ])


def _fv(v, mx):
    """Formatiranje vrednosti za os."""
    if mx >= 500_000: return f"{v/1000:.0f}k"
    if mx >= 5_000:   return f"{v/1000:.1f}k"
    if mx >= 100:     return f"{v:.0f}"
    return f"{v:.1f}"


def _chart_hist(vals, title, xlabel, color=(33, 150, 243)) -> bytes | None:
    if not vals: return None
    W, H, ML, MR, MT, MB = 700, 400, 65, 20, 48, 52
    CW, CH = W - ML - MR, H - MT - MB
    bins = 22
    mn, mx = min(vals), max(vals)
    if mn == mx: mx = mn + 1
    step = (mx - mn) / bins
    counts = [0] * bins
    for v in vals:
        counts[min(int((v - mn) / step), bins - 1)] += 1
    max_c = max(counts) or 1
    cv = _CV(W, H)
    for i in range(5):
        cv.hl(ML, W - MR, MT + i * CH // 4, (215, 215, 220))
    bw = CW // bins
    for i, cnt in enumerate(counts):
        if cnt == 0: continue
        bh = max(1, int(cnt / max_c * CH))
        cv.rect(ML + i * bw + 1, H - MB - bh, bw - 2, bh, color)
    med, avg = statistics.median(vals), statistics.mean(vals)
    def xv(v): return ML + int((v - mn) / (mx - mn) * CW)
    cv.vl(xv(med), MT, H - MB, (200, 50, 50), 2)
    cv.vl(xv(avg), MT, H - MB, (240, 140, 0), 2)
    for i in range(7):
        tv = mn + i * (mx - mn) / 6
        tx = ML + int(i * CW / 6)
        lbl = _fv(tv, mx)
        cv.txt(tx - cv.tw(lbl) // 2, H - MB + 6, lbl, (70, 70, 80))
    for i in range(5):
        cnt_v = int(max_c * (4 - i) / 4)
        cv.txt(max(0, ML - len(str(cnt_v)) * 6 - 4), MT + i * CH // 4 - 3,
               str(cnt_v), (70, 70, 80))
    cv.hl(ML, W - MR, H - MB, (70, 70, 80), 2)
    cv.vl(ML, MT, H - MB, (70, 70, 80), 2)
    tw = cv.tw(title, 2); cv.txt((W - tw) // 2, 10, title, (30, 30, 40), 2)
    xl = cv.tw(xlabel); cv.txt((W - xl) // 2, H - 14, xlabel, (70, 70, 80))
    cv.hl(W - MR - 130, W - MR - 115, MT + 10, (200, 50, 50), 2)
    cv.txt(W - MR - 110, MT + 4, "mediana", (200, 50, 50))
    cv.hl(W - MR - 130, W - MR - 115, MT + 22, (240, 140, 0), 2)
    cv.txt(W - MR - 110, MT + 16, "povprecje", (240, 140, 0))
    return cv.to_png()


def _chart_scatter(xs, ys, title, xlabel, color=(33, 150, 243)) -> bytes | None:
    if not xs: return None
    W, H, ML, MR, MT, MB = 700, 400, 65, 20, 48, 52
    CW, CH = W - ML - MR, H - MT - MB
    xmn, xmx = min(xs), max(xs); ymn, ymx = min(ys), max(ys)
    if xmn == xmx: xmx = xmn + 1
    if ymn == ymx: ymx = ymn + 1
    def px_x(x): return ML + int((x - xmn) / (xmx - xmn) * CW)
    def px_y(y): return H - MB - int((y - ymn) / (ymx - ymn) * CH)
    cv = _CV(W, H)
    for i in range(5):
        cv.hl(ML, W - MR, MT + i * CH // 4, (215, 215, 220))
        cv.vl(ML + i * CW // 4, MT, H - MB, (215, 215, 220))
    mxs, mys = statistics.mean(xs), statistics.mean(ys)
    slope = sum((x - mxs) * (y - mys) for x, y in zip(xs, ys)) / (sum((x - mxs) ** 2 for x in xs) or 1)
    inter = mys - slope * mxs
    cv.line(ML, px_y(slope * xmn + inter), W - MR, px_y(slope * xmx + inter),
            (220, 50, 50), 2)
    step = max(1, len(xs) // 800)
    for x, y in zip(xs[::step], ys[::step]):
        cv.dot(px_x(x), px_y(y), 3, color)
    cv.hl(ML, W - MR, H - MB, (70, 70, 80), 2)
    cv.vl(ML, MT, H - MB, (70, 70, 80), 2)
    for i in range(7):
        tv = xmn + i * (xmx - xmn) / 6
        tx = ML + int(i * CW / 6)
        lbl = _fv(tv, xmx)
        cv.txt(tx - cv.tw(lbl) // 2, H - MB + 6, lbl, (70, 70, 80))
    for i in range(5):
        tv = ymn + (4 - i) * (ymx - ymn) / 4
        lbl = _fv(tv, ymx)
        cv.txt(max(0, ML - cv.tw(lbl) - 4), MT + i * CH // 4 - 3, lbl, (70, 70, 80))
    r_v = pearson(xs, ys)
    r_s = f" (r={r_v:+.3f})" if r_v is not None else ""
    full = title + r_s
    cv.txt((W - cv.tw(full, 2)) // 2, 10, full, (30, 30, 40), 2)
    cv.txt((W - cv.tw(xlabel)) // 2, H - 14, xlabel, (70, 70, 80))
    return cv.to_png()


def _chart_bar(labels, values, title) -> bytes | None:
    if not labels: return None
    n = len(labels)
    W, ML, MR, MT, MB = 700, 175, 80, 48, 30
    H = max(280, MT + n * 30 + MB)
    CW, CH = W - ML - MR, H - MT - MB
    vmx = max(values) or 1
    cv = _CV(W, H)
    for i in range(5):
        cv.vl(ML + i * CW // 4, MT, H - MB, (215, 215, 220))
    bh = max(8, CH // n - 4)
    for i, (lbl, v) in enumerate(zip(labels, values)):
        by = MT + i * (CH // n) + 2
        bw2 = max(1, int(v / vmx * CW))
        cv.rect(ML, by, bw2, bh, (76, 175, 80))
        v_lbl = _fv(v, vmx)
        cv.txt(ML + bw2 + 3, by + bh // 2 - 3, v_lbl, (70, 70, 80))
        short = _asc(lbl[:22])
        cv.txt(max(0, ML - cv.tw(short) - 3), by + bh // 2 - 3, short, (40, 40, 50))
    cv.hl(ML, W - MR, H - MB, (70, 70, 80), 2)
    cv.vl(ML, MT, H - MB, (70, 70, 80), 2)
    cv.txt((W - cv.tw(title, 2)) // 2, 10, title, (30, 30, 40), 2)
    return cv.to_png()


def _export_docx(rows, lok_groups, sel_grafi, output_path):
    """Izvozi analizo v Word dokument (.docx)."""
    try:
        from docx import Document
        from docx.shared import Inches, Cm, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io as _io
    except ImportError:
        print("✗  Za DOCX izvoz namesti: py -m pip install python-docx")
        return
    print("   Generiram Word dokument …")
    doc = Document()
    for sec in doc.sections:
        sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)
        sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)

    # Naslov
    h0 = doc.add_heading("Analiza nepremičnin", 0)
    h0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Generirano: {datetime.now().strftime('%d.%m.%Y %H:%M')}  "
                    f"|  Število analiziranih oglasov: {len(rows)}")
    run.italic = True
    doc.add_paragraph()

    def _hdr(tbl, cols):
        cells = tbl.rows[0].cells
        for i, t in enumerate(cols):
            cells[i].text = t
            for para in cells[i].paragraphs:
                for r in para.runs: r.bold = True

    def _row(tbl, vals):
        c = tbl.add_row().cells
        for i, v in enumerate(vals): c[i].text = str(v)

    # ── 1. Osnovna statistika ──────────────────────────────────────────────
    doc.add_heading("1. Osnovna statistika", 1)
    hdrs = ["Metrika", "N", "Min", "Max", "Povprečje", "Mediana", "Std. dev."]
    st = doc.add_table(rows=1, cols=len(hdrs)); st.style = "Table Grid"
    _hdr(st, hdrs)
    for key, label, suf in [("Cena", "Cena (€)", " €"),
                              ("VelikostM2", "Površina (m²)", " m²"),
                              ("ZemljisteM2", "Površina zemljišča (m²)", " m²"),
                              ("CenaNaM2", "Cena/m² (€/m²)", " €/m²")]:
        vals = col(rows, key)
        if not vals: continue
        n = len(vals)
        _row(st, [label, n,
                  f"{min(vals):,.0f}{suf}", f"{max(vals):,.0f}{suf}",
                  f"{statistics.mean(vals):,.0f}{suf}",
                  f"{statistics.median(vals):,.0f}{suf}",
                  f"{(statistics.stdev(vals) if n > 1 else 0):,.0f}{suf}"])
    lv = [r2["LetoGradnje"] for r2 in rows if r2["LetoGradnje"] and r2["LetoGradnje"] > 1900]
    if lv:
        n = len(lv)
        _row(st, ["Leto gradnje", n, f"{min(lv):.0f}", f"{max(lv):.0f}",
                  f"{statistics.mean(lv):.1f}", f"{statistics.median(lv):.0f}",
                  f"{(statistics.stdev(lv) if n > 1 else 0):.1f}"])
    doc.add_paragraph()

    # ── 2. Top lokacije ────────────────────────────────────────────────────
    doc.add_heading("2. Top lokacije (po številu oglasov)", 1)
    top15 = sorted(lok_groups.items(), key=lambda x: -len(x[1]))[:15]
    if top15:
        lt = doc.add_table(rows=1, cols=4); lt.style = "Table Grid"
        _hdr(lt, ["Lokacija", "N", "Povp. cena (€)", "Mediana (€)"])
        for lok, cene in top15:
            _row(lt, [lok, len(cene),
                      f"{statistics.mean(cene):,.0f}",
                      f"{statistics.median(cene):,.0f}"])
    doc.add_paragraph()

    # ── 3. Po vrsti objekta ────────────────────────────────────────────────
    doc.add_heading("3. Po vrsti objekta", 1)
    vg: dict = defaultdict(list)
    for r2 in rows: vg[r2["VrstaObjekta"] or "?"].append(r2["Cena"])
    vt = doc.add_table(rows=1, cols=3); vt.style = "Table Grid"
    _hdr(vt, ["Vrsta objekta", "N", "Povp. cena (€)"])
    for v, cene in sorted(vg.items(), key=lambda x: -len(x[1])):
        _row(vt, [v, len(cene), f"{statistics.mean(cene):,.0f}"])
    doc.add_paragraph()

    # ── 4. Korelacije ──────────────────────────────────────────────────────
    doc.add_heading("4. Korelacije s ceno", 1)
    kt = doc.add_table(rows=1, cols=2); kt.style = "Table Grid"
    _hdr(kt, ["Spremenljivka", "Pearsonov r"])
    for key, label in [("VelikostM2", "Površina"), ("ZemljisteM2", "Površina zemljišča"),
                        ("LetoGradnje", "Leto gradnje"), ("CenaNaM2", "Cena/m²")]:
        pairs = [(r2[key], r2["Cena"]) for r2 in rows if r2[key] is not None]
        if len(pairs) > 3:
            rv = pearson([p[0] for p in pairs], [p[1] for p in pairs])
            if rv is not None:
                _row(kt, [label, f"{rv:+.4f}"])
    doc.add_paragraph()

    # ── 5. Grafi ──────────────────────────────────────────────────────────
    if sel_grafi:
        doc.add_heading("5. Grafi", 1)
        fig_num = [0]

        def _add_chart(png_bytes, caption):
            if png_bytes is None: return
            fig_num[0] += 1
            doc.add_picture(_io.BytesIO(png_bytes), width=Inches(5.8))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp = doc.add_paragraph(f"Slika {fig_num[0]} – {caption}")
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cp.runs:
                run.italic = True; run.font.size = Pt(9)
            doc.add_paragraph()

        if "hist_cen" in sel_grafi:
            _add_chart(_chart_hist(col(rows, "Cena"),
                                   "Porazdelitev cen", "Cena (e)"),
                       "Porazdelitev cen nepremičnin")

        if "hist_cm2" in sel_grafi:
            _add_chart(_chart_hist(col(rows, "CenaNaM2"),
                                   "Porazdelitev cene na m2", "Cena/m2",
                                   (255, 87, 34)),
                       "Porazdelitev cene na m²")

        if "scatter_m2" in sel_grafi:
            prs = [(r2["VelikostM2"], r2["Cena"] / 1000)
                   for r2 in rows if r2["VelikostM2"] and r2["Cena"]]
            _add_chart(_chart_scatter([x for x, _ in prs], [y for _, y in prs],
                                      "Povrsina vs Cena", "Povrsina (m2)"),
                       "Korelacija med površino in ceno")

        if "scatter_leto" in sel_grafi:
            prs = [(r2["LetoGradnje"], r2["Cena"] / 1000)
                   for r2 in rows if r2["LetoGradnje"] and r2["LetoGradnje"] > 1900 and r2["Cena"]]
            _add_chart(_chart_scatter([x for x, _ in prs], [y for _, y in prs],
                                      "Leto gradnje vs Cena", "Leto gradnje",
                                      (156, 39, 176)),
                       "Korelacija med letom gradnje in ceno")

        if "bar_lok" in sel_grafi:
            lok_m = {k: statistics.median(v) for k, v in lok_groups.items() if len(v) >= 2}
            top12 = sorted(lok_m.items(), key=lambda x: x[1])[-12:]
            _add_chart(_chart_bar([l for l, _ in top12], [v for _, v in top12],
                                  "Mediana cen po lokacijah (top 12)"),
                       "Mediana cen po lokacijah")

    doc.save(output_path)
    print(f"✓  Word dokument shranjen: {output_path}")


# Generate selected charts
os.makedirs("grafi", exist_ok=True)
saved_files = []

cena_vals = col(rows, "Cena")
if "hist_cen" in sel_grafi and cena_vals:
    p = os.path.join("grafi", "hist_cen.svg")
    svg_hist_chart(cena_vals, "Porazdelitev cen", "Cena (€)", p)
    saved_files.append("hist_cen.svg")

cm2_vals = col(rows, "CenaNaM2")
if "hist_cm2" in sel_grafi and cm2_vals:
    p = os.path.join("grafi", "hist_cena_m2.svg")
    svg_hist_chart(cm2_vals, "Porazdelitev cene na m²", "€/m²", p, color="#FF5722")
    saved_files.append("hist_cena_m2.svg")

both_m2 = [(r["VelikostM2"], r["Cena"]) for r in rows if r["VelikostM2"] and r["Cena"]]
if "scatter_m2" in sel_grafi and both_m2:
    p = os.path.join("grafi", "scatter_cena_m2.svg")
    svg_scatter([x for x, _ in both_m2], [y / 1000 for _, y in both_m2],
                "Cena vs. Površina", "Površina (m²)", "Cena (1000 €)", p)
    saved_files.append("scatter_cena_m2.svg")

both_leto = [(r["LetoGradnje"], r["Cena"]) for r in rows
             if r["LetoGradnje"] and r["LetoGradnje"] > 1900 and r["Cena"]]
if "scatter_leto" in sel_grafi and both_leto:
    p = os.path.join("grafi", "scatter_leto_cena.svg")
    svg_scatter([x for x, _ in both_leto], [y / 1000 for _, y in both_leto],
                "Leto gradnje vs. Cena", "Leto gradnje", "Cena (1000 €)", p, color="#9C27B0")
    saved_files.append("scatter_leto_cena.svg")

lok_med = {k: statistics.median(v) for k, v in lok_groups.items() if len(v) >= 2}
top12 = sorted(lok_med.items(), key=lambda x: x[1])[-12:]
if "bar_lok" in sel_grafi and top12:
    p = os.path.join("grafi", "bar_lokacije.svg")
    svg_bar_chart([l for l, _ in top12], [v for _, v in top12],
                  "Mediana cene po lokacijah (top 12)", "Lokacija", "Mediana (€)", p)
    saved_files.append("bar_lokacije.svg")

if saved_files:
    print(f"\n✓  SVG grafi shranjeni v: grafi/")
    for f in saved_files:
        print(f"   • {f}")
else:
    print("\n⚠  Ni generiranih grafov (noben graf ni bil izbran ali ni podatkov).")

print(f"\n  Analiza končana – {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}")

# ── DOCX izvoz ────────────────────────────────────────────────────────────────
if args.docx:
    docx_path = args.docx_izhod or os.path.join(
        os.path.dirname(os.path.abspath(csv_path)) if csv_path else ".",
        f"analiza_{datetime.now().strftime('%Y%m%d_%H%M')}.docx")
    _export_docx(rows, lok_groups, sel_grafi, docx_path)

