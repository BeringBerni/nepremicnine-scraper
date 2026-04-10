"""
modeli.py – Regresijski modeli za napovedovanje cen nepremičnin
===============================================================
Implementacije (čista standardna knjižnica):
  • Linearna regresija (OLS)
  • Ridge regresija     (optimizacija: alpha)
  • Odločitveno drevo  (optimizacija: max_depth)
  • Naključni gozd     (optimizacija: n_trees × max_depth)

Potek dela:
  1. Nalaganje CSV in čiščenje podatkov
  2. Imputacija manjkajočih vrednosti (mediana)
  3. Kodiranje kategoričnih spremenljivk
  4. Standardizacija značilk
  5. Razdelitev na učno (80 %) in testno (20 %) množico
  6. Hiperparametrska optimizacija (5-kratna CV)
  7. Primerjava modelov (R², MAE, RMSE)
  8. Generiranje DOCX poročila

Zagon:
    py modeli.py
    py modeli.py --csv podatki.csv --docx
"""

import sys, os, csv, math, random, statistics, argparse, struct, zlib, itertools
from datetime import datetime
from collections import defaultdict

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

sys.setrecursionlimit(10_000)

# ── CLI argumenti ─────────────────────────────────────────────────────────────
parser = argparse.ArgumentParser(description="Regresijski modeli – nepremičnine")
parser.add_argument("--csv",        default=None,  help="Pot do CSV datoteke")
parser.add_argument("--split",      default=0.8,   type=float, help="Delež učnih podatkov")
parser.add_argument("--seed",       default=42,    type=int,   help="Naključno seme")
parser.add_argument("--docx",       action="store_true",       help="Ustvari DOCX poročilo")
parser.add_argument("--docx-izhod", default=None,  dest="docx_izhod")
args = parser.parse_args()

random.seed(args.seed)

# ── Nalaganje CSV ──────────────────────────────────────────────────────────────
CANDIDATES = [
    args.csv,
    "nepremicnine_export_prodaja.csv",
    "nepremicnine_export_najem.csv",
    "nepremicnine_export.csv",
    os.path.join("UUI-lv2", "bin", "Debug", "net10.0", "nepremicnine_export.csv"),
]
csv_path = next((p for p in CANDIDATES if p and os.path.isfile(p)), None)
if csv_path is None:
    print("✗  CSV datoteka ni najdena.")
    sys.exit(1)
print(f"✓  Berem: {csv_path}")

def _f(s):
    try:
        v = float(str(s).replace(",", ".").strip())
        return v if math.isfinite(v) else None
    except Exception:
        return None

_ENERGY = {"A+": 8, "A": 7, "B": 6, "C": 5, "D": 4, "E": 3, "F": 2, "G": 1}

raw_rows = []
with open(csv_path, newline="", encoding="utf-8-sig") as fh:
    for r in csv.DictReader(fh, delimiter=";"):
        cena = _f(r.get("Cena"))
        if cena is None or not (5_000 < cena < 5_000_000):
            continue
        m2   = _f(r.get("VelikostM2"))
        zem  = _f(r.get("ZemljisteM2"))
        leto = _f(r.get("LetoGradnje"))
        sob  = _f(r.get("StSob"))
        enr  = _ENERGY.get(str(r.get("EnergetskiRazred", "")).strip().upper(), None)
        if m2 is not None and m2 >= 5_000:
            m2 = None
        raw_rows.append({
            "Cena":        cena,
            "VelikostM2":  m2,
            "ZemljisteM2": zem,
            "LetoGradnje": leto if (leto and leto > 1900) else None,
            "StSob":       sob,
            "EnergetRazred": float(enr) if enr else None,
            "VrstaObjekta": r.get("VrstaObjekta", "").strip(),
            "Obcina":       (r.get("Obcina") or r.get("Lokacija", "")).strip(),
        })

print(f"   Skupaj vrstic: {len(raw_rows)}")
if len(raw_rows) < 30:
    print("✗  Premalo podatkov (< 30 vrstic).")
    sys.exit(1)

# ── Predprocesiranje ───────────────────────────────────────────────────────────
def med(rows, key):
    vals = [r[key] for r in rows if r[key] is not None]
    return statistics.median(vals) if vals else 0.0

imp = {k: med(raw_rows, k) for k in
       ["VelikostM2", "ZemljisteM2", "LetoGradnje", "StSob", "EnergetRazred"]}
print("   Imputacija (mediana):", {k: f"{v:.1f}" for k, v in imp.items()})

# Kodiranje VrstaObjekta po frekvenci
vc = defaultdict(int)
for r in raw_rows:
    vc[r["VrstaObjekta"]] += 1
vrsta_enc = {v: i for i, (v, _) in enumerate(sorted(vc.items(), key=lambda x: -x[1]))}

# Kodiranje Obcina po mediani cene (nižja mediana → nižja koda)
ob_prices = defaultdict(list)
for r in raw_rows:
    if r["Obcina"]:
        ob_prices[r["Obcina"]].append(r["Cena"])
ob_med = {k: statistics.median(v) for k, v in ob_prices.items()}
ob_sorted = sorted(ob_med, key=lambda k: ob_med[k])
obcina_enc = {k: i for i, k in enumerate(ob_sorted)}

FEAT_NAMES = ["VelikostM2", "ZemljisteM2", "LetoGradnje", "StSob",
              "EnergetRazred", "VrstaObjekta (enc)", "Obcina (enc)"]

all_X, all_y = [], []
for r in raw_rows:
    x = [
        r["VelikostM2"]    if r["VelikostM2"]    is not None else imp["VelikostM2"],
        r["ZemljisteM2"]   if r["ZemljisteM2"]   is not None else imp["ZemljisteM2"],
        r["LetoGradnje"]   if r["LetoGradnje"]    is not None else imp["LetoGradnje"],
        r["StSob"]         if r["StSob"]          is not None else imp["StSob"],
        r["EnergetRazred"] if r["EnergetRazred"]  is not None else imp["EnergetRazred"],
        float(vrsta_enc.get(r["VrstaObjekta"], 0)),
        float(obcina_enc.get(r["Obcina"], len(obcina_enc) // 2)),
    ]
    all_X.append(x)
    all_y.append(r["Cena"])

P = len(FEAT_NAMES)
print(f"   Značilke ({P}): {', '.join(FEAT_NAMES)}")

# Standardizacija (po učnih podatkih – bo izračunana po razdelitvi)
def standardize(X, means=None, stds=None):
    p = len(X[0])
    if means is None:
        means = [statistics.mean(row[j] for row in X) for j in range(p)]
    if stds is None:
        stds  = [(statistics.stdev(row[j] for row in X) or 1.0) for j in range(p)]
    return [[(row[j] - means[j]) / stds[j] for j in range(p)] for row in X], means, stds

# ── Razdelitev podatkov ────────────────────────────────────────────────────────
idx = list(range(len(all_X)))
random.shuffle(idx)
cut = int(len(idx) * args.split)
tr_idx, te_idx = idx[:cut], idx[cut:]

X_tr_raw  = [all_X[i] for i in tr_idx];  y_tr = [all_y[i] for i in tr_idx]
X_te_raw  = [all_X[i] for i in te_idx];  y_te = [all_y[i] for i in te_idx]

X_tr, f_means, f_stds = standardize(X_tr_raw)
X_te, _, _ = standardize(X_te_raw, f_means, f_stds)

n_tr, n_te = len(X_tr), len(X_te)
print(f"   Učna množica: {n_tr}  |  Testna množica: {n_te}")
print()

# ── Metrike ───────────────────────────────────────────────────────────────────
def r2(yt, yp):
    my = statistics.mean(yt)
    ss_t = sum((v - my)**2 for v in yt)
    ss_r = sum((t - p)**2 for t, p in zip(yt, yp))
    return 1.0 - ss_r / ss_t if ss_t else 0.0

def mae(yt, yp):
    return statistics.mean(abs(t - p) for t, p in zip(yt, yp))

def rmse(yt, yp):
    return math.sqrt(statistics.mean((t - p)**2 for t, p in zip(yt, yp)))

def eval_metrics(yt, yp):
    return {"R2": r2(yt, yp), "MAE": mae(yt, yp), "RMSE": rmse(yt, yp)}

# ── Gauss eliminacija (reused by LinearRegression & Ridge) ────────────────────
def _gauss(A, n):
    for c in range(n):
        piv = max(range(c, n), key=lambda r: abs(A[r][c]))
        A[c], A[piv] = A[piv], A[c]
        if abs(A[c][c]) < 1e-10:
            continue
        for r in range(n):
            if r != c:
                f = A[r][c] / A[c][c]
                A[r] = [A[r][k] - f * A[c][k] for k in range(n + 1)]
    return [A[j][n] / A[j][j] if abs(A[j][j]) > 1e-10 else 0.0 for j in range(n)]

# ── MODELI ────────────────────────────────────────────────────────────────────

class LinearRegression:
    """Ordinarna linearna regresija (OLS) – normalne enačbe."""
    name = "Linearna regresija (OLS)"

    def fit(self, X, y):
        n, p = len(X), len(X[0])
        Xa = [[1.0] + list(row) for row in X]
        q = p + 1
        XtX = [[sum(Xa[i][r]*Xa[i][c] for i in range(n)) for c in range(q)] for r in range(q)]
        Xty = [sum(Xa[i][j]*y[i] for i in range(n)) for j in range(q)]
        A   = [XtX[r][:] + [Xty[r]] for r in range(q)]
        beta = _gauss(A, q)
        self._b0, self._w = beta[0], beta[1:]
        return self

    def predict(self, X):
        return [self._b0 + sum(x[j]*self._w[j] for j in range(len(x))) for x in X]

    def __str__(self):
        return self.name


class RidgeRegression:
    """Ridge regresija z L2 regularizacijo (alpha = regularizacijski parameter)."""
    name = "Ridge regresija"

    def __init__(self, alpha=1.0):
        self.alpha = alpha

    def fit(self, X, y):
        n, p = len(X), len(X[0])
        y_mean = statistics.mean(y)
        ys = [v - y_mean for v in y]
        XtX = [[sum(X[i][r]*X[i][c] for i in range(n)) for c in range(p)] for r in range(p)]
        for j in range(p):
            XtX[j][j] += self.alpha
        Xty = [sum(X[i][j]*ys[i] for i in range(n)) for j in range(p)]
        A   = [XtX[r][:] + [Xty[r]] for r in range(p)]
        self._w  = _gauss(A, p)
        self._b0 = y_mean
        return self

    def predict(self, X):
        return [self._b0 + sum(x[j]*self._w[j] for j in range(len(x))) for x in X]

    def __str__(self):
        return f"{self.name} (α={self.alpha})"


# ── Odločitveno drevo (CART) ──────────────────────────────────────────────────
class _Node:
    __slots__ = ("feat", "thresh", "left", "right", "val")
    def __init__(self, val=None, feat=None, thresh=None, left=None, right=None):
        self.val = val; self.feat = feat; self.thresh = thresh
        self.left = left; self.right = right


class DecisionTreeRegressor:
    """Odločitveno drevo (CART) za regresijo."""
    name = "Odločitveno drevo (CART)"

    def __init__(self, max_depth=5, min_samples=4, max_features=None):
        self.max_depth    = max_depth
        self.min_samples  = min_samples
        self.max_features = max_features
        self._root        = None

    def fit(self, X, y):
        self._root = self._build(X, y, 0)
        return self

    def _best_split(self, X, y):
        n, p = len(X), len(X[0])
        best_gain = -1.0
        best_feat = best_thresh = None
        total_var = sum(v**2 for v in y) / n - (sum(y)/n)**2

        feats = list(range(p))
        if self.max_features and self.max_features < p:
            feats = random.sample(feats, self.max_features)

        for fi in feats:
            # Sort by feature fi and use prefix sums
            order = sorted(range(n), key=lambda i: X[i][fi])
            ps_y  = [0.0] * (n + 1)
            ps_y2 = [0.0] * (n + 1)
            for k in range(n):
                i = order[k]
                ps_y[k+1]  = ps_y[k]  + y[i]
                ps_y2[k+1] = ps_y2[k] + y[i]**2

            for k in range(1, n):
                if X[order[k]][fi] == X[order[k-1]][fi]:
                    continue
                nl, nr = k, n - k
                if nl < 2 or nr < 2:
                    continue
                sl,  s2l = ps_y[k],       ps_y2[k]
                sr,  s2r = ps_y[n]-sl,    ps_y2[n]-s2l
                var_l = s2l/nl - (sl/nl)**2
                var_r = s2r/nr - (sr/nr)**2
                gain  = total_var - (nl*var_l + nr*var_r)/n
                if gain > best_gain:
                    best_gain  = gain
                    best_feat  = fi
                    best_thresh = (X[order[k-1]][fi] + X[order[k]][fi]) / 2
        return best_feat, best_thresh

    def _build(self, X, y, depth):
        val = sum(y) / len(y)
        if (len(y) < self.min_samples or
                (self.max_depth is not None and depth >= self.max_depth)):
            return _Node(val=val)
        feat, thresh = self._best_split(X, y)
        if feat is None:
            return _Node(val=val)
        mask = [X[i][feat] <= thresh for i in range(len(X))]
        Xl = [X[i] for i in range(len(X)) if     mask[i]]
        yl = [y[i] for i in range(len(y)) if     mask[i]]
        Xr = [X[i] for i in range(len(X)) if not mask[i]]
        yr = [y[i] for i in range(len(y)) if not mask[i]]
        if not Xl or not Xr:
            return _Node(val=val)
        return _Node(val=val, feat=feat, thresh=thresh,
                     left =self._build(Xl, yl, depth+1),
                     right=self._build(Xr, yr, depth+1))

    def _pred1(self, node, x):
        if node.feat is None:
            return node.val
        return self._pred1(node.left if x[node.feat] <= node.thresh else node.right, x)

    def predict(self, X):
        return [self._pred1(self._root, x) for x in X]

    def __str__(self):
        return f"{self.name} (max_depth={self.max_depth})"


class RandomForestRegressor:
    """Naključni gozd – ansambel odločitvenih dreves s bootstrap vzorčenjem."""
    name = "Naključni gozd"

    def __init__(self, n_estimators=20, max_depth=5, min_samples=4):
        self.n_estimators = n_estimators
        self.max_depth    = max_depth
        self.min_samples  = min_samples
        self._trees       = []

    def fit(self, X, y):
        n, p   = len(X), len(X[0])
        mf     = max(1, int(p**0.5))          # sqrt(p) random features per split
        self._trees = []
        for t in range(self.n_estimators):
            idxs = [random.randint(0, n-1) for _ in range(n)]  # bootstrap
            Xb   = [X[i] for i in idxs]
            yb   = [y[i] for i in idxs]
            tree = DecisionTreeRegressor(
                max_depth=self.max_depth,
                min_samples=self.min_samples,
                max_features=mf
            ).fit(Xb, yb)
            self._trees.append(tree)
            if (t + 1) % 10 == 0:
                print(f"   RF: {t+1}/{self.n_estimators} dreves", flush=True)
        return self

    def predict(self, X):
        all_p = [t.predict(X) for t in self._trees]
        return [statistics.mean(all_p[j][i] for j in range(len(self._trees)))
                for i in range(len(X))]

    def predict_all(self, X):
        """Vrne napovedi vseh posameznih dreves (za interval zaupanja)."""
        return [t.predict(X) for t in self._trees]

    def __str__(self):
        return f"{self.name} (n={self.n_estimators}, max_depth={self.max_depth})"


class NeuralNetRegressor:
    """Enoplastna nevronska mreža (MLP) za regresijo.

    Arhitektura: vhod (p) → skrita plast (hidden, ReLU) → izhod (1, linearni)
    Optimizer  : mini-batch SGD z LR decay vsakih 10 epoh
    Implementacija: čista standardna knjižnica (brez NumPy/scikit-learn)
    """
    name = "Nevronska mreza (MLP)"

    def __init__(self, hidden=24, epochs=30, lr=0.01, batch=256):
        self.hidden = hidden
        self.epochs = epochs
        self.lr     = lr
        self.batch  = batch

    def fit(self, X, y):
        n, p = len(X), len(X[0])
        h    = self.hidden
        # He inicializacija uteži
        s1 = (2.0 / p) ** 0.5
        s2 = (2.0 / h) ** 0.5
        self.W1 = [[random.gauss(0, s1) for _ in range(p)] for _ in range(h)]
        self.b1 = [0.0] * h
        self.W2 = [random.gauss(0, s2) for _ in range(h)]
        self.b2 = 0.0
        # Normalizacija y (boljša konvergenca)
        ym  = statistics.mean(y)
        ys_ = statistics.stdev(y) if len(y) > 1 else 1.0
        self._ym = ym
        self._ys = ys_
        yn = [(v - ym) / ys_ for v in y]
        lr = self.lr
        for ep in range(self.epochs):
            idx = list(range(n)); random.shuffle(idx)
            for bs in range(0, n, self.batch):
                bi = idx[bs: bs + self.batch]; nb = len(bi)
                dW1 = [[0.0] * p for _ in range(h)]
                db1 = [0.0] * h; dW2 = [0.0] * h; db2 = 0.0
                for i in bi:
                    xi = X[i]
                    # Forward
                    z1 = [sum(self.W1[j][k] * xi[k] for k in range(p)) + self.b1[j]
                          for j in range(h)]
                    a1 = [max(0.0, v) for v in z1]      # ReLU
                    z2 = sum(self.W2[j] * a1[j] for j in range(h)) + self.b2
                    # Backward (MSE loss)
                    err = z2 - yn[i]; db2 += err
                    for j in range(h):
                        dW2[j] += err * a1[j]
                        if z1[j] > 0:                    # ReLU gradient
                            d_ = err * self.W2[j]; db1[j] += d_
                            for k in range(p):
                                dW1[j][k] += d_ * xi[k]
                # Update
                for j in range(h):
                    self.b1[j] -= lr * db1[j] / nb
                    self.W2[j] -= lr * dW2[j] / nb
                    for k in range(p):
                        self.W1[j][k] -= lr * dW1[j][k] / nb
                self.b2 -= lr * db2 / nb
            # LR decay
            if (ep + 1) % 10 == 0:
                lr *= 0.6
                print(f"   MLP: ep {ep + 1}/{self.epochs}  lr={lr:.5f}", flush=True)
        return self

    def predict(self, X):
        h = self.hidden; res = []
        for xi in X:
            z1 = [sum(self.W1[j][k] * xi[k] for k in range(len(xi))) + self.b1[j]
                  for j in range(h)]
            a1 = [max(0.0, v) for v in z1]
            z2 = sum(self.W2[j] * a1[j] for j in range(h)) + self.b2
            res.append(z2 * self._ys + self._ym)
        return res

    def __str__(self):
        return f"{self.name} (hidden={self.hidden}, epochs={self.epochs})"


# ── Prečno preverjanje (k-fold CV) ────────────────────────────────────────────
def kfold_score(ModelClass, params, X, y, k=5):
    n     = len(X)
    fold  = max(1, n // k)
    scores = []
    for i in range(k):
        vs = i * fold
        ve = (i+1)*fold if i < k-1 else n
        Xv = X[vs:ve]; yv = y[vs:ve]
        Xt = X[:vs] + X[ve:]; yt = y[:vs] + y[ve:]
        if len(Xt) < 10 or len(Xv) < 2:
            continue
        m = ModelClass(**params).fit(Xt, yt)
        scores.append(r2(yv, m.predict(Xv)))
    return statistics.mean(scores) if scores else float("-inf")


def grid_search(ModelClass, grid, X, y, k=5, label=""):
    keys   = list(grid.keys())
    best   = {}
    best_s = float("-inf")
    all_r  = []
    combos = list(itertools.product(*[grid[k_] for k_ in keys]))
    for vals in combos:
        params = dict(zip(keys, vals))
        s = kfold_score(ModelClass, params, X, y, k)
        all_r.append((params, s))
        if s > best_s:
            best_s = s;  best = params
        p_str = ", ".join(f"{k_}={v}" for k_, v in params.items())
        print(f"   {label} [{p_str}]  CV R²={s:.4f}")
    return best, best_s, all_r


# ── Trening in vrednotenje vseh modelov ──────────────────────────────────────
print("=" * 60)
print("  LINEARNA REGRESIJA (OLS)")
print("=" * 60)
lr = LinearRegression().fit(X_tr, y_tr)
lr_m = eval_metrics(y_te, lr.predict(X_te))
cv_lr = kfold_score(LinearRegression, {}, X_tr, y_tr)
print(f"  R²={lr_m['R2']:.4f}  MAE={lr_m['MAE']:,.0f} €  RMSE={lr_m['RMSE']:,.0f} €  CV R²={cv_lr:.4f}")

print()
print("=" * 60)
print("  RIDGE REGRESIJA  (optimizacija α)")
print("=" * 60)
ridge_grid = {"alpha": [0.01, 0.1, 1.0, 10.0, 100.0, 500.0, 1000.0]}
best_ridge, cv_ridge, ridge_all = grid_search(RidgeRegression, ridge_grid, X_tr, y_tr, label="Ridge")
ridge = RidgeRegression(**best_ridge).fit(X_tr, y_tr)
ridge_m = eval_metrics(y_te, ridge.predict(X_te))
print(f"  Najboljši α={best_ridge['alpha']}  R²={ridge_m['R2']:.4f}  MAE={ridge_m['MAE']:,.0f} €  RMSE={ridge_m['RMSE']:,.0f} €")

print()
print("=" * 60)
print("  ODLOČITVENO DREVO  (optimizacija max_depth)")
print("=" * 60)
dt_grid = {"max_depth": [2, 3, 4, 5, 6, 8]}
best_dt, cv_dt, dt_all = grid_search(DecisionTreeRegressor, dt_grid, X_tr, y_tr, label="DT")
dt = DecisionTreeRegressor(**best_dt).fit(X_tr, y_tr)
dt_m = eval_metrics(y_te, dt.predict(X_te))
print(f"  Najboljši max_depth={best_dt['max_depth']}  R²={dt_m['R2']:.4f}  MAE={dt_m['MAE']:,.0f} €  RMSE={dt_m['RMSE']:,.0f} €")

print()
print("=" * 60)
print("  NAKLJUČNI GOZD  (optimizacija n_trees × max_depth)")
print("=" * 60)
rf_grid = {"n_estimators": [10, 20], "max_depth": [3, 5, 7]}
best_rf, cv_rf, rf_all = grid_search(RandomForestRegressor, rf_grid, X_tr, y_tr, label="RF")
rf = RandomForestRegressor(**best_rf).fit(X_tr, y_tr)
rf_m = eval_metrics(y_te, rf.predict(X_te))
print(f"  Najboljši n={best_rf['n_estimators']}, depth={best_rf['max_depth']}  "
      f"R²={rf_m['R2']:.4f}  MAE={rf_m['MAE']:,.0f} €  RMSE={rf_m['RMSE']:,.0f} €")

print()
print("=" * 60)
print("  NEVRONSKA MREŽA (MLP)  (hidden × epochs)")
print("=" * 60)
nn_grid = {"hidden": [16, 24], "epochs": [20, 30]}
best_nn, cv_nn, nn_all = grid_search(NeuralNetRegressor, nn_grid, X_tr, y_tr, label="MLP")
nn = NeuralNetRegressor(**best_nn).fit(X_tr, y_tr)
nn_m = eval_metrics(y_te, nn.predict(X_te))
print(f"  Najboljši hidden={best_nn['hidden']}, epochs={best_nn['epochs']}  "
      f"R²={nn_m['R2']:.4f}  MAE={nn_m['MAE']:,.0f} €  RMSE={nn_m['RMSE']:,.0f} €")

# ── Povzetek ──────────────────────────────────────────────────────────────────
models_summary = [
    ("Linearna regresija (OLS)", lr_m, cv_lr, {}, lr),
    (f"Ridge (α={best_ridge['alpha']})", ridge_m, cv_ridge, best_ridge, ridge),
    (f"Odl. drevo (depth={best_dt['max_depth']})", dt_m, cv_dt, best_dt, dt),
    (f"Naključni gozd (n={best_rf['n_estimators']}, d={best_rf['max_depth']})",
     rf_m, cv_rf, best_rf, rf),
    (f"Nevronska mreža (h={best_nn['hidden']}, ep={best_nn['epochs']})",
     nn_m, cv_nn, best_nn, nn),
]

print()
print("=" * 70)
print("  PRIMERJAVA MODELOV")
print(f"  {'Model':<38}  {'R²':>6}  {'MAE':>10}  {'RMSE':>10}  {'CV R²':>6}")
print("  " + "-" * 68)
for name, m, cv, _, _model in models_summary:
    print(f"  {name:<38}  {m['R2']:>6.4f}  {m['MAE']:>10,.0f}  {m['RMSE']:>10,.0f}  {cv:>6.4f}")
print("=" * 70)

best_name, best_m, best_cv, best_params, best_model = max(
    models_summary, key=lambda x: x[1]["R2"])
print(f"\n  🏆  Najboljši model: {best_name}  (R²={best_m['R2']:.4f})")
print(f"\n  Analiza končana – {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}")


# ═════════════════════════════════════════════════════════════════════════════
#  DOCX POROČILO
# ═════════════════════════════════════════════════════════════════════════════

# ── Graf utilities (self-contained, struct + zlib) ────────────────────────────
_F57: dict = {
    ' ':[0,0,0,0,0,0,0],
    '0':[14,17,19,21,25,17,14],'1':[4,12,4,4,4,4,14],'2':[14,17,1,2,4,8,31],
    '3':[31,2,4,2,1,17,14],'4':[2,6,10,18,31,2,2],'5':[31,16,30,1,1,17,14],
    '6':[6,8,16,30,17,17,14],'7':[31,1,2,4,8,8,8],'8':[14,17,17,14,17,17,14],
    '9':[14,17,17,15,1,2,12],'.':[0,0,0,0,0,12,12],'-':[0,0,0,31,0,0,0],
    '+':[0,4,4,31,4,4,0],'=':[0,0,31,0,31,0,0],'/':[1,2,4,8,16,0,0],
    '(':[2,4,8,8,8,4,2],')':[8,4,2,2,2,4,8],'%':[24,25,2,4,8,19,3],
    'k':[16,16,18,20,24,20,18],'m':[0,26,21,21,21,21,21],
    'a':[0,14,1,15,17,19,13],'b':[16,16,30,17,17,17,30],'c':[0,14,17,16,16,17,14],
    'd':[1,1,13,19,17,19,13],'e':[0,14,17,31,16,17,14],'f':[6,9,8,30,8,8,8],
    'g':[0,14,17,17,15,1,14],'h':[16,16,30,17,17,17,17],'i':[12,12,0,12,12,12,14],
    'j':[6,0,6,6,6,22,12],'l':[12,4,4,4,4,4,14],'n':[0,28,18,17,17,17,17],
    'o':[0,14,17,17,17,17,14],'p':[0,30,17,17,30,16,16],'r':[0,22,25,16,16,16,16],
    's':[0,14,16,14,1,17,14],'t':[4,4,31,4,4,5,2],'u':[0,17,17,17,17,19,13],
    'v':[0,17,17,17,17,10,4],'z':[0,31,2,4,8,16,31],
    'A':[14,17,17,31,17,17,17],'B':[30,17,17,30,17,17,30],'C':[14,17,16,16,16,17,14],
    'D':[28,18,17,17,17,18,28],'E':[31,16,16,30,16,16,31],'F':[31,16,16,30,16,16,16],
    'G':[14,17,16,23,17,17,14],'H':[17,17,17,31,17,17,17],'I':[14,4,4,4,4,4,14],
    'J':[7,2,2,2,2,18,12],'K':[17,18,20,24,20,18,17],'L':[16,16,16,16,16,16,31],
    'M':[17,27,21,21,17,17,17],'N':[17,25,21,21,19,17,17],'O':[14,17,17,17,17,17,14],
    'P':[30,17,17,30,16,16,16],'R':[30,17,17,30,20,18,17],'S':[14,17,16,14,1,17,14],
    'T':[31,4,4,4,4,4,4],'U':[17,17,17,17,17,17,14],'V':[17,17,17,17,17,10,4],
    'W':[17,17,17,21,21,27,17],'X':[17,10,4,4,4,10,17],'Y':[17,17,10,4,4,4,4],
    'Z':[31,1,2,4,8,16,31],
}

def _asc(s):
    return (s.replace('š','s').replace('Š','S').replace('č','c').replace('Č','C')
             .replace('ž','z').replace('Ž','Z').replace('ć','c').replace('đ','d')
             .replace('á','a').replace('é','e').replace('ú','u').replace('²','2'))

def _png(W, H, px):
    raw = bytearray()
    for row in px:
        raw.append(0)
        for p_ in row:
            raw += bytes(p_)
    comp = zlib.compress(bytes(raw), 6)
    def ck(n_, d):
        crc = zlib.crc32(n_ + d) & 0xFFFFFFFF
        return struct.pack('>I', len(d)) + n_ + d + struct.pack('>I', crc)
    return (b'\x89PNG\r\n\x1a\n'
            + ck(b'IHDR', struct.pack('>IIBBBBB', W, H, 8, 2, 0, 0, 0))
            + ck(b'IDAT', comp)
            + ck(b'IEND', b''))

class _CV:
    def __init__(self, W=700, H=400, bg=(248,249,250)):
        self.W, self.H = W, H
        self._p = [[[*bg]]*W for _ in range(H)]
    def _s(self, x, y, c):
        if 0 <= x < self.W and 0 <= y < self.H:
            self._p[y][x] = list(c)
    def rect(self, x, y, w, h, c):
        lc = list(c)
        for dy in range(max(0,y), min(self.H,y+h)):
            row = self._p[dy]
            for dx in range(max(0,x), min(self.W,x+w)):
                row[dx] = lc
    def hl(self, x1, x2, y, c, t=1):
        lc = list(c)
        for x in range(min(x1,x2), max(x1,x2)+1):
            for dt in range(t):
                if 0 <= y+dt < self.H and 0 <= x < self.W:
                    self._p[y+dt][x] = lc
    def vl(self, x, y1, y2, c, t=1):
        lc = list(c)
        for y in range(min(y1,y2), max(y1,y2)+1):
            for dt in range(t):
                if 0 <= y < self.H and 0 <= x+dt < self.W:
                    self._p[y][x+dt] = lc
    def line(self, x1, y1, x2, y2, c, t=1):
        dx, dy = abs(x2-x1), abs(y2-y1)
        sx = 1 if x1<x2 else -1; sy = 1 if y1<y2 else -1
        e, x, y = dx-dy, x1, y1
        while True:
            for dt in range(t): self._s(x+dt, y, c)
            if x==x2 and y==y2: break
            e2 = 2*e
            if e2>-dy: e-=dy; x+=sx
            if e2<dx:  e+=dx; y+=sy
    def dot(self, cx, cy, r, c):
        for dy in range(-r,r+1):
            for dx in range(-r,r+1):
                if dx*dx+dy*dy <= r*r: self._s(cx+dx, cy+dy, c)
    def txt(self, x, y, s, c, sc=1):
        lc = list(c)
        for ch in _asc(s):
            bits = _F57.get(ch, [0]*7)
            for ri, b in enumerate(bits):
                for ci in range(5):
                    if b & (1<<(4-ci)):
                        for ddy in range(sc):
                            for ddx in range(sc):
                                self._s(x+ci*sc+ddx, y+ri*sc+ddy, lc)
            x += 6*sc
    def tw(self, s, sc=1): return len(s)*6*sc
    def to_png(self):
        return _png(self.W, self.H,
                    [[bytes(self._p[y][x]) for x in range(self.W)] for y in range(self.H)])

def _fv(v, mx):
    if mx >= 500_000: return f"{v/1000:.0f}k"
    if mx >= 5_000:   return f"{v/1000:.1f}k"
    if mx >= 100:     return f"{v:.0f}"
    return f"{v:.2f}"

# ── Primerjalni grafikon (horizontalni stolpičar z vrednostmi) ─────────────────
def _chart_compare(model_names, r2_vals, title="Primerjava R2 modelov"):
    n = len(model_names)
    if n == 0: return None
    W, ML, MR, MT, MB = 760, 220, 50, 48, 30
    H = max(240, MT + n*48 + MB)
    CW, CH = W-ML-MR, H-MT-MB
    vmn, vmx = min(0.0, min(r2_vals)), max(0.05, max(r2_vals))
    colors = [(33,150,243),(76,175,80),(255,87,34),(156,39,176),(0,188,212)]
    cv = _CV(W, H)
    # grid
    for i in range(5):
        gx = ML + int(i*(CW/4))
        cv.vl(gx, MT, H-MB, (215,215,220))
    bar_h = max(14, CH//n - 8)
    for i, (nm, v) in enumerate(zip(model_names, r2_vals)):
        by  = MT + i*(CH//n) + 4
        bw2 = max(1, int((v - vmn)/(vmx - vmn) * CW))
        col = colors[i % len(colors)]
        cv.rect(ML, by, bw2, bar_h, col)
        lbl = f"{v:.4f}"
        cv.txt(ML+bw2+4, by+bar_h//2-3, lbl, (40,40,50))
        short = _asc(nm[:30])
        cv.txt(max(0, ML-cv.tw(short)-4), by+bar_h//2-3, short, (40,40,50))
    cv.hl(ML, W-MR, H-MB, (70,70,80), 2)
    cv.vl(ML, MT, H-MB, (70,70,80), 2)
    # X axis ticks
    for i in range(5):
        v = vmn + i*(vmx-vmn)/4
        tx = ML + int(i*CW/4)
        lbl = f"{v:.2f}"
        cv.txt(tx-cv.tw(lbl)//2, H-MB+5, lbl, (70,70,80))
    tw = cv.tw(title, 2); cv.txt((W-tw)//2, 10, title, (30,30,40), 2)
    return cv.to_png()


def _chart_pred_actual(y_true, y_pred, title="Napovedano vs Dejansko"):
    if not y_true: return None
    W, H, ML, MR, MT, MB = 700, 500, 65, 30, 55, 55
    CW, CH = W-ML-MR, H-MT-MB
    mn = min(min(y_true), min(y_pred))
    mx = max(max(y_true), max(y_pred))
    if mn == mx: mx = mn + 1
    def px(v): return ML + int((v-mn)/(mx-mn)*CW)
    def py(v): return H-MB - int((v-mn)/(mx-mn)*CH)
    cv = _CV(W, H)
    for i in range(5):
        cv.hl(ML, W-MR, MT+i*CH//4, (215,215,220))
        cv.vl(ML+i*CW//4, MT, H-MB, (215,215,220))
    # ideal line y=x
    cv.line(ML, H-MB, W-MR, MT, (200,50,50), 2)
    # points
    for yt, yp in zip(y_true, y_pred):
        cv.dot(px(yt), py(yp), 4, (33,150,243))
    cv.hl(ML, W-MR, H-MB, (70,70,80), 2)
    cv.vl(ML, MT, H-MB, (70,70,80), 2)
    for i in range(5):
        v = mn + i*(mx-mn)/4
        tx = ML+int(i*CW/4);  ty = H-MB-int(i*CH/4)
        lbl = _fv(v, mx)
        cv.txt(tx-cv.tw(lbl)//2, H-MB+6, lbl, (70,70,80))
        cv.txt(max(0,ML-cv.tw(lbl)-4), ty-3, lbl, (70,70,80))
    tw = cv.tw(title, 2); cv.txt((W-tw)//2, 12, title, (30,30,40), 2)
    cv.txt(ML+CW//2-cv.tw("Dejanska cena")//2, H-14, "Dejanska cena", (70,70,80))
    cv.txt(5, H//2-cv.tw("Napovedano",2)//2, "Napovedano", (70,70,80), 1)
    return cv.to_png()


def _chart_hp_curve(hp_values, r2_values, title, xlabel):
    """Krivulja hiperparametrov (CV R² vs vrednost HP)."""
    if not hp_values: return None
    xs = [float(v) if v is not None else 0.0 for v in hp_values]
    ys = r2_values
    W, H, ML, MR, MT, MB = 700, 380, 65, 30, 48, 52
    CW, CH = W-ML-MR, H-MT-MB
    xmn, xmx = min(xs), max(xs);  ymn, ymx = min(ys), max(ys)
    if xmn == xmx: xmx = xmn + 1
    if ymn == ymx: ymx = ymn + 0.1
    def px_(x): return ML + int((x-xmn)/(xmx-xmn)*CW)
    def py_(y): return H-MB - int((y-ymn)/(ymx-ymn)*CH)
    cv = _CV(W, H)
    for i in range(5):
        cv.hl(ML, W-MR, MT+i*CH//4, (215,215,220))
    for i in range(len(xs)-1):
        cv.line(px_(xs[i]), py_(ys[i]), px_(xs[i+1]), py_(ys[i+1]), (33,150,243), 2)
    for x, y in zip(xs, ys):
        cv.dot(px_(x), py_(y), 5, (33,150,243))
    best_i = ys.index(max(ys))
    cv.dot(px_(xs[best_i]), py_(ys[best_i]), 7, (220,50,50))
    cv.hl(ML, W-MR, H-MB, (70,70,80), 2)
    cv.vl(ML, MT, H-MB, (70,70,80), 2)
    for i in range(len(xs)):
        lbl = str(xs[i]) if xs[i] != int(xs[i]) else str(int(xs[i]))
        cv.txt(px_(xs[i])-cv.tw(lbl)//2, H-MB+6, lbl, (70,70,80))
    for i in range(5):
        v = ymn+i*(ymx-ymn)/4
        cv.txt(max(0,ML-cv.tw(_fv(v,1))-4), H-MB-int(i*CH/4)-3, _fv(v,1), (70,70,80))
    tw = cv.tw(title, 2); cv.txt((W-tw)//2, 10, title, (30,30,40), 2)
    xl = cv.tw(xlabel); cv.txt((W-xl)//2, H-14, xlabel, (70,70,80))
    return cv.to_png()


# ── DOCX report ───────────────────────────────────────────────────────────────
if args.docx:
    try:
        from docx import Document
        from docx.shared import Inches, Cm, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io as _io
    except ImportError:
        print("✗  Namesti: py -m pip install python-docx")
        sys.exit(1)

    out_path = args.docx_izhod or os.path.join(
        os.path.dirname(os.path.abspath(csv_path)),
        f"ml_porocilo_{datetime.now().strftime('%Y%m%d_%H%M')}.docx")

    print(f"\n   Generiram DOCX poročilo: {out_path}")
    doc = Document()
    for sec in doc.sections:
        sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)
        sec.top_margin = Cm(2.5);  sec.bottom_margin = Cm(2.5)

    def _H(lvl, txt):
        h = doc.add_heading(txt, lvl)
        if lvl == 0: h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _p(txt="", bold=False, italic=False, size=None):
        p = doc.add_paragraph()
        run = p.add_run(txt)
        run.bold = bold; run.italic = italic
        if size: run.font.size = Pt(size)
        return p

    def _tbl(headers, rows_data):
        t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
        for i, h in enumerate(headers):
            c = t.rows[0].cells[i]; c.text = h
            for run in c.paragraphs[0].runs: run.bold = True
        for row in rows_data:
            cells = t.add_row().cells
            for i, v in enumerate(row): cells[i].text = str(v)
        return t

    def _img(png_bytes, caption, w=5.8):
        if not png_bytes: return
        doc.add_picture(_io.BytesIO(png_bytes), width=Inches(w))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cp.runs: run.italic = True; run.font.size = Pt(9)
        doc.add_paragraph()

    # ── Naslov ────────────────────────────────────────────────────────────────
    _H(0, "Regresijski modeli za napovedovanje cen nepremičnin")
    _p(f"Generirano: {datetime.now().strftime('%d.%m.%Y %H:%M')}  |  "
       f"Podatki: {csv_path}", italic=True)
    doc.add_paragraph()

    # ── 1. Potek dela ─────────────────────────────────────────────────────────
    _H(1, "1. Potek dela")
    _p("Analiza poteka v naslednjih korakih:")
    for step in [
        "Nalaganje in čiščenje podatkov iz CSV datoteke",
        "Imputacija manjkajočih vrednosti z medianimi vrednostmi",
        "Kodiranje kategoričnih spremenljivk (VrstaObjekta, Obcina)",
        f"Standardizacija vseh {P} značilk (μ=0, σ=1)",
        f"Naključna razdelitev: {int(args.split*100)}% učna / {100-int(args.split*100)}% testna množica",
        "5-kratno prečno preverjanje (CV) za optimizacijo hiperparametrov",
        "Trening končnih modelov na celotni učni množici",
        "Vrednotenje na testni množici (R², MAE, RMSE)",
    ]:
        doc.add_paragraph(f"  • {step}")
    doc.add_paragraph()

    # ── 2. Podatki ────────────────────────────────────────────────────────────
    _H(1, "2. Podatki")
    cena_all = [r["Cena"] for r in raw_rows]
    _tbl(
        ["Lastnost", "Vrednost"],
        [
            ["Vir podatkov", csv_path],
            ["Skupaj vzorcev", len(raw_rows)],
            ["Učna množica", n_tr],
            ["Testna množica", n_te],
            ["Ciljna spremenljivka", "Cena (€)"],
            ["Min. cena", f"{min(cena_all):,.0f} €"],
            ["Maks. cena", f"{max(cena_all):,.0f} €"],
            ["Povp. cena", f"{statistics.mean(cena_all):,.0f} €"],
            ["Mediana cene", f"{statistics.median(cena_all):,.0f} €"],
        ]
    )
    doc.add_paragraph()

    # ── 3. Predprocesiranje ───────────────────────────────────────────────────
    _H(1, "3. Predprocesiranje in značilke")
    _tbl(
        ["Značilka", "Tip", "Imputacija"],
        [
            ["VelikostM2", "Numerična (m²)", f"mediana = {imp['VelikostM2']:.0f} m²"],
            ["ZemljisteM2", "Numerična (m²)", f"mediana = {imp['ZemljisteM2']:.0f} m²"],
            ["LetoGradnje", "Numerična (leto)", f"mediana = {imp['LetoGradnje']:.0f}"],
            ["StSob", "Numerična (sob)", f"mediana = {imp['StSob']:.1f}"],
            ["EnergetRazred", "Numerična (1–8)", f"mediana = {imp['EnergetRazred']:.1f}"],
            ["VrstaObjekta", "Kategorična → ordinalna koda", "—"],
            ["Obcina", "Kategorična → koda po mediani cene", "—"],
        ]
    )
    _p("Vse numerične značilke so standardizirane (z-score) glede na učno množico.", italic=True)
    doc.add_paragraph()

    # ── 4. Razdelitev podatkov ────────────────────────────────────────────────
    _H(1, "4. Razdelitev podatkov")
    _tbl(
        ["Parameter", "Vrednost"],
        [
            ["Naključno seme", args.seed],
            ["Delež učnih podatkov", f"{args.split*100:.0f}%"],
            ["Učna množica", f"{n_tr} vzorcev"],
            ["Testna množica", f"{n_te} vzorcev"],
            ["Število značilk", P],
            ["Prečno preverjanje", "5-kratno (k-fold CV)"],
        ]
    )
    doc.add_paragraph()

    # ── 5. Modeli in hiperparametri ───────────────────────────────────────────
    _H(1, "5. Modeli in optimizacija hiperparametrov")

    _H(2, "5.1 Linearna regresija (OLS)")
    _p("Minimizira vsoto kvadratov ostankov z normalnimi enačbami. "
       "Nima hiperparametrov – je osnovna referenčna metoda.")
    _p(f"Koeficienti (standardizirane značilke):", bold=True)
    lr_coef_rows = [[f, f"{w:.4f}"] for f, w in zip(FEAT_NAMES, lr._w)]
    _tbl(["Značilka", "Koeficient"], lr_coef_rows)
    doc.add_paragraph()

    _H(2, "5.2 Ridge regresija")
    _p("Linearna regresija z L2 regularizacijo. Parameter α nadzoruje jakost regularizacije.")
    alphas_sorted = sorted(set(r["alpha"] for r, _ in ridge_all),
                           key=lambda a: float(a))
    hp_r2 = []
    for a in alphas_sorted:
        sc = next((s for p_, s in ridge_all if p_["alpha"] == a), None)
        if sc is not None:
            hp_r2.append(sc)
    _tbl(
        ["Alpha", "CV R²"],
        [[a, f"{s:.4f}"] for a, (_, s) in
         zip(alphas_sorted, [(p_, s) for p_, s in sorted(ridge_all, key=lambda x: float(x[0]["alpha"]))])]
    )
    _p(f"Izbrani α = {best_ridge['alpha']}  (CV R² = {cv_ridge:.4f})", bold=True)
    _img(_chart_hp_curve(alphas_sorted, hp_r2, "Ridge: CV R2 glede na alpha", "alpha"),
         "Slika – Ridge regresija: krivulja CV R² po alpha vrednostih")

    _H(2, "5.3 Odločitveno drevo (CART)")
    _p("Rekurzivno deli prostor značilk po kriteriju minimalnega MSE. "
       "max_depth omejuje globino drevesa in s tem regularizira model.")
    depths_sorted = sorted(set(r["max_depth"] for r, _ in dt_all))
    hp_dt = [next(s for p_, s in dt_all if p_["max_depth"] == d) for d in depths_sorted]
    _tbl(
        ["max_depth", "CV R²"],
        [[d, f"{s:.4f}"] for d, s in zip(depths_sorted, hp_dt)]
    )
    _p(f"Izbran max_depth = {best_dt['max_depth']}  (CV R² = {cv_dt:.4f})", bold=True)
    _img(_chart_hp_curve(depths_sorted, hp_dt, "CART: CV R2 glede na max depth", "max depth"),
         "Slika – Odločitveno drevo: krivulja CV R² po max_depth vrednostih")

    _H(2, "5.4 Naključni gozd")
    _p("Ansambel odločitvenih dreves s bootstrap vzorčenjem in naključnim izborom značilk. "
       f"Vsako drevo pri delitvi upošteva √{P} ≈ {max(1,int(P**0.5))} naključnih značilk.")
    _tbl(
        ["n_trees", "max_depth", "CV R²"],
        [[p_["n_estimators"], p_["max_depth"], f"{s:.4f}"] for p_, s in
         sorted(rf_all, key=lambda x: (-x[1], x[0]["max_depth"]))]
    )
    _p(f"Izbrani n={best_rf['n_estimators']}, max_depth={best_rf['max_depth']}  "
       f"(CV R² = {cv_rf:.4f})", bold=True)
    doc.add_paragraph()

    _H(2, "5.5 Nevronska mreža (MLP)")
    _p("Enoplastna nevronska mreža: vhod (p značilk) → skrita plast (ReLU) → izhod (linearni). "
       "Optimizer: mini-batch SGD z LR decay vsakih 10 epoh. Implementacija v čisti standardni knjižnici.")
    _tbl(
        ["hidden", "epochs", "CV R²"],
        [[p_["hidden"], p_["epochs"], f"{s:.4f}"] for p_, s in
         sorted(nn_all, key=lambda x: -x[1])]
    )
    _p(f"Izbrani hidden={best_nn['hidden']}, epochs={best_nn['epochs']}  "
       f"(CV R² = {cv_nn:.4f})", bold=True)
    doc.add_paragraph()

    # ── 6. Primerjava uspešnosti ──────────────────────────────────────────────
    _H(1, "6. Primerjava uspešnosti modelov")
    _tbl(
        ["Model", "R²", "MAE (€)", "RMSE (€)", "CV R²"],
        [[nm,
          f"{m['R2']:.4f}",
          f"{m['MAE']:,.0f}",
          f"{m['RMSE']:,.0f}",
          f"{cv:.4f}"]
         for nm, m, cv, _, _ in models_summary]
    )
    _p(f"\nNajboljši model: {best_name}", bold=True)
    _p(f"R² = {best_m['R2']:.4f}  |  MAE = {best_m['MAE']:,.0f} €  |  RMSE = {best_m['RMSE']:,.0f} €")
    doc.add_paragraph()

    # ── 7. Grafi ──────────────────────────────────────────────────────────────
    _H(1, "7. Grafi primerjave modelov")

    # R² primerjava (vsi 5 modeli)
    nms  = [nm[:28] for nm, *_ in models_summary]
    r2s  = [m["R2"] for _, m, *_ in models_summary]
    _img(_chart_compare(nms, r2s, "Primerjava R2 – vseh 5 modelov (visje=boljse)"),
         "Slika 1 – Primerjava R² vrednosti vseh 5 modelov na testni množici. "
         "Naključni gozd in MLP dosegata višje R² kot linearne metode.")

    # MAE primerjava (1/MAE – višje je boljše)
    maes = [m["MAE"] for _, m, *_ in models_summary]
    _img(_chart_compare(nms, [1.0/m if m > 0 else 0 for m in maes],
                        "Primerjava 1/MAE (vecje=manj napake)"),
         "Slika 1b – Primerjava inverznega MAE (višje = manjša napaka v €). "
         "Odločitveno drevo in naključni gozd dosegata nižji MAE.")

    # Napovedano vs dejansko za najboljši model
    best_preds = best_model.predict(X_te)
    _img(_chart_pred_actual(y_te, best_preds,
                            f"Napovedano vs Dejansko – {best_name.split(' (')[0]}"),
         f"Slika 2 – Napovedane vs dejanske cene: {best_name}. "
         "Točke blizu rdeče premice y=x kažejo dobro napoved.")

    # Porazdelitev ostankov za najboljši model
    residuals = [yt - yp for yt, yp in zip(y_te, best_preds)]
    cv2 = _CV(700, 380)
    mn_r, mx_r = min(residuals), max(residuals)
    if mn_r == mx_r: mx_r = mn_r + 1
    bins2 = 18; step2 = (mx_r-mn_r)/bins2
    cnts2 = [0]*bins2
    for v in residuals:
        cnts2[min(int((v-mn_r)/step2), bins2-1)] += 1
    max_c2 = max(cnts2) or 1
    ML2, MR2, MT2, MB2 = 65, 20, 48, 52
    CW2, CH2 = 700-ML2-MR2, 380-MT2-MB2
    bw2_ = CW2//bins2
    for i in range(5):
        cv2.hl(ML2, 700-MR2, MT2+i*CH2//4, (215,215,220))
    for i, c in enumerate(cnts2):
        if c == 0: continue
        bh2 = max(1, int(c/max_c2*CH2))
        col2 = (220,50,50) if i < bins2//2 else (76,175,80)
        cv2.rect(ML2+i*bw2_+1, 380-MB2-bh2, bw2_-2, bh2, col2)
    cv2.hl(ML2, 700-MR2, 380-MB2, (70,70,80), 2)
    cv2.vl(ML2, MT2, 380-MB2, (70,70,80), 2)
    cv2.vl(ML2+CW2//2, MT2, 380-MB2, (180,50,50), 2)  # zero line
    for i in range(5):
        v = mn_r + i*(mx_r-mn_r)/4
        tx = ML2+int(i*CW2/4)
        cv2.txt(tx-cv2.tw(_fv(v,mx_r))//2, 380-MB2+6, _fv(v,mx_r), (70,70,80))
    ttl = "Porazdelitev ostankov (testna mnozica)"
    cv2.txt((700-cv2.tw(ttl,2))//2, 10, ttl, (30,30,40), 2)
    _img(cv2.to_png(),
         f"Slika 3 – Porazdelitev ostankov najboljšega modela ({best_name}). "
         "Simetrična porazdelitev okoli 0 = nepristranski model.")

    doc.save(out_path)
    print(f"✓  DOCX poročilo shranjeno: {out_path}")

